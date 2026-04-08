# -*- coding: utf-8 -*-
"""Generate Claude_Meeting_04_2026-04-08.docx"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Claude\Session Summaries\Claude_Meeting_04_2026-04-08.docx"

doc = Document()

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("Claude Meeting 04 — Session Summary", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p = doc.add_paragraph("Date: 2026-04-08  |  Project: Claude Manager  |  Session: Claude S4")
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ── Section helper ────────────────────────────────────────────────────────────
def h1(text):
    doc.add_heading(text, level=1)

def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)

# ── 1. Completed ──────────────────────────────────────────────────────────────
h1("✅ Completed This Session")
bullet("ClaudeManager NSSM service confirmed RUNNING (was marked PAUSED in brief — self-resolved)")
bullet("GitHub remote added to C:\\Claude and pushed to Quiddity-Innovations/CLAUDE-MANAGER (new private repo created)")
bullet("/api/ping test added to test_dashboard_api.py — 10/10 tests passing")
bullet("maia.db explored: 38 tables, 6 bots (Maia=active), 30 people, 903 conversations")
bullet("NEXUS .gitignore updated — API key files, LOGS/, data/ now ignored (security fix)")
bullet("All loose files committed and pushed: Maia (8 files), NEXUS (6 files), OC (11 files), CLAUDE-MANAGER (4 files)")
bullet("Session naming convention saved to all project memory dirs (C--QI, C--OC, C--Users-renne-Downloads)")
bullet("Google Calendar events created: Meeting 04 summary (today, yellow) + Meeting 05 prep (tomorrow, orange)")
bullet("LATEST.md and status.json updated for Meeting 05 handoff")

# ── 2. Issue Discovered ───────────────────────────────────────────────────────
h1("⚠️ Issue Discovered — MCP Worktree Instability")
bullet("MCP tools (openspace, sqlite-maia, sqlite-naya, git) do NOT load reliably in worktree sessions")
bullet("Root cause: Claude Code worktree sessions (e.g. claude/funny-margulis) cannot hold MCP server processes stable")
bullet("Pattern: MCPs appear briefly in system-reminder, then disconnect before ToolSearch resolves their schema")
bullet("Workaround confirmed: MCP tasks deferred to Claude S5 — must open Claude Code from C:\\Claude root (not a worktree)")
bullet("Fix option A: Add MCP config to project-level .claude/settings.json")
bullet("Fix option B: Register MCPs in C:\\Users\\renne\\.claude\\settings.json (global) instead of .claude.json")

# ── 3. Next Up ────────────────────────────────────────────────────────────────
h1("🔄 Next Up — Claude S5: Claude Manager: Fix MCP + First Real Queries")
numbered("Open Claude Code from C:\\Claude root — confirm all 5 MCPs load and stay stable")
numbered("Run OpenSpace search_skills on fetch-ai-news, git-commit, session-summary")
numbered("First real sqlite-maia MCP query: list_tables → read bots, config, conversations")
numbered("First real git MCP use: git_log on C:\\QI via mcp__git__git_log")
numbered("NEXUS installer weekend test prep — target 2026-04-12: wipe DB/cache, run wizard")

# ── 4. In Development ─────────────────────────────────────────────────────────
h1("🚀 In Development")
bullet("MCP ecosystem: 5 servers registered but session-context stability needs hardening (Claude S5)")
bullet("Maia Phase 3: items 7.4, 2.8, 7.5, service restart still pending")
bullet("NEXUS installer Phase 3: proper .exe/.msi after weekend test passes")
bullet("OpenClaw watchdog (oc-watchdog.sh) newly committed — needs testing in WSL")

# ── 5. Files Modified ─────────────────────────────────────────────────────────
h1("📁 Files Modified This Session")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"

# Header row
hdr = table.rows[0].cells
hdr[0].text = "File"
hdr[1].text = "Project"
hdr[2].text = "Action"
for cell in hdr:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'D5E8F0')
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

rows = [
    (r"C:\Claude\Tests\test_dashboard_api.py", "CLAUDE-MANAGER", "Added /api/ping test — 10/10 passing"),
    (r"C:\Claude\status.json", "CLAUDE-MANAGER", "Updated with Meeting 04 results"),
    (r"C:\Claude\Session Summaries\LATEST.md", "CLAUDE-MANAGER", "Updated for Meeting 05 handoff"),
    (r"C:\QI\maia_gradio.py", "Maia", "Committed + pushed"),
    (r"C:\QI\maia_i18n.py", "Maia", "Committed + pushed"),
    (r"C:\QI\maia_server.py", "Maia", "Committed + pushed"),
    (r"C:\QI\knowledge\nexus_curated.jsonl", "Maia", "RAG seed data (8 entries) committed"),
    (r"C:\NEXUS\.gitignore", "NEXUS", "API keys + LOGS/ + data/ ignored"),
    (r"C:\OC\keep-wsl-alive.ps1", "OC", "Updated + pushed"),
    (r"C:\OC\oc-watchdog.sh", "OC", "Added + pushed"),
    ("Memory files (3 project dirs)", "All Projects", "Session naming convention saved"),
]

for file, project, action in rows:
    row = table.add_row().cells
    row[0].text = file
    row[1].text = project
    row[2].text = action

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
