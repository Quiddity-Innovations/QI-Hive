# -*- coding: utf-8 -*-
"""
Claude Manager — Session Documentation Generator
Creates / updates Implementation Log, Meeting Minutes, Version History, and Session Summary .docx
Run: python C:\Claude\scripts\create_session_docs.py
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, RGBColor

DOCS_DIR = Path(r"C:\Claude\DOCUMENTATION")
SUMM_DIR = Path(r"C:\Claude\Session Summaries")
DOCS_DIR.mkdir(exist_ok=True)
SUMM_DIR.mkdir(exist_ok=True)

NOW     = datetime.now()
TODAY   = NOW.strftime("%Y-%m-%d")
DTSTAMP = NOW.strftime("%Y-%m-%d %H%M")

# ── Helpers ───────────────────────────────────────────────────────────────────

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def table_row(table, cells):
    row = table.add_row()
    for i, val in enumerate(cells):
        row.cells[i].text = val
    return row


# ── Session Summary .docx ─────────────────────────────────────────────────────

def create_session_summary():
    path = SUMM_DIR / f"Claude_Meeting_02_extended_{DTSTAMP}.docx"
    doc = Document()

    heading(doc, "QI Claude Manager — Meeting 02 Extended Session Summary")
    para(doc, f"Date: {TODAY}  |  Session: Claude_Meeting_02_extended  |  Generated: {NOW.strftime('%H:%M')}")
    doc.add_paragraph()

    heading(doc, "✅ Completed This Session", 2)
    items = [
        "Tester agent scaffolded: soul.md, skills.md, config.json at C:\\Claude\\Agents\\tester\\",
        "Test stack installed: pytest, httpx, playwright (Chromium), locust, pytest-json-report",
        "C:\\Claude\\Tests\\ built — 6 test files covering all 4 active QI services",
        "test_smoke.py: 4-project ping test, graceful skip if offline",
        "test_dashboard_api.py: full CRUD cycle, all 5 pages, all 5 API endpoints — 8/8 pass",
        "test_maia_api.py, test_naya_api.py, test_nexus_api.py: per-project API tests with skip-if-offline",
        "test_dashboard_ui.py: Playwright headless — sidebar, cards, board columns, guide page",
        "load/locustfile.py: Locust load scenarios for all 4 services",
        "run_tests.py: unified runner — saves JSON results + auto-creates kanban tasks for failures",
        "Dashboard /tests page: Smoke / API / UI / Run All buttons, live results panel",
        "POST /api/tests/run endpoint: executes tests in subprocess, returns JSON summary to browser",
        "Tester agent added to agent_roster in status.json (7 agents total)",
        "agent_icons + Add Task modal updated to include tester",
        "Smoke test ran: 3/4 pass — Naya root 404 confirmed as real bug",
        "Task t49ec78 auto-created on kanban board: 'Fix failing test: Naya root 404'",
        "QI_Claude_Manager_Guide.md updated: Tester agent + Tests page documented",
        "LATEST.md updated with full Meeting 03 agenda and file manifest",
        "Google Calendar events created: today (yellow) + Meeting 03 agenda (orange)",
        "sync_tasks() verified working end-to-end: health check → tasks.json → board",
    ]
    for item in items:
        bullet(doc, item)

    heading(doc, "🔄 Next Up (Meeting 03)", 2)
    nexts = [
        "1. sc start ClaudeManager (admin terminal) — make dashboard persist as NSSM service",
        "2. Fix Naya root 404 — add GET / route to naya_server.py",
        "3. Install SQLite MCP — pip install mcp-sqlite + register in .claude.json",
        "4. Install Git MCP — npm install -g @cyanheads/git-mcp-server + register",
        "5. Agent workflow test — Architect→Builder→Inspector chain on a real small task",
        "6. NEXUS integration — wire Scout to /scout/digest for AI news digest",
        "7. Skill evolution — test OpenSpace skill-discovery against existing QI skill",
    ]
    for n in nexts:
        bullet(doc, n)

    heading(doc, "🚀 In Development", 2)
    for item in [
        "Claude Manager full agent team — 7/7 scaffolded, workflow tests pending",
        "NEXUS AI news Scout integration with dashboard",
        "OpenSpace skill evolution pipeline",
        "Maia RAG knowledge base (ChromaDB)",
        "OpenClaw Koe voice layer (shared across projects)",
    ]:
        bullet(doc, item)

    heading(doc, "⚠️ Known Issues (on kanban board)", 2)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Task"; hdr[1].text = "Project"; hdr[2].text = "Priority"
    data = [
        ("Fix failing test: Naya root 404",     "Naya",          "High"),
        ("Commit loose files in Naya",           "Naya",          "Medium"),
        ("Commit loose files in NEXUS",          "NEXUS",         "Medium"),
    ]
    for row in data:
        table_row(t, row)
    doc.add_paragraph()

    heading(doc, "📁 Files Created / Updated", 2)
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    h2 = t2.rows[0].cells
    h2[0].text = "File"; h2[1].text = "Action"
    files = [
        (r"C:\Claude\Agents\tester\soul.md",          "CREATED"),
        (r"C:\Claude\Agents\tester\skills.md",         "CREATED"),
        (r"C:\Claude\Agents\tester\config.json",       "CREATED"),
        (r"C:\Claude\Tests\conftest.py",               "CREATED"),
        (r"C:\Claude\Tests\test_smoke.py",             "CREATED"),
        (r"C:\Claude\Tests\test_maia_api.py",          "CREATED"),
        (r"C:\Claude\Tests\test_naya_api.py",          "CREATED"),
        (r"C:\Claude\Tests\test_nexus_api.py",         "CREATED"),
        (r"C:\Claude\Tests\test_dashboard_api.py",     "CREATED"),
        (r"C:\Claude\Tests\test_dashboard_ui.py",      "CREATED"),
        (r"C:\Claude\Tests\run_tests.py",              "CREATED"),
        (r"C:\Claude\Tests\load\locustfile.py",        "CREATED"),
        (r"C:\Claude\Dashboard\server.py",             "UPDATED"),
        (r"C:\UNIVERSAL\QI_Claude_Manager_Guide.md",   "UPDATED"),
        (r"C:\Claude\status.json",                     "UPDATED"),
        (r"C:\Claude\Session Summaries\LATEST.md",     "UPDATED"),
    ]
    for f in files:
        table_row(t2, f)

    doc.save(path)
    print(f"✅ Session summary saved: {path}")
    return path


# ── Implementation Log ────────────────────────────────────────────────────────

def update_implementation_log():
    path = DOCS_DIR / "Claude_Manager_Implementation_Log.docx"

    if path.exists():
        doc = Document(path)
    else:
        doc = Document()
        heading(doc, "Claude Manager — Implementation Log")
        para(doc, "Quiddity Innovations | All significant builds and changes")
        doc.add_paragraph()

    # Add new entry
    heading(doc, f"[{TODAY}] Meeting 02 Extended — Tester Agent + Test Suite", 2)

    sections = {
        "Tester Agent": [
            "C:\\Claude\\Agents\\tester\\ — soul.md, skills.md, config.json",
            "Model: Haiku (default), Sonnet (max). Cross-project scope.",
            "Tools: pytest, httpx, playwright, locust",
            "Escalation: critical → ops, dev failures → builder",
        ],
        "Test Infrastructure": [
            "pip install playwright locust pytest httpx pytest-json-report",
            "playwright install chromium",
            "C:\\Claude\\Tests\\ — 6 test files + locustfile.py + run_tests.py",
            "Test targets: Maia :8001, Naya :8002, NEXUS :8010, Dashboard :8600",
            "All API tests skip gracefully if target service is offline",
        ],
        "Dashboard /tests Page": [
            "New nav item: Tests (bi-bug icon)",
            "4 buttons: Smoke, API Tests, UI Tests, Run All",
            "Results panel: pass/fail/skip counts + per-test row detail",
            "POST /api/tests/run — subprocess runner, returns JSON summary",
            "render_tests() reads C:\\Claude\\Tests\\results\\latest.json",
        ],
        "Task Automation": [
            "run_tests.py auto-creates kanban tasks for every failing test",
            "Deduplication: never creates a task if title already on board",
            "Smoke test result: 3/4 pass — Naya root 404 → task t49ec78 created",
        ],
    }

    for section_title, bullets in sections.items():
        para(doc, section_title, bold=True)
        for b in bullets:
            bullet(doc, b)
    doc.add_paragraph()

    doc.save(path)
    print(f"✅ Implementation log updated: {path}")


# ── Meeting Minutes ────────────────────────────────────────────────────────────

def update_meeting_minutes():
    path = DOCS_DIR / "Claude_Manager_Meeting_Minutes.docx"

    if path.exists():
        doc = Document(path)
    else:
        doc = Document()
        heading(doc, "Claude Manager — Meeting Minutes")
        para(doc, "Quiddity Innovations | Session records")
        doc.add_paragraph()

    heading(doc, f"Meeting 02 Extended — {TODAY}", 2)
    para(doc, "Attendees: Renne Santiago, Claude (Agent Manager)")
    para(doc, "Topic: Tester agent design and test infrastructure build")
    doc.add_paragraph()

    heading(doc, "Decisions Made", 3)
    decisions = [
        "Test stack: pytest + httpx (API), Playwright (UI), Locust (load) — all free, all Python",
        "Rejected: Selenium (replaced by Playwright), Cypress (JS-only), TestComplete/Katalon (paid), JMeter (Java), Robot Framework (unnecessary overhead)",
        "Tester has cross-project mandate — only agent allowed to test ALL QI services",
        "Failed tests auto-create kanban tasks — no manual triage needed",
        "Tester model: Haiku default (fast/cheap for test runs), Sonnet max for analysis",
        "C:\\Claude has no git repo yet — to be initialized in Meeting 03",
    ]
    for d in decisions:
        bullet(doc, d)

    heading(doc, "Issues Found", 3)
    issues = [
        "Naya root / returns HTTP 404 — FastAPI server has no root route. Task filed.",
        "Naya and NEXUS have uncommitted git changes — to be committed.",
        "C:\\Claude not yet a git repository.",
    ]
    for i in issues:
        bullet(doc, i)

    heading(doc, "Next Session (Meeting 03)", 3)
    nexts = [
        "sc start ClaudeManager (admin) — persist dashboard as service",
        "Fix Naya root 404",
        "Install SQLite MCP + Git MCP",
        "Run Architect→Builder→Inspector workflow on a real task",
        "Wire Scout to NEXUS /scout/digest",
        "Test OpenSpace skill-discovery",
    ]
    for n in nexts:
        bullet(doc, n)
    doc.add_paragraph()

    doc.save(path)
    print(f"✅ Meeting minutes updated: {path}")


# ── Version History ────────────────────────────────────────────────────────────

def update_version_history():
    path = DOCS_DIR / "Claude_Manager_Version_History.docx"

    if path.exists():
        doc = Document(path)
    else:
        doc = Document()
        heading(doc, "Claude Manager — Version History")
        para(doc, "Quiddity Innovations | Code change log")
        doc.add_paragraph()
        # Add header row for initial table
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        h = t.rows[0].cells
        h[0].text = "Date"; h[1].text = "Version"; h[2].text = "File"; h[3].text = "Change"
        doc.save(path)
        doc = Document(path)

    heading(doc, f"v2.1 — {TODAY} — Tester Agent + /tests Dashboard Page", 2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Date"; hdr[1].text = "File"; hdr[2].text = "Change Type"; hdr[3].text = "Summary"
    changes = [
        (TODAY, "Dashboard/server.py",              "Feature",  "Added /tests page, POST /api/tests/run, tester in agent_icons + modal, nav item"),
        (TODAY, "Tests/test_smoke.py",              "New",      "4-service smoke test with skip-if-offline"),
        (TODAY, "Tests/test_dashboard_api.py",      "New",      "Full CRUD + page tests, 8 assertions"),
        (TODAY, "Tests/test_maia_api.py",           "New",      "Maia API tests (read-only, skip if offline)"),
        (TODAY, "Tests/test_naya_api.py",           "New",      "Naya API tests (read-only, skip if offline)"),
        (TODAY, "Tests/test_nexus_api.py",          "New",      "NEXUS API tests + synthesize/scout endpoint checks"),
        (TODAY, "Tests/test_dashboard_ui.py",       "New",      "Playwright headless UI tests with screenshot on failure"),
        (TODAY, "Tests/run_tests.py",               "New",      "Unified runner: saves JSON + auto-creates kanban tasks"),
        (TODAY, "Tests/load/locustfile.py",         "New",      "Locust load scenarios for all 4 services"),
        (TODAY, "Tests/conftest.py",                "New",      "Shared httpx session fixtures"),
        (TODAY, "Agents/tester/soul.md",            "New",      "Tester agent identity + cross-project test mandate"),
        (TODAY, "Agents/tester/skills.md",          "New",      "Smoke, full API, UI, load, write-test skills"),
        (TODAY, "Agents/tester/config.json",        "New",      "Haiku default, Sonnet max, all test targets registered"),
        (TODAY, "status.json",                      "Updated",  "7 agents, tester in roster, feature registry updated"),
        (TODAY, "QI_Claude_Manager_Guide.md",       "Updated",  "Tester agent row, Tests page section added"),
        (TODAY, "Session Summaries/LATEST.md",      "Updated",  "Meeting 03 agenda, full file manifest"),
    ]
    for row in changes:
        table_row(t, row)
    doc.add_paragraph()

    doc.save(path)
    print(f"✅ Version history updated: {path}")


if __name__ == "__main__":
    print("\n" + "="*60)
    print("  CLAUDE MANAGER — DOCUMENTATION UPDATE")
    print("="*60 + "\n")
    create_session_summary()
    update_implementation_log()
    update_meeting_minutes()
    update_version_history()
    print("\n✅ All documentation complete.\n")
