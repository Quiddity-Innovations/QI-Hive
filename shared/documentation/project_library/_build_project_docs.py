# -*- coding: utf-8 -*-
"""
QI Hive — Project Status Library builder (template v2, 2026-08-06).

For each project: copy the raw INTRO source set (status_*.json/.md + *.svg) into
  project_library\<Name>\source\
and compile a polished Word document
  project_library\<Name>\<Name>_ProjectStatus_<DATE>.docx
mirroring the dashboard Project Status page (Overview, Blueprint, Feature
Status Business, Feature Status Dev, Code Explained, Future Enhancements,
Tech Stack, Docs).

Template v2 — "QI Report" look:
  cover page with accent bar + meta table, contents page, PART chips,
  ruled section headings, accent-header tables with banded rows, callouts,
  running header/footer with page numbers. Palette is the existing QI doc
  scheme (QI blue #0D6EA6 + status colors) — deliberately NOT the BU look.

Reusable for the remaining QI projects — add entries to PROJECTS and re-run.
"""
from __future__ import annotations
import json, os, re, shutil, sys, tempfile
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATE = "2026-08-06"
LIB = Path(r"C:\QIH\shared\documentation\project_library")

# ── QI Report palette (hex, no #) ──────────────────────────────────────────────
ACCENT      = "0D6EA6"   # QI blue (existing library accent)
ACCENT_DARK = "0A5580"
INK         = "1F2937"   # body text
MUTED       = "64748B"
BORDER      = "C9D4DE"
BAND        = "F1F5F9"   # zebra band
CALLOUT_BG  = "E9F2F8"
CODE_BG     = "1E293B"   # slate-900, matches dashboard code block
CODE_FG     = "E2E8F0"

HEAD_FONT = "Segoe UI"
BODY_FONT = "Calibri"
MONO_FONT = "Consolas"

def _rgb(hexs):  # "0D6EA6" -> RGBColor
    return RGBColor(int(hexs[0:2], 16), int(hexs[2:4], 16), int(hexs[4:6], 16))

# pid -> (display_name, intro_dir). Builds only those whose INTRO has status_intro.md
# (skip-if-not-ready guard in __main__), so this can be re-run safely each wave.
PROJECTS = {
    # Batch 1
    "autopdf": ("AutoPDF", Path(r"C:\AutoPDF\INTRO")),
    "mapsnap": ("MapSnap", Path(r"C:\MapSnap\INTRO")),
    "nexus":   ("NEXUS",   Path(r"C:\NEXUS\INTRO")),
    "qi_hive": ("QI Hive", Path(r"C:\QIH\INTRO")),
    # Ecosystem
    "naya":         ("Naya",                 Path(r"C:\NAYA\INTRO")),
    "qi_brain":     ("QI Brain",             Path(r"C:\QIH\engine\brain\INTRO")),
    "easyflow":     ("EasyFlow",             Path(r"C:\EasyFlow\INTRO")),
    "openclaw":     ("OpenClaw",             Path(r"C:\OC\INTRO")),
    "tubescout":    ("TubeScout",            Path(r"C:\TUBESCOUT\INTRO")),
    "cognibase":    ("CogniBase",            Path(r"C:\CogniBase\INTRO")),
    "personalsong": ("PersonalSong Studio",  Path(r"C:\PersonalSong\INTRO")),
    "avatarstudio": ("AvatarStudio",         Path(r"C:\1-AI\APPS\AvatarStudio\INTRO")),
    "akiyascout":   ("AkiyaScout",           Path(r"C:\AkiyaScout\INTRO")),
    "lotterywiz":   ("LotteryWiz",           Path(r"C:\Lottery Wiz\INTRO")),
    "claude_voice": ("Claude Voice",         Path(r"C:\CLAUDE\Claude Voice\INTRO")),
    "claude_manager":("Claude Manager",      Path(r"C:\CLAUDE\INTRO")),
    "mq":           ("MQ",                   Path(r"C:\MQ\INTRO")),
    "m2v":          ("M2V",                  Path(r"C:\M2V\INTRO")),
    "gamez":        ("Gamez",                Path(r"C:\Gamez\INTRO")),
    "filehq":       ("FileHQ",               Path(r"C:\NAYA\filehq\INTRO")),
    "cypherminer":  ("CypherMiner",          Path(r"C:\CypherMiner\INTRO")),
    "digitization": ("Digitization Cost Tool", Path(r"C:\Users\renne\Downloads\DIGITIZATION COSTS\INTRO")),
    # Batch 3 — added 2026-08-02 (component-inventory follow-up)
    "playdeck":     ("PlayDeck",     Path(r"C:\PlayDeck\INTRO")),
    "bakeoff":      ("Bakeoff",      Path(r"C:\QIP\Bakeoff\INTRO")),
    "qi_connector": ("QI Connector", Path(r"C:\QIP\Connector\INTRO")),
}

STATUS_LABEL = {
    "live": "Live", "partial": "Partial", "planned": "Planned",
    "disabled": "Disabled", "pending": "Pending",
}
STATUS_RGB = {
    "live": _rgb("1B7A33"), "partial": _rgb("B87A00"),
    "planned": _rgb("0D6EA6"), "disabled": _rgb("666666"),
    "pending": _rgb("B87A00"),
}

PARTS = [
    ("Overview & Features",          "What the project is, who it serves, and its headline capabilities."),
    ("Architecture & Blueprint",     "System architecture and data-model diagrams."),
    ("Feature Status — Business",    "Capability-level status from the stakeholder perspective."),
    ("Feature Status — Developer",   "Component-level status mapped to files and functions."),
    ("Code Explained",               "Real source excerpts showing which code creates each feature."),
    ("Future Enhancements",          "Planned work, grouped by priority."),
    ("Technology Stack",             "Technologies, roles, licences, and versions in use."),
    ("Documentation & Source Index", "Where every related document and source file lives."),
]

# ══ low-level XML helpers ══════════════════════════════════════════════════════
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(f"w:{k}"), str(v))
    return e

def shade_cell(cell, fill_hex):
    cell._tc.get_or_add_tcPr().append(_el("w:shd", val="clear", fill=fill_hex))

def shade_run(run, fill_hex):
    run._r.get_or_add_rPr().append(_el("w:shd", val="clear", fill=fill_hex))

def para_border(par, edge="bottom", sz=12, color=ACCENT, space=4):
    pPr = par._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr"); pPr.append(pBdr)
    pBdr.append(_el(f"w:{edge}", val="single", sz=sz, space=space, color=color))

def table_borders(t, edges):
    """edges: dict edge->(sz,color) — only listed edges get a border."""
    tblPr = t._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        if edge in edges:
            sz, color = edges[edge]
            borders.append(_el(f"w:{edge}", val="single", sz=sz, color=color))
        else:
            borders.append(_el(f"w:{edge}", val="none", sz=0, color="auto"))
    tblPr.append(borders)

def table_cell_margins(t, top=40, bottom=40, left=90, right=90):
    tblPr = t._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for k, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        mar.append(_el(f"w:{k}", w=v, type="dxa"))
    tblPr.append(mar)

def add_field(par, instr, bold=False, size=8.5, color=MUTED):
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), instr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    fonts = _el("w:rFonts", ascii=BODY_FONT, hAnsi=BODY_FONT); rPr.append(fonts)
    rPr.append(_el("w:sz", val=int(size * 2)))
    rPr.append(_el("w:color", val=color))
    if bold: rPr.append(OxmlElement("w:b"))
    r.append(rPr)
    t = OxmlElement("w:t"); t.text = "1"; r.append(t)
    fld.append(r)
    par._p.append(fld)

# ══ template building blocks ═══════════════════════════════════════════════════
def set_base_styles(doc):
    n = doc.styles["Normal"]
    n.font.name = BODY_FONT; n.font.size = Pt(10.5); n.font.color.rgb = _rgb(INK)
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.12
    specs = {  # style: (size, bold, color, before, after, border_rule)
        "Heading 1": (20, True, INK,        18, 8,  True),
        "Heading 2": (13.5, True, ACCENT_DARK, 14, 5, False),
        "Heading 3": (11,  True, INK,        10, 4, False),
    }
    for name, (size, bold, color, before, after, rule) in specs.items():
        st = doc.styles[name]
        st.font.name = HEAD_FONT; st.font.size = Pt(size)
        st.font.bold = bold; st.font.color.rgb = _rgb(color)
        st.font.italic = False
        # kill the blue underline inheritance from the default theme
        el = st.element.get_or_add_rPr()
        pf = st.paragraph_format
        pf.space_before = Pt(before); pf.space_after = Pt(after)
        pf.keep_with_next = True

def h1(doc, text):
    h = doc.add_heading(text, level=1)
    para_border(h, "bottom", sz=14, color=ACCENT, space=6)
    return h

def part_chip(doc, idx):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"  PART {idx}  ")
    r.font.name = HEAD_FONT; r.font.size = Pt(8.5); r.bold = True
    r.font.color.rgb = _rgb("FFFFFF")
    shade_run(r, ACCENT)
    return p

def accent_bar(doc, color=ACCENT, sz=36):
    """Thick horizontal accent rule (used on the cover)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    para_border(p, "bottom", sz=sz, color=color, space=0)
    return p

def callout(doc, text, fill=CALLOUT_BG, edge_color=ACCENT):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_borders(t, {"left": (24, edge_color)})
    table_cell_margins(t, top=90, bottom=90, left=140, right=140)
    cell = t.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(9.5); r.font.color.rgb = _rgb(INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def style_data_table(t, header=True, size=9.5):
    """Accent header row + banded body + horizontal rules. Call after filling."""
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_borders(t, {"top": (8, ACCENT), "bottom": (8, BORDER), "insideH": (4, BORDER)})
    table_cell_margins(t)
    for i, row in enumerate(t.rows):
        if header and i == 0:
            # repeat the header row when the table breaks across pages
            row._tr.get_or_add_trPr().append(_el("w:tblHeader", val="true"))
        for cell in row.cells:
            if header and i == 0:
                shade_cell(cell, ACCENT)
            elif i % 2 == (0 if header else 1):
                shade_cell(cell, BAND)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                for run in p.runs:
                    if run.font.size is None:
                        run.font.size = Pt(size)
                    if header and i == 0:
                        run.bold = True
                        run.font.color.rgb = _rgb("FFFFFF")
                        run.font.name = HEAD_FONT
                        run.font.size = Pt(size - 0.5)

def setup_page(doc, name, header_text=None):
    sect = doc.sections[0]
    sect.top_margin = Inches(0.9); sect.bottom_margin = Inches(0.8)
    sect.left_margin = Inches(0.95); sect.right_margin = Inches(0.95)
    sect.different_first_page_header_footer = True
    right_edge = sect.page_width - sect.left_margin - sect.right_margin

    hdr = sect.header
    hp = hdr.paragraphs[0]; hp.text = ""
    hp.paragraph_format.tab_stops.add_tab_stop(right_edge, WD_TAB_ALIGNMENT.RIGHT)
    r = hp.add_run(header_text if header_text else f"{name} — Project Status")
    r.font.name = HEAD_FONT; r.font.size = Pt(8.5); r.bold = True
    r.font.color.rgb = _rgb(ACCENT_DARK)
    r2 = hp.add_run("\tQI Hive · Project Documentation Library")
    r2.font.name = BODY_FONT; r2.font.size = Pt(8.5); r2.font.color.rgb = _rgb(MUTED)
    para_border(hp, "bottom", sz=6, color=BORDER, space=4)

    ftr = sect.footer
    fp = ftr.paragraphs[0]; fp.text = ""
    fp.paragraph_format.tab_stops.add_tab_stop(right_edge, WD_TAB_ALIGNMENT.RIGHT)
    para_border(fp, "top", sz=6, color=BORDER, space=4)
    r = fp.add_run(f"Quiddity Innovations · Generated {DATE}")
    r.font.size = Pt(8.5); r.font.color.rgb = _rgb(MUTED)
    fp.add_run("\t").font.size = Pt(8.5)
    r = fp.add_run("Page ")
    r.font.size = Pt(8.5); r.font.color.rgb = _rgb(MUTED)
    add_field(fp, "PAGE", bold=True, size=8.5, color=ACCENT_DARK)
    r = fp.add_run(" of ")
    r.font.size = Pt(8.5); r.font.color.rgb = _rgb(MUTED)
    add_field(fp, "NUMPAGES", size=8.5, color=MUTED)

def cover_page(doc, name, intro: Path):
    accent_bar(doc, sz=48)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Q U I D D I T Y   I N N O V A T I O N S")
    r.font.name = HEAD_FONT; r.font.size = Pt(10); r.bold = True
    r.font.color.rgb = _rgb(MUTED)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    r = p.add_run("QI Hive · Project Documentation Library")
    r.font.name = HEAD_FONT; r.font.size = Pt(10); r.font.color.rgb = _rgb(ACCENT)

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(name)
    r.font.name = HEAD_FONT; r.font.size = Pt(38); r.bold = True
    r.font.color.rgb = _rgb(INK)
    p = doc.add_paragraph()
    r = p.add_run("Project Status Report")
    r.font.name = HEAD_FONT; r.font.size = Pt(17); r.font.color.rgb = _rgb(ACCENT)
    para_border(p, "bottom", sz=10, color=BORDER, space=10)

    for _ in range(3):
        doc.add_paragraph()
    meta = [
        ("Prepared by", "QI Hive — automated documentation pipeline"),
        ("Generated",   DATE),
        ("Source set",  str(intro)),
        ("Classification", "Internal — QI Ecosystem"),
    ]
    t = doc.add_table(rows=0, cols=2)
    for k, v in meta:
        c = t.add_row().cells
        r = c[0].paragraphs[0].add_run(k.upper())
        r.font.name = HEAD_FONT; r.font.size = Pt(8); r.bold = True
        r.font.color.rgb = _rgb(MUTED)
        r = c[1].paragraphs[0].add_run(v)
        r.font.size = Pt(9.5)
        if k == "Source set":
            r.font.name = MONO_FONT; r.font.size = Pt(8.5)
    t.columns[0].width = Inches(1.5); t.columns[1].width = Inches(5.1)
    for row in t.rows:
        row.cells[0].width = Inches(1.5); row.cells[1].width = Inches(5.1)
    table_borders(t, {"insideH": (4, BORDER)})
    table_cell_margins(t, top=60, bottom=60, left=0, right=90)

    doc.add_paragraph(); doc.add_paragraph()
    callout(doc,
        "This report mirrors the live Project Status page in the QI Hive dashboard. "
        "It is generated from the project's INTRO source set; the raw status_*.json / "
        ".md / .svg files are preserved in the accompanying source\\ folder.")
    doc.add_page_break()

def contents_page(doc):
    h1(doc, "Contents")
    t = doc.add_table(rows=0, cols=2)
    for i, (title, desc) in enumerate(PARTS, 1):
        c = t.add_row().cells
        r = c[0].paragraphs[0].add_run(f"{i:02d}")
        r.font.name = HEAD_FONT; r.font.size = Pt(12); r.bold = True
        r.font.color.rgb = _rgb(ACCENT)
        p = c[1].paragraphs[0]
        r = p.add_run(title)
        r.font.name = HEAD_FONT; r.font.size = Pt(11); r.bold = True
        r.font.color.rgb = _rgb(INK)
        p2 = c[1].add_paragraph()
        r = p2.add_run(desc)
        r.font.size = Pt(9); r.font.color.rgb = _rgb(MUTED)
        p2.paragraph_format.space_after = Pt(8)
    t.columns[0].width = Inches(0.55); t.columns[1].width = Inches(6.05)
    for row in t.rows:
        row.cells[0].width = Inches(0.55); row.cells[1].width = Inches(6.05)
    table_borders(t, {"insideH": (4, BORDER)})
    table_cell_margins(t, top=80, bottom=40, left=0, right=60)
    doc.add_page_break()

def part_heading(doc, idx, title):
    part_chip(doc, idx)
    return h1(doc, title)

# ── inline markdown (**bold**, `code`) ─────────────────────────────────────────
def add_inline(par, text):
    for piece in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            r = par.add_run(piece[2:-2]); r.bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            r = par.add_run(piece[1:-1]); r.font.name = MONO_FONT; r.font.size = Pt(9.5)
        else:
            par.add_run(piece)

# ── markdown table block -> docx table ─────────────────────────────────────────
def flush_md_table(doc, rows):
    if not rows:
        return
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [c for c in cells if not all(set(x) <= set("-: ") and x for x in c)]
    if not cells:
        return
    ncol = max(len(c) for c in cells)
    t = doc.add_table(rows=0, cols=ncol)
    for i, row in enumerate(cells):
        tr = t.add_row().cells
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            add_inline(tr[j].paragraphs[0], val)
    style_data_table(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ── status_intro.md -> docx ────────────────────────────────────────────────────
def render_intro_md(doc, md_path: Path, part_idx=None):
    text = md_path.read_text(encoding="utf-8")
    tbl_buf = []
    first_h1_done = False
    for line in text.splitlines():
        s = line.rstrip("\n")
        if s.lstrip().startswith("|"):
            tbl_buf.append(s); continue
        if tbl_buf:
            flush_md_table(doc, tbl_buf); tbl_buf = []
        if s.startswith("# "):
            if part_idx and not first_h1_done:
                part_chip(doc, part_idx)
                first_h1_done = True
            h1(doc, s[2:].strip())
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
        elif s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet"); add_inline(p, s[2:])
        elif s.strip().startswith("*") and s.strip().endswith("*") and not s.strip().startswith("**"):
            p = doc.add_paragraph(); r = p.add_run(s.strip().strip("*")); r.italic = True
            r.font.size = Pt(8.5); r.font.color.rgb = _rgb(MUTED)
        elif s.strip():
            p = doc.add_paragraph(); add_inline(p, s)
    if tbl_buf:
        flush_md_table(doc, tbl_buf)
    if part_idx and not first_h1_done:
        # md had no H1 of its own — still give the section a proper heading
        pass

# ── status badge run ───────────────────────────────────────────────────────────
def status_run(cell, status):
    s = (status or "").lower()
    p = cell.paragraphs[0]
    r = p.add_run(STATUS_LABEL.get(s, status or ""))
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = STATUS_RGB.get(s, _rgb(INK))

def header_row(t, labels):
    cells = t.add_row().cells
    for c, lab in zip(cells, labels):
        c.paragraphs[0].add_run(lab)

def add_code_block(doc, code, language=""):
    """Render a code excerpt as a shaded, monospaced single-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_borders(t, {"left": (24, ACCENT)})
    table_cell_margins(t, top=90, bottom=90, left=140, right=140)
    cell = t.cell(0, 0)
    shade_cell(cell, CODE_BG)
    cell.paragraphs[0].text = ""
    lines = (code or "").split("\n")
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        r.font.name = MONO_FONT; r.font.size = Pt(8.5)
        r.font.color.rgb = _rgb(CODE_FG)
    return t

# ── SVG embed (svglib best-effort, else reference path) ────────────────────────
def embed_svg(doc, svg_path: Path, tmpdir: str):
    try:
        from svglib.svglib import svg2rlg
        from reportlab.graphics import renderPM
        drawing = svg2rlg(str(svg_path))
        png = os.path.join(tmpdir, svg_path.stem + ".png")
        renderPM.drawToFile(drawing, png, fmt="PNG")
        width = min(6.3, (drawing.width or 600) / 96.0)
        doc.add_picture(png, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(); rr = cap.add_run(svg_path.name)
        rr.italic = True; rr.font.size = Pt(8.5); rr.font.color.rgb = _rgb(MUTED)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    except Exception:
        p = doc.add_paragraph()
        r = p.add_run(f"[Diagram: {svg_path.name} — view in the dashboard Blueprint tab or open ")
        r.font.size = Pt(9)
        r2 = p.add_run(str(svg_path)); r2.font.name = MONO_FONT; r2.font.size = Pt(9)
        p.add_run("]").font.size = Pt(9)
        return False

# ── per-project build ──────────────────────────────────────────────────────────
def build(pid, name, intro: Path):
    out_dir = LIB / name
    src_dir = out_dir / "source"
    src_dir.mkdir(parents=True, exist_ok=True)

    # 1) copy raw source set
    copied = []
    for f in sorted(intro.glob("*.json")) + sorted(intro.glob("status_intro.md")) + sorted(intro.glob("*.svg")):
        shutil.copy2(f, src_dir / f.name); copied.append(f.name)

    # 2) load data
    def rj(n):
        p = intro / n
        return json.loads(p.read_text(encoding="utf-8")) if p.exists() else None
    biz = rj("status_features_business.json") or []
    dev = rj("status_features_dev.json") or []
    code_data = rj("status_code.json") or {}
    fut = rj("status_future.json") or {}
    tech = rj("status_techstack.json") or {}
    docs = rj("status_documentation.json") or {}

    doc = Document()
    set_base_styles(doc)
    setup_page(doc, name)

    cover_page(doc, name, intro)
    contents_page(doc)

    # PART 1 — Overview & Features (from status_intro.md)
    render_intro_md(doc, intro / "status_intro.md", part_idx=1)
    doc.add_page_break()

    # PART 2 — Blueprint
    part_heading(doc, 2, "Architecture & Blueprint")
    svgs = sorted(intro.glob("*.svg"))
    with tempfile.TemporaryDirectory() as tmp:
        if not svgs:
            doc.add_paragraph("No architecture diagrams available.")
        for svg in svgs:
            doc.add_heading(svg.stem.replace("_", " ").title(), level=2)
            embed_svg(doc, svg, tmp)
            doc.add_paragraph()
        doc.add_page_break()

        # PART 3 — Feature Status (Business)
        part_heading(doc, 3, "Feature Status — Business")
        for cat in biz:
            doc.add_heading(cat.get("category", ""), level=2)
            t = doc.add_table(rows=0, cols=3)
            header_row(t, ["Capability", "What it does", "Status"])
            for f in cat.get("features", []):
                c = t.add_row().cells
                c[0].paragraphs[0].add_run(f.get("name", "")).bold = True
                desc = f.get("description", "")
                if f.get("notes"):
                    desc += f"\n({f['notes']})"
                c[1].text = desc
                status_run(c[2], f.get("status"))
            style_data_table(t)
            doc.add_paragraph()
        doc.add_page_break()

        # PART 4 — Feature Status (Dev)
        part_heading(doc, 4, "Feature Status — Developer / Technical")
        for cat in dev:
            doc.add_heading(cat.get("category", ""), level=2)
            t = doc.add_table(rows=0, cols=4)
            header_row(t, ["Component", "File / Function", "Status", "Detail"])
            for f in cat.get("features", []):
                c = t.add_row().cells
                c[0].paragraphs[0].add_run(f.get("name", "")).bold = True
                fr = c[1].paragraphs[0].add_run(f.get("file", "")); fr.font.name = MONO_FONT; fr.font.size = Pt(8.5)
                status_run(c[2], f.get("status"))
                c[3].text = f.get("detail", "")
            style_data_table(t)
            doc.add_paragraph()
        doc.add_page_break()

        # PART 5 — Code Explained (code -> feature)
        part_heading(doc, 5, "Code Explained — Code → Feature")
        csecs = code_data.get("sections", [])
        callout(doc, code_data.get("intro")
                or "Real source excerpts showing which part of the code creates each feature.")
        if not csecs:
            doc.add_paragraph("No annotated code snippets available for this project.")
        for sec in csecs:
            doc.add_heading(sec.get("category", ""), level=2)
            for sn in sec.get("snippets", []):
                p = doc.add_paragraph()
                p.paragraph_format.keep_with_next = True
                p.add_run(sn.get("title", "")).bold = True
                if sn.get("feature"):
                    rr = p.add_run("   implements: " + sn["feature"])
                    rr.font.size = Pt(9); rr.font.color.rgb = STATUS_RGB["live"]
                if sn.get("file"):
                    fp = doc.add_paragraph()
                    fp.paragraph_format.keep_with_next = True
                    fr = fp.add_run(sn["file"]); fr.font.name = MONO_FONT
                    fr.font.size = Pt(8.5); fr.font.color.rgb = _rgb(MUTED)
                add_code_block(doc, sn.get("code", ""), sn.get("language", ""))
                ep = doc.add_paragraph(); ep.add_run(sn.get("explanation", "")).font.size = Pt(9.5)
                doc.add_paragraph()
        doc.add_page_break()

        # PART 6 — Future Enhancements
        part_heading(doc, 6, "Future Enhancements")
        for cat in fut.get("categories", []):
            h = doc.add_heading(level=2)
            h.add_run(cat.get("name", ""))
            pr = (cat.get("priority", "") or "").upper()
            if pr:
                rr = h.add_run(f"   [{pr} PRIORITY]"); rr.font.size = Pt(9.5)
                rr.font.color.rgb = {"HIGH": _rgb("C02020"),
                                     "MEDIUM": _rgb("B87A00"),
                                     "LOW": _rgb("666666")}.get(pr, _rgb(INK))
            for item in cat.get("items", []):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(item.get("title", "")).bold = True
                if item.get("detail"):
                    p.add_run(" — " + item["detail"])
        doc.add_page_break()

        # PART 7 — Technology Stack
        part_heading(doc, 7, "Technology Stack")
        t = doc.add_table(rows=0, cols=5)
        header_row(t, ["Layer", "Technology", "Role", "License", "Version"])
        last = None
        for row in tech.get("table", []):
            c = t.add_row().cells
            layer = row.get("layer", "")
            c[0].paragraphs[0].add_run("" if layer == last else layer).bold = True
            last = layer
            c[1].paragraphs[0].add_run(row.get("technology", "")).bold = True
            c[2].text = row.get("role", "")
            c[3].text = row.get("license", "")
            c[4].text = row.get("version", "")
        style_data_table(t)
        doc.add_paragraph()
        descs = tech.get("descriptions", [])
        if descs:
            doc.add_heading("Technology Deep Dives", level=2)
            for d in descs:
                ttl = d.get("title") or d.get("technology") or ""
                body = d.get("body") or d.get("text") or ""
                doc.add_heading(ttl, level=3)
                doc.add_paragraph(body)
        doc.add_page_break()

        # PART 8 — Documentation index
        part_heading(doc, 8, "Documentation & Source Index")
        for sec in docs.get("sections", []):
            doc.add_heading(sec.get("name", ""), level=2)
            for d in sec.get("documents", []):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(d.get("title", "")).bold = True
                typ = d.get("type", "")
                if typ:
                    rr = p.add_run(f"  [{typ}]"); rr.font.size = Pt(8.5); rr.font.color.rgb = _rgb(MUTED)
                loc = (d.get("location", "") or "") + (d.get("file", "") or "")
                if loc:
                    p2 = doc.add_paragraph(); p2.paragraph_format.left_indent = Inches(0.5)
                    rr = p2.add_run(loc); rr.font.name = MONO_FONT; rr.font.size = Pt(8.5)
                if d.get("description"):
                    p3 = doc.add_paragraph(); p3.paragraph_format.left_indent = Inches(0.5)
                    p3.add_run(d["description"]).font.size = Pt(9)

        out_path = out_dir / f"{name.replace(' ', '_')}_ProjectStatus_{DATE}.docx"
        doc.save(str(out_path))
    return out_path, copied, {
        "biz_cats": len(biz), "biz_feats": sum(len(c.get("features", [])) for c in biz),
        "dev_cats": len(dev), "dev_feats": sum(len(c.get("features", [])) for c in dev),
        "tech_rows": len(tech.get("table", [])), "svgs": len(svgs),
        "doc_count": sum(len(s.get("documents", [])) for s in docs.get("sections", [])),
    }


if __name__ == "__main__":
    LIB.mkdir(parents=True, exist_ok=True)
    only = set(a.lower() for a in sys.argv[1:])
    for pid, (name, intro) in PROJECTS.items():
        if only and pid not in only:
            continue
        if not (intro / "status_intro.md").exists():
            print(f"SKIP {name}: not ready ({intro}\\status_intro.md missing)")
            continue
        try:
            out, copied, stats = build(pid, name, intro)
            print(f"OK  {name}: {out}")
            print(f"    source files copied: {len(copied)}  | {stats}")
        except Exception as e:
            import traceback
            print(f"FAIL {name}: {e}")
            traceback.print_exc()
