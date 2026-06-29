# -*- coding: utf-8 -*-
"""Build the CogniBase library in BOTH editions from one canonical markdown source.
  - Renne edition  : per-doc .docx + combined .docx + combined .pdf   (CogniBase_BU_Library_v2.*)
  - QI edition     : combined .docx + combined .pdf                    (CogniBase_BU_Library_v2_qi.*)
Editions differ only by author byline + a few QI phrasings, applied as render-time substitutions.
"""
import os, re, glob, sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

WS   = r"C:\Users\renne\Downloads\CogniBase_BU_Library_v2"
DOCS = os.path.join(WS, "docs")
BU_RED = RGBColor(0xCC, 0x00, 0x00); GREY = RGBColor(0x60, 0x60, 0x60)
DATE = "2026-06-28"

EDITIONS = [
    dict(suffix='',    author='Renne Santiago',       label='Author', perdoc=True,  subs=[]),
    dict(suffix='_qi', author='Quiddity Innovations',  label='By',     perdoc=False, subs=[
        ('Renne Santiago', 'Quiddity Innovations'),
        ('the product-family backbone', 'the QI product-family backbone'),
        ('Ops/JSON (standard health endpoints)', 'Ops/JSON (QI standard)'),
    ]),
]

def apply_subs(text, subs):
    for a, b in subs: text = text.replace(a, b)
    return text

def add_inline(p, text):
    for seg in re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text):
        if not seg: continue
        if seg.startswith('**') and seg.endswith('**'):
            r = p.add_run(seg[2:-2]); r.bold = True
        elif seg.startswith('`') and seg.endswith('`'):
            r = p.add_run(seg[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
        elif seg.startswith('*') and seg.endswith('*'):
            r = p.add_run(seg[1:-1]); r.italic = True
        else:
            p.add_run(seg)

def shade_header(table):
    for cell in table.rows[0].cells:
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.bold = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'CC0000'); tcPr.append(shd)
    for r in table.rows[1:]:
        for cell in r.cells:
            for par in cell.paragraphs:
                for run in par.runs: run.font.size = Pt(9)

def render_text(doc, md_text):
    lines = md_text.splitlines()
    i = 0; in_code = False; code_buf = []
    while i < len(lines):
        ln = lines[i]
        if ln.strip().startswith('```'):
            if in_code:
                p = doc.add_paragraph(); r = p.add_run('\n'.join(code_buf)); r.font.name='Consolas'; r.font.size=Pt(8.5)
                code_buf = []; in_code = False
            else: in_code = True
            i += 1; continue
        if in_code: code_buf.append(ln); i += 1; continue
        if re.match(r'^\|', ln) and i+1 < len(lines) and re.match(r'^\|[\s:|-]+\|', lines[i+1]):
            header = [c.strip() for c in ln.strip('|').split('|')]
            i += 2; rows = []
            while i < len(lines) and lines[i].startswith('|'):
                rows.append([c.strip() for c in lines[i].strip('|').split('|')]); i += 1
            t = doc.add_table(rows=1, cols=len(header)); t.style = 'Light Grid Accent 1'
            for j, h in enumerate(header):
                cell = t.rows[0].cells[j]; cell.text=''; add_inline(cell.paragraphs[0], h)
            for row in rows:
                cells = t.add_row().cells
                for j, c in enumerate(row):
                    if j < len(cells): cells[j].text=''; add_inline(cells[j].paragraphs[0], c)
            shade_header(t); doc.add_paragraph(); continue
        if ln.startswith('### '):
            h = doc.add_heading(ln[4:], 3)
            for r in h.runs: r.font.color.rgb = BU_RED
        elif ln.startswith('## '):
            h = doc.add_heading(ln[3:], 2)
            for r in h.runs: r.font.color.rgb = BU_RED
        elif ln.startswith('# '):
            p = doc.add_paragraph(); r = p.add_run(ln[2:]); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=BU_RED
        elif ln.startswith('- '):
            add_inline(doc.add_paragraph(style='List Bullet'), ln[2:])
        elif ln.startswith('> '):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3); add_inline(p, ln[2:])
            for r in p.runs: r.italic = True; r.font.color.rgb = GREY
        elif ln.strip() == '---':
            pass
        elif ln.strip():
            add_inline(doc.add_paragraph(), ln)
        i += 1

def base_styles(doc, author):
    n = doc.styles['Normal']; n.font.name = 'Calibri'; n.font.size = Pt(10.5)
    cp = doc.core_properties
    cp.author = author; cp.last_modified_by = author
    cp.title = 'CogniBase — BU-Aligned Document Library'; cp.category = f'Authored by {author}'

def build_edition(md_files, ed):
    print(f"== edition {ed['suffix'] or '(renne)'} : author={ed['author']} ==")
    if ed['perdoc']:
        for md in md_files:
            text = apply_subs(open(md, encoding='utf-8').read(), ed['subs'])
            d = Document(); base_styles(d, ed['author']); render_text(d, text); d.save(md[:-3] + '.docx')
        print(f"  wrote {len(md_files)} per-doc docx")
    C = Document(); base_styles(C, ed['author'])
    appx = [m for m in md_files if 'Appendix' in os.path.basename(m)]
    docs_label = str(len(md_files) - len(appx)) + (f" + {len(appx)} {'appendix' if len(appx)==1 else 'appendices'}" if appx else "")
    t = C.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('CogniBase'); r.font.size = Pt(34); r.bold = True; r.font.color.rgb = BU_RED
    s = C.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.add_run('BU-Aligned Document Library').font.size = Pt(18)
    for label, val in [('Prepared for', 'Boston University IS&T — AI & Data Engineering'),
                       (ed['label'], ed['author']), ('Date', DATE),
                       ('Edition', 'v2 · Frame A (BU-vision-led)'),
                       ('Documents', docs_label)]:
        p = C.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        a = p.add_run(f'{label}:  '); a.bold = True; b = p.add_run(val); b.font.color.rgb = GREY
    C.add_page_break()
    h = C.add_heading('Contents', 1)
    for r in h.runs: r.font.color.rgb = BU_RED
    for md in md_files:
        title = apply_subs(open(md, encoding='utf-8').readline().lstrip('# ').strip(), ed['subs'])
        C.add_paragraph(title, style='List Number')
    C.add_page_break()
    for k, md in enumerate(md_files):
        render_text(C, apply_subs(open(md, encoding='utf-8').read(), ed['subs']))
        if k != len(md_files) - 1: C.add_page_break()
    combined = os.path.join(WS, f"CogniBase_BU_Library_v2{ed['suffix']}_combined.docx")
    C.save(combined); print('  combined docx:', os.path.basename(combined))
    pdf = os.path.join(WS, f"CogniBase_BU_Library_v2{ed['suffix']}.pdf")
    try:
        import win32com.client as w
        word = w.Dispatch('Word.Application'); word.Visible = False
        doc = word.Documents.Open(combined); doc.SaveAs(pdf, FileFormat=17); doc.Close(); word.Quit()
        print('  PDF:', os.path.basename(pdf), os.path.getsize(pdf), 'bytes')
    except Exception as e:
        print('  PDF export failed:', e)

def build():
    md_files = sorted(glob.glob(os.path.join(DOCS, '[0-9][0-9]_*.md')))
    print('found', len(md_files), 'markdown docs')
    for ed in EDITIONS:
        build_edition(md_files, ed)

if __name__ == '__main__':
    build()
