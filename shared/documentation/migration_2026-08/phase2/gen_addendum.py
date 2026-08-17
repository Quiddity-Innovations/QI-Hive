"""Addendum to the 2026-08-09 migration summary - the straggler hunt."""
import sys
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

OUT = (r"C:\QIH\shared\documentation\session_summaries"
       r"\QIH_Migration_Addendum_2026-08-09_1100.docx")

d = Document()
d.add_heading("QI Migration - Addendum: releasing C:\\1-AI", 0)
p = d.add_paragraph()
p.add_run("Date: ").bold = True
p.add_run("2026-08-09, after the MCP repoint\n")
p.add_run("Trigger: ").bold = True
p.add_run("p3e_rename_1ai.ps1 refused to rename C:\\1-AI - its safety gate found "
          "13 processes still mapping the old interpreter. That gate did its "
          "job; a forced rename would have failed messily or half-succeeded.")

d.add_heading("What the 13 processes turned out to be", 1)
t = d.add_table(rows=1, cols=3)
t.style = "Light Grid Accent 1"
h = t.rows[0].cells
h[0].text = "Group"
h[1].text = "Root cause"
h[2].text = "Fix"
for a, b, c in [
    ("6 Claude Voice helpers (bridge_responder x3, realtime, session_watch, "
     "meeting_server, voice_tray, voice_button)",
     "Detached grandchildren of the QI_ClaudeVoice* services, started "
     "2026-08-08 20:09 - hours before the service restart. NSSM stopped the "
     "parent; these survived untracked.",
     "Killed directly, services bounced."),
    ("The reason they respawned on the OLD interpreter",
     "session_hook.py correctly uses sys.executable - it inherits whatever ran "
     "it. The hardcoding was in the Claude Code hooks: 12 .claude/settings.json "
     "files invoked C:\\1-AI\\APPS\\PYTHON\\python.exe directly, including the "
     "user-level one with 8 occurrences. The dead path propagated from hook to "
     "supervisor to every helper.",
     "159 files rewritten across 15 projects (settings.json hooks, Claude Voice "
     ".bat launchers, install scripts). 0 left stale."),
    ("oc-keepalive-daemon.py",
     "Child of OC-Keepalive-Service - an NSSM service that does NOT use the "
     "QI_ prefix and runs its own C:\\OC\\nssm.exe. Every scan in this "
     "migration, and the previous session's inventory, enumerated 'QI_*' only, "
     "so it was invisible throughout.",
     "Repointed. Required a direct registry write: the standardised nssm binary "
     "returns 'OpenService(): Access is denied' against a service registered by "
     "a different nssm, and failed SILENTLY - the first attempt reported success "
     "while the registry kept the old value."),
    ("4 MCP servers (qi_registry_mcp x3, brain/mcp.py)",
     "Children of claude.exe. Editing ~/.claude.json does NOT restart Claude "
     "Code - an incorrect assumption in the earlier handoff. The config on disk "
     "is correct (0 occurrences of 1-AI); the running processes hold the old "
     "one in memory.",
     "Requires quitting Claude Code completely. This is the only remaining "
     "blocker."),
]:
    cells = t.add_row().cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c

d.add_heading("Additional services found outside the QI_ convention", 1)
d.add_paragraph(
    "Searching for NSSM-hosted services rather than for the QI_ prefix found "
    "four that the convention misses entirely:")
t2 = d.add_table(rows=1, cols=3)
t2.style = "Light Grid Accent 1"
hh = t2.rows[0].cells
hh[0].text = "Service"
hh[1].text = "State"
hh[2].text = "Note"
for a, b, c in [
    ("OC-Keepalive-Service", "Running",
     "Was still on the old interpreter. Repointed. Uses C:\\OC\\nssm.exe."),
    ("ClaudeManager", "Stopped",
     "Was still on the old interpreter. Repointed."),
    ("NEXUSTunnel", "Running",
     "DUPLICATE of QI_NEXUSTunnel - both registered, both target port 7880. "
     "Leftover from the QI_ renaming, never removed."),
    ("NayaTunnel", "Stopped",
     "DUPLICATE of QI_NayaTunnel. Same leftover pattern."),
]:
    cells = t2.add_row().cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c

d.add_paragraph()
d.add_paragraph(
    "All four violate the QI_ prefix rule in CLAUDE.md and are absent from "
    "QI_Service_Registry.md. Renaming an NSSM service means remove + reinstall, "
    "so that is left as Renne's decision rather than done automatically. The "
    "duplicate tunnels are the more urgent of the two issues.")

d.add_heading("State now", 1)
for t3 in [
    "Processes pinning C:\\1-AI: 4, all of them MCP servers under claude.exe.",
    "Services: 49 running, 7 stopped across QI_* and the four non-QI services. "
    "The QI_ estate is unchanged at 47 running / 5 stopped.",
    "~/.claude.json verified: 0 occurrences of 1-AI, duplicates removed.",
    "159 source/config files repointed, 0 stale.",
]:
    d.add_paragraph(t3, style="List Bullet")

d.add_heading("The one remaining step", 1)
d.add_paragraph(
    "Quit Claude Code completely - every window, and confirm no claude.exe "
    "remains in Task Manager. Then, from a plain elevated PowerShell:")
d.add_paragraph(
    'gsudo powershell -NoProfile -ExecutionPolicy Bypass -File '
    '"C:/QIH/shared/documentation/migration_2026-08/phase2/p3e_rename_1ai.ps1" '
    '-Apply', style="Intense Quote")
d.add_paragraph(
    "It gates on nothing running from C:\\1-AI, renames the tree to "
    "C:\\1-AI.RETIRED_2026-08-09, recreates ONLY C:\\1-AI\\APPS\\PYTHON as a "
    "junction to C:\\Program Files\\Python311 so the remaining venvs keep "
    "working, then verifies four venvs, the whole service estate and four "
    "health endpoints. Renne deletes the retired tree by hand once satisfied.")

d.add_heading("Still outstanding after that", 1)
for t4 in [
    "Reinstall VSCode and LM Studio; then drop the last C:\\1-AI\\APPS\\VSCode "
    "PATH entry, deliberately left so 'code' keeps working until then.",
    "Remove the duplicate NEXUSTunnel / NayaTunnel services.",
    "Dedupe the Brain's project_state: 53 rows for 32 projects.",
    "Reconcile fidelityanalyzer (in Brain, not in qi_registry.json) and "
    "MailBrain (on disk, in neither).",
    "Phase 4, starting with AkiyaScout as the rehearsal move.",
]:
    d.add_paragraph(t4, style="List Bullet")

d.save(OUT)
print("Saved: " + OUT)
