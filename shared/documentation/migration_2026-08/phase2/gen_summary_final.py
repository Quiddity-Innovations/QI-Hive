"""Generate the final 2026-08-09 QIH migration session summary .docx."""
import sys
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

OUT = (r"C:\QIH\shared\documentation\session_summaries"
       r"\QIH_Migration_Summary_2026-08-09_0930.docx")

d = Document()
d.add_heading("QI Migration - Phases 2-5 - Session Summary", 0)

p = d.add_paragraph()
p.add_run("Date: ").bold = True
p.add_run("2026-08-09\n")
p.add_run("Model: ").bold = True
p.add_run("Claude Opus 5\n")
p.add_run("Outcome: ").bold = True
p.add_run("PHASE 2 COMPLETE. Phase 3 substantially done (AvatarStudio moved, "
          "all environment references cleaned); the final deletion of C:\\1-AI "
          "is blocked only on the MCP repoint that ends the session. Phase 5 "
          "plan delivered. Phase 4 planned, not executed. Machine healthy "
          "throughout: 47 services running, 5 stopped - exact parity with the "
          "pre-migration baseline.")

# ---------------------------------------------------------------- phase 2
d.add_heading("Phase 2 - Python relocation: COMPLETE", 1)

d.add_heading("The move itself", 2)
for t in [
    "Restore point #143 'QI migration Phase2.1 pre-PythonMove' created first.",
    "robocopy C:\\1-AI\\APPS\\PYTHON -> C:\\Program Files\\Python311: 110,763 "
    "files, 4.305 GB, 0 failures. venvs/ and testenv/ excluded.",
    "Per-user bundle {1da2e09b-199c-4def-9a99-93a8c1b8ddf2} uninstalled, then "
    "python-3.11.9-amd64.exe installed with InstallAllUsers=1. Both returned "
    "3010 (success, reboot advised). All nine MSI components now registered "
    "PerMachine. THE MACHINE WAS NOT REBOOTED.",
    "Verified zero pending file-rename operations touch Python or C:\\1-AI - "
    "the 150 queued entries are unrelated Firefox/OneDrive/Edge leftovers.",
    "New interpreter: Python 3.11.9, sys.prefix correct, 999 site-packages, "
    "13/13 import smoke test passes, no C:\\1-AI on sys.path.",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("Everything repointed", 2)
t = d.add_table(rows=1, cols=3)
t.style = "Light Grid Accent 1"
h = t.rows[0].cells
h[0].text = "What"
h[1].text = "Count"
h[2].text = "Verification"
for a, b, c in [
    ("NSSM services", "31", "All restarted sequentially. 28 returned to Running "
     "and were confirmed executing the NEW interpreter; 3 left stopped because "
     "they were stopped before the migration. 0 failures."),
    ("Scheduled tasks", "12", "0 stale tasks remain machine-wide. Eight hid the "
     "path inside conhost --headless arguments."),
    ("Console-script launchers", "149", "Each self-verified with zipfile after "
     "rewrite; 0 failures. pip, uvicorn, chroma, mcp-server-sqlite all execute."),
    ("Registry values", "12", "HKLM and HKCU PythonCore corrected; py -0p now "
     "resolves to C:\\Program Files\\Python311. Keys exported to .reg first."),
    ("Environment variables", "8", "HF_HOME, HUGGINGFACE_HUB_CACHE, "
     "TRANSFORMERS_CACHE, TORCH_HOME, DIFFUSERS_CACHE, PIP_CACHE_DIR plus both "
     "PATHs. Only the VSCode PATH entry remains, deliberately."),
    ("Service-backing venvs", "2", "CogniBase 117/117 packages, service Running, "
     "health HTTP 200. Headroom 147/147 with torch 2.13.0+cpu, Running, "
     "listening on 9020."),
]:
    cells = t.add_row().cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c

# ---------------------------------------------------------------- findings
d.add_heading("Problems found that no handoff mentioned", 1)
d.add_paragraph(
    "Each of these would have caused a silent breakage after C:\\1-AI was "
    "deleted. All were caught before any destructive step.")

t2 = d.add_table(rows=1, cols=2)
t2.style = "Light Grid Accent 1"
h2 = t2.rows[0].cells
h2[0].text = "Finding"
h2[1].text = "Why it mattered"
for a, b in [
    ("14 stale venvs, not 1",
     "The handoff named only C:\\CogniBase\\.venv. C:\\CLAUDE\\Tools\\"
     "headroom_env is what QI_Headroom actually runs - invisible to a scan of "
     "service Application paths, because the stale interpreter is one level "
     "down inside the venv."),
    ("149 of 152 console-script .exe launchers had the interpreter path baked "
     "into the binary",
     "A file copy does not fix these. mcp-server-sqlite.exe, used by two MCP "
     "configs, was among them."),
    ("The obvious shebang rewrite corrupts those launchers",
     "Scanning forward for b'#!' hits a false positive inside the t64 stub's "
     "own error-string table, which contains the literal text \"Expected to "
     "find '!' following '#' in shebang line\". The result died with "
     "0xC0000005. Caught by testing on a scratch copy. Fixed by anchoring on "
     "the PK zip magic and scanning backwards."),
    ("Six more environment variables pointed into C:\\1-AI",
     "Including BOTH Machine and User PATH. Deleting C:\\1-AI would have left "
     "every new shell resolving 'python' to a dead path. PIP_CACHE_DIR also "
     "explained the mysterious 593 MB under models\\huggingface."),
    ("Three scheduled tasks were already broken before this migration",
     "QI_NightlyBackup pointed at C:\\UNIVERSAL\\qi_brain\\tools\\backup.py - "
     "the file moved to C:\\QIH\\engine\\brain\\tools\\backup.py during the "
     "UNIVERSAL->QIH migration and the task was never updated. MaiaReconcile "
     "and MaiaRevertMiMo both invoked C:\\QI\\.venv, which does not exist. All "
     "three repaired."),
    ("Non-elevated scheduled-task enumeration is incomplete",
     "QI_NightlyBackup and Maia\\Maia Tunnel Watchdog are invisible to a "
     "non-admin Get-ScheduledTask. An early non-elevated scan reported them as "
     "absent; the elevated pass found and fixed them. The handoff's count of "
     "12 was right."),
    ("The Brain holds 53 project rows for 32 projects",
     "21 duplicate pairs: each project stored once as a full record under its "
     "snake_case id and once as a status-only stub under its display name. "
     "This is why AvatarStudio appeared to be missing from the dev view - the "
     "stub row carries no phase. Renne asked for this to be fixed as a "
     "separate pass after the migration."),
]:
    c = t2.add_row().cells
    c[0].text = a
    c[1].text = b

# ---------------------------------------------------------------- phase 3
d.add_heading("Phase 3 - retire C:\\1-AI: substantially done", 1)
for t in [
    "AvatarStudio copied to C:\\APPS\\AvatarStudio: 5.8 GB, 50,523 files, exact "
    "file-count match, venv functional. Original retained until deletion.",
    "qi_registry.json avatarstudio.path and paths.logs updated to C:\\APPS\\"
    "AvatarStudio and verified against disk. Six doc/config references rewritten.",
    "All HuggingFace / torch / diffusers / pip cache variables repointed to "
    "D:\\AI, which has 1.19 TB free.",
    "Both PATHs cleaned; Python entries rewritten rather than dropped.",
    "Model asset evaluation (at Renne's request): the ONLY model file over 1 MB "
    "anywhere in C:\\1-AI is nomic-embed-text-v1.5.Q4_K_M.gguf (80 MB), bundled "
    "inside LM Studio's own webpack resources. It is a GGUF text-embedding "
    "model; ComfyUI has no node that consumes one, and reinstalling LM Studio "
    "restores it. NOTHING to move to ComfyUI. The 593 MB under "
    "models\\huggingface is a regenerable pip HTTP cache - discard.",
    "D:\\Review: cleaned up by Renne. The remaining folder and D:\\Review itself "
    "stay in place for future cleanup. Not touched by migration automation.",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("What still blocks the deletion", 2)
for t in [
    "17 processes still execute C:\\1-AI\\APPS\\PYTHON: the qi-brain and "
    "qi-registry MCP servers spawned by Claude Code, plus Claude Voice helpers "
    "and the OC keepalive daemon. Windows will not delete an in-use image, so "
    "the directory cannot go until those exit.",
    "Repointing the MCP servers in ~/.claude.json restarts Claude Code and "
    "releases them - which is exactly why that step is scheduled last.",
    "VSCode and LM Studio still need reinstalling to standard locations. Their "
    "program files are the only other content left in C:\\1-AI.",
]:
    d.add_paragraph(t, style="List Bullet")

# ---------------------------------------------------------------- 4 and 5
d.add_heading("Phase 4 - planned, not executed", 1)
for t in [
    "qi_paths.py delivered and tested at C:\\QIH\\engine\\common\\qi_paths.py. "
    "Resolves code/data/config/cache/logs per app, creates directories on "
    "access, supports shared QI_DATA_DIR and per-app QI_<APP>_DATA_DIR "
    "overrides. Both override modes verified.",
    "Reference scan: about 70,000 path references across 21 project roots. The "
    "bulk are regenerated artifacts, snapshots and logs that must NOT be "
    "rewritten - a blind find/replace would corrupt the rollback data kept to "
    "recover from a bad migration.",
    "Junction-backed incremental move strategy with a 17-app ordering by blast "
    "radius: AkiyaScout first as a rehearsal, PlayDeck last as instructed. QIH "
    "stays at C:\\QIH.",
    "The QI_Headroom failure this session is the argument for Phase 4 in "
    "miniature: it broke because NSSM writes its logs INSIDE the venv, so "
    "renaming the venv took the log directory with it. Separating data from "
    "code prevents exactly that.",
    "File: migration_2026-08\\PHASE_4_Consolidation_Plan.md",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("Phase 5 - packaging plan: COMPLETE (plan only)", 1)
for t in [
    "PyInstaller --onedir over Nuitka: build-iteration speed matters more than "
    "binary purity now; revisit only if source protection becomes a real "
    "commercial requirement.",
    "Inno Setup over WiX for Tier A: no enterprise MSI or SCCM requirement "
    "exists, and NSSM registration is a trivial shell-out in Inno.",
    "Unsigned now; buy EV rather than OV at first external release, because EV "
    "gets immediate SmartScreen trust while OV must accrue reputation.",
    "Scope narrowed to 5 shippable apps. The QIH engine is explicitly excluded.",
    "File: migration_2026-08\\PHASE_5_Installer_Packaging_Plan.md",
]:
    d.add_paragraph(t, style="List Bullet")

# ---------------------------------------------------------------- next
d.add_heading("Next Up", 1)
for i, t in enumerate([
    "Repoint qi-brain and qi-registry in ~/.claude.json to C:\\Program Files\\"
    "Python311\\python.exe. This restarts Claude Code.",
    "In the new session: delete C:\\1-AI\\APPS\\PYTHON and replace it with a "
    "junction to C:\\Program Files\\Python311. This is the agreed bridge that "
    "keeps the 12 remaining stale venvs working with zero re-downloads.",
    "Reinstall VSCode and LM Studio to standard locations; then fix the last "
    "VSCode PATH entry.",
    "Delete the rest of C:\\1-AI (AvatarStudio original, the 593 MB pip cache, "
    "the two stray files).",
    "Dedupe the Brain's project_state: 53 rows -> 32, merging the display-name "
    "stubs away. Also reconcile fidelityanalyzer (in Brain, absent from "
    "qi_registry.json) and MailBrain (on disk, in neither).",
    "Begin Phase 4 with AkiyaScout as the rehearsal move.",
], 1):
    d.add_paragraph(f"{i}. {t}", style="List Number")

d.add_heading("Rollback assets created this session", 1)
for t in [
    "Restore point #143.",
    "rollback\\registry\\HKLM_Python.reg and HKCU_Python.reg.",
    "rollback\\PATH_Machine_before.txt and PATH_User_before.txt.",
    "C:\\CogniBase\\.venv.old and C:\\CLAUDE\\Tools\\headroom_env.old.",
    "149 Scripts\\*.exe.bak files.",
    "*.bak-phase3 alongside every rewritten doc and qi_registry.json.",
    "venv_freeze\\*.txt - pip freeze of both rebuilt venvs.",
]:
    d.add_paragraph(t, style="List Bullet")

d.save(OUT)
print("Saved: " + OUT)
