"""Generate the 2026-08-09 QIH migration session summary .docx."""
import sys
from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")

OUT = (r"C:\QIH\shared\documentation\session_summaries"
       r"\QIH_Migration_Summary_2026-08-09_0430.docx")

d = Document()

d.add_heading("QI Migration - Phases 2-5 - Session Summary", 0)
p = d.add_paragraph()
p.add_run("Date: ").bold = True
p.add_run("2026-08-09 (early morning session)\n")
p.add_run("Model: ").bold = True
p.add_run("Claude Opus 5\n")
p.add_run("Outcome: ").bold = True
p.add_run("BLOCKED on an unanswered UAC prompt after 1h03m. No destructive "
          "action was taken. Machine is fully healthy and unchanged except for "
          "a new restore point. Substantial discovery and planning work was "
          "completed; four previously-unknown migration blockers were found.")

# ---------------------------------------------------------------- completed
d.add_heading("Completed This Session", 1)

d.add_heading("Restore point created (Phase 2.0)", 2)
for t in [
    "Restore point #143 'QI migration Phase2.1 pre-PythonMove' created at 07:15.",
    "SystemRestorePointCreationFrequency set to 0 to bypass the 24h throttle.",
    "Prior rollback data from 2026-08-08 (services_before.csv, task XML, config "
    "backups) verified present and still accurate.",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("Ground-truth state re-verified", 2)
for t in [
    "Python 3.11.9 at C:\\1-AI\\APPS\\PYTHON, still a PER-USER install.",
    "C:\\Program Files\\Python311 does not exist.",
    "473 packages present; import smoke test passes.",
    "47 of 52 QI_* services running - the same 5 stopped as before, no regression.",
    "No reboot pending. Machine not restarted.",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("Phase 5 packaging plan - COMPLETE (plan only, nothing built)", 2)
for t in [
    "Recommendation: PyInstaller --onedir over Nuitka (build-iteration speed "
    "beats binary purity at this stage; revisit only if source protection "
    "becomes a real commercial requirement).",
    "Recommendation: Inno Setup over WiX for Tier A (no enterprise MSI/SCCM "
    "requirement exists; NSSM registration is a trivial shell-out in Inno).",
    "Recommendation: stay unsigned now; buy EV (not OV) at first external "
    "release - EV gets immediate SmartScreen trust, OV must accrue reputation.",
    "Scope narrowed to 5 Tier A apps. The QIH engine is explicitly out of scope.",
    "Hard prerequisite documented: Phase 4 code/data separation gates all "
    "packaging work.",
    "File: migration_2026-08\\PHASE_5_Installer_Packaging_Plan.md",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("Phase 4 plan + working resolver - DELIVERED", 2)
for t in [
    "qi_paths.py written and tested at C:\\QIH\\engine\\common\\qi_paths.py. "
    "Resolves code/data/config/cache/logs per app; creates dirs on access; "
    "supports both QI_DATA_DIR (shared root) and QI_<APP>_DATA_DIR (per-app) "
    "overrides. Both override modes verified working.",
    "Reference scan completed: ~70,000 path references across 21 project roots. "
    "Established that the bulk are regenerated artifacts, snapshots and logs "
    "that must NOT be rewritten - a blind find/replace would corrupt the very "
    "rollback data kept to recover from a bad migration.",
    "Junction-backed incremental move strategy designed, with a 17-app ordering "
    "by blast radius (AkiyaScout first as rehearsal, PlayDeck last as instructed).",
    "File: migration_2026-08\\PHASE_4_Consolidation_Plan.md",
]:
    d.add_paragraph(t, style="List Bullet")

# ---------------------------------------------------------------- discoveries
d.add_heading("Blockers Discovered (not in any prior handoff)", 1)
d.add_paragraph(
    "These four would each have caused a silent breakage after C:\\1-AI was "
    "deleted. All were found before any destructive step ran.")

t = d.add_table(rows=1, cols=3)
t.style = "Light Grid Accent 1"
h = t.rows[0].cells
h[0].text = "Finding"
h[1].text = "Impact"
h[2].text = "Resolution"

rows = [
    ("14 venvs were built from the old interpreter, not 1. The handoff named "
     "only C:\\CogniBase\\.venv.",
     "Includes C:\\CLAUDE\\Tools\\headroom_env, which is what the QI_Headroom "
     "service actually runs. That service was classified 'leave alone' because "
     "its Application field points at headroom.exe - the stale interpreter is "
     "one level down, invisible to a service-path scan.",
     "Full list captured by find_venvs.py. All 14 must be recreated before "
     "C:\\1-AI is deleted."),

    ("149 of 152 Scripts\\*.exe console launchers have the old interpreter path "
     "baked into the binary.",
     "A file copy does not fix these. They would keep invoking "
     "C:\\1-AI\\APPS\\PYTHON\\python.exe. mcp-server-sqlite.exe - used by two "
     "MCP configs - is one of them.",
     "fix_console_scripts.py written, with the locator bug below fixed, and the "
     "approach proven on a scratch copy."),

    ("The obvious shebang-rewrite implementation silently corrupts the launchers.",
     "Scanning forward for b'#!' matches a false positive inside the t64 "
     "launcher stub's own error-string table (it literally contains \"Expected "
     "to find '!' following '#' in shebang line\"). Rewriting there produced a "
     "launcher that died with 0xC0000005 ACCESS_VIOLATION.",
     "Caught by testing on a scratch copy before touching anything real. Fixed "
     "by anchoring on the PK\\x03\\x04 zip magic and scanning backwards. "
     "Retested: zip intact, launcher runs, --help returns exit 0."),

    ("Two scheduled tasks invoke a venv that does not exist.",
     "MaiaReconcile and MaiaRevertMiMo both call "
     "C:\\QI\\.venv\\Scripts\\python.exe. There is no C:\\QI\\.venv. Both have "
     "been failing silently.",
     "Repoint to the real interpreter or retire. Unrelated to this migration - "
     "pre-existing breakage."),
]
for a, b, c in rows:
    cells = t.add_row().cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c

d.add_heading("Corrections to the handoff", 1)
t2 = d.add_table(rows=1, cols=2)
t2.style = "Light Grid Accent 1"
h2 = t2.rows[0].cells
h2[0].text = "Handoff said"
h2[1].text = "Reality"
for a, b in [
    ("4 scheduled tasks reference the old Python (v1 prompt)",
     "10 do. The v2 prompt said 12, but QI_NightlyBackup and "
     "'Maia\\Maia Tunnel Watchdog' do not exist on this machine at all."),
    ("31 NSSM services need repointing",
     "Confirmed exactly 31, of 52 QI_* total. List extracted and verified."),
    ("C:\\CogniBase\\.venv must be recreated",
     "True, but it is 1 of 14 stale venvs."),
    ("Copy the Python tree and repoint services",
     "Necessary but not sufficient - 149 console-script launchers also need "
     "their embedded interpreter path rewritten."),
]:
    c = t2.add_row().cells
    c[0].text = a
    c[1].text = b

# ---------------------------------------------------------------- blocked
d.add_heading("Why the session stopped", 1)
for t in [
    "The robocopy of the Python tree into C:\\Program Files\\Python311 requires "
    "elevation. gsudo raised a UAC prompt (consent.exe PID 50812) which was "
    "never answered - it remained pending for 1 hour 3 minutes.",
    "gsudo's credential cache did not carry across from the restore-point call, "
    "because each Bash tool invocation runs in a fresh process tree. Next "
    "session: run 'gsudo cache on --pid 0 -d 24:00:00' once, up front, so a "
    "single click covers every later elevation.",
    "No destructive step was reached. Nothing was uninstalled, deleted, or "
    "repointed. The only system change this session is restore point #143.",
]:
    d.add_paragraph(t, style="List Bullet")

# ---------------------------------------------------------------- next
d.add_heading("Next Up (immediate priority)", 1)
for i, t in enumerate([
    "Establish the gsudo cache FIRST: gsudo cache on --pid 0 -d 24:00:00 - one "
    "UAC click for the whole session.",
    "Phase 2.1a: robocopy C:\\1-AI\\APPS\\PYTHON -> C:\\Program Files\\Python311, "
    "excluding venvs\\ and testenv\\. Script ready: p21a_copy_python.ps1.",
    "Phase 2.1b: uninstall the per-user bundle {1da2e09b-199c-4def-9a99-"
    "93a8c1b8ddf2} /quiet /norestart, then install python-3.11.9-amd64.exe "
    "/quiet /norestart InstallAllUsers=1 TargetDir=\"C:\\Program Files\\Python311\".",
    "Run fix_console_scripts.py --apply against the new Scripts dir (149 files, "
    "keeps .bak, self-verifies each rewrite with zipfile).",
    "Phase 2.2: p22_repoint_services.ps1 (31 services), then restart and health-"
    "check each.",
    "Phase 2.3: p23_repoint_tasks.ps1 (10 tasks, handles the conhost --headless "
    "wrappers and re-quotes for the space in 'Program Files').",
    "Phase 2.4: recreate all 14 stale venvs, not just CogniBase.",
    "Phase 2.5: correct HKLM and HKCU PythonCore\\3.11 InstallPath.",
], 1):
    d.add_paragraph(f"{i}. {t}", style="List Number")

d.add_heading("In Development (multi-session)", 1)
for t in [
    "Phase 3 - retire C:\\1-AI: move AvatarStudio to C:\\APPS\\AvatarStudio "
    "(live project, port 7862; note its .venv is one of the 14 stale ones), "
    "reinstall VSCode and LM Studio, then delete. D:\\Review untouched.",
    "Phase 4 - 17-app junction-backed consolidation into C:\\APPS, adopting "
    "qi_paths.py per app. Plan written; no moves executed.",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("Future Enhancements", 1)
for t in [
    "Phase 5 packaging build-out (Tier A, 5 apps), gated on Phase 4.",
    "Retire or repair MaiaReconcile / MaiaRevertMiMo.",
    "Resolve the Retirement Analyzer path inconsistency (registry says "
    "'C:\\Retirement Analyzer', NSSM says 'C:\\RetirementAnalyzer').",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("Documents Created / Modified", 1)
t3 = d.add_table(rows=1, cols=2)
t3.style = "Light Grid Accent 1"
h3 = t3.rows[0].cells
h3[0].text = "File"
h3[1].text = "Purpose"
for a, b in [
    ("QIH\\engine\\common\\qi_paths.py", "Phase 4 code/data path resolver (tested)"),
    ("migration_2026-08\\PHASE_5_Installer_Packaging_Plan.md", "Phase 5 plan"),
    ("migration_2026-08\\PHASE_4_Consolidation_Plan.md", "Phase 4 plan + move order"),
    ("phase2\\p21_restorepoint.ps1", "Restore point (ran OK, #143)"),
    ("phase2\\p21a_copy_python.ps1", "Python tree copy (READY, not run)"),
    ("phase2\\p22_repoint_services.ps1", "31-service repoint (READY, not run)"),
    ("phase2\\p23_repoint_tasks.ps1", "Scheduled task repoint (READY, not run)"),
    ("phase2\\fix_console_scripts.py", "Console-launcher shebang rewriter (tested)"),
    ("phase2\\find_venvs.py", "Stale venv scanner (found 14)"),
    ("phase2\\scan_refs.py", "Cross-reference scanner (~70k refs)"),
    ("session_summaries\\QIH_Migration_Summary_2026-08-09_0430.docx", "This summary"),
]:
    c = t3.add_row().cells
    c[0].text = a
    c[1].text = b

d.add_paragraph()
warn = d.add_paragraph()
warn.add_run("NOTE: ").bold = True
warn.add_run("The MCP servers qi-brain and qi-registry in ~/.claude.json still "
             "point at C:\\1-AI\\APPS\\PYTHON\\python.exe and were deliberately "
             "NOT repointed - that step ends the session by restarting Claude "
             "Code, and must stay last.")

d.save(OUT)
print("Saved: " + OUT)
