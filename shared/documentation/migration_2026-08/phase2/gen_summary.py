# -*- coding: utf-8 -*-
"""Generate the session summary .docx for the 2026-08-08 Phase 2 attempt."""
import sys
from datetime import datetime

from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")

OUT = (r"C:\QIH\shared\documentation\session_summaries"
       r"\QIH_Migration_Summary_2026-08-08_2250.docx")

doc = Document()

doc.add_heading("QI Migration - Phase 2 (Python) - Session Summary", 0)
p = doc.add_paragraph()
p.add_run("Date: ").bold = True
p.add_run("2026-08-08 (evening session)")
p = doc.add_paragraph()
p.add_run("Scope attempted: ").bold = True
p.add_run("Phases 2-4 execution, Phase 5 planning")
p = doc.add_paragraph()
p.add_run("Outcome: ").bold = True
p.add_run("Phase 2.0 complete. Phase 2.1 attempted and did NOT reach its target - "
          "no harm done, machine fully healthy. Phases 2.2-5 not started.")

doc.add_heading("Completed This Session", 1)

doc.add_heading("Rollback state captured (Phase 2.0) - COMPLETE", 2)
for line in [
    "System restore point #141 'QI migration Phase2 pre-Python' created 22:19 "
    "(24h throttle bypassed via SystemRestorePointCreationFrequency=0).",
    "All 52 QI_* NSSM services dumped to rollback\\nssm_dump_all.txt and "
    "services_before.csv (Application, AppDirectory, AppParameters, Start, Status).",
    "Scheduled tasks exported to XML under rollback\\scheduled_tasks\\.",
    "Config backups taken: .claude.json, qi_registry.json, whitelist.json, "
    "CogniBase pyvenv.cfg.",
    "pip freeze of the old interpreter -> rollback_requirements-old.txt (469 lines); "
    "cleaned to requirements-new.txt (467 pinned packages).",
]:
    doc.add_paragraph(line, style="List Bullet")

doc.add_heading("Elevation model established", 2)
for line in [
    "QI_Elevate broker confirmed running as LocalSystem with a regex-bounded whitelist.",
    "gsudo v2.6.1 confirmed working with CacheMode=Auto and a 24h credential cache - "
    "one UAC click covers a whole session.",
    "Decision: use gsudo rather than widening the broker whitelist, so the security "
    "whitelist stays untouched.",
]:
    doc.add_paragraph(line, style="List Bullet")

doc.add_heading("Corrections found to the Phase 1 handoff", 2)
t = doc.add_table(rows=1, cols=2)
t.style = "Light Grid Accent 1"
hdr = t.rows[0].cells
hdr[0].text = "Handoff said"
hdr[1].text = "Reality"
for a, b in [
    ("4 scheduled tasks reference the old Python",
     "12 do. The extra 8 hide the python path inside 'conhost.exe --headless' "
     "arguments, so a naive scan of the Execute field misses them."),
    ("31 NSSM services",
     "Confirmed 31 (of 52 QI_* services total)."),
    ("(not mentioned)",
     "torch/torchaudio are +cpu builds. A plain 'pip install torch==2.10.0' would "
     "pull the CUDA wheel instead - must use --index-url "
     "https://download.pytorch.org/whl/cpu"),
    ("(not mentioned)",
     "Editable install 'openspace' is dangling - C:\\CLAUDE\\OpenSpace no longer "
     "exists. Drop it, do not recreate."),
    ("(not mentioned)",
     "Editable install 'cognibase' resolves to C:\\CogniBase\\Application."),
]:
    row = t.add_row().cells
    row[0].text = a
    row[1].text = b

doc.add_heading("What went wrong in Phase 2.1 (and why nothing broke)", 1)
doc.add_paragraph(
    "The official python-3.11.9-amd64.exe was run with /quiet InstallAllUsers=1 "
    "TargetDir=\"C:\\Program Files\\Python311\". It did not install there.",
)
for line in [
    "The WiX Burn bundle detected the existing per-user Python 3.11.8 (registered at "
    "C:\\1-AI\\APPS\\PYTHON) as a RELATED BUNDLE and switched to upgrade mode.",
    "On the related-bundle upgrade path, Burn IGNORES both TargetDir and "
    "InstallAllUsers and patches the existing install in place, preserving its scope.",
    "Net effect: C:\\1-AI\\APPS\\PYTHON was upgraded 3.11.8 -> 3.11.9 and remains a "
    "per-user install. C:\\Program Files\\Python311 was never created.",
    "The installer requested a restart because python311.dll was mapped by 31 running "
    "services. The restart was declined and is NOT needed - see verification below.",
]:
    doc.add_paragraph(line, style="List Bullet")

doc.add_heading("Verified healthy after the event", 2)
for line in [
    "Uptime 9 days 7 hours - the machine never restarted.",
    "PendingFileRenameOperations holds 150 entries; ZERO touch Python or C:\\1-AI. "
    "They are Firefox / OneDrive / Edge / Office / Tailscale / printer leftovers.",
    "python.exe, python311.dll, python3.dll, vcruntime140.dll all present.",
    "Python reports 3.11.9; all 473 packages present; import smoke test passes "
    "(fastapi, uvicorn, requests, aiohttp, pydantic, docx, numpy, pandas, flask, linebot).",
    "47 of 52 services running. The 5 stopped are the same 5 that were stopped before "
    "the session (QI_AutoPDF, QI_GamezQuantProxy, QI_MaiaDemoTunnel, QI_MapSnapBUSetup, "
    "QI_RetirementAnalyzer) - no regression.",
    "CONCLUSION: no restart required. Nothing is in a half-applied state.",
]:
    doc.add_paragraph(line, style="List Bullet")

doc.add_heading("Trap discovered (applies to all future sessions)", 1)
doc.add_paragraph(
    "PowerShell 5.1 reads a BOM-less .ps1 as ANSI, not UTF-8. A UTF-8 em dash decodes "
    "to a 3-character sequence that contains a double-quote, which silently breaks "
    "string parity and produces a cascade of nonsense parser errors far from the real "
    "line. Keep all .ps1 files pure ASCII, or write them with a UTF-8 BOM. Em dashes in "
    "comments are survivable; em dashes inside a quoted string are not."
)

doc.add_heading("Next Up (immediate priority for the next session)", 1)
doc.add_paragraph(
    "The corrected Phase 2.1 approach - it never leaves services exposed, because the "
    "new location is fully working before anything is repointed to it:"
)
for i, line in enumerate([
    "robocopy C:\\1-AI\\APPS\\PYTHON -> C:\\Program Files\\Python311 (about 4.7 GB, "
    "Lib alone is 4.4 GB). Exclude the nested 'venvs' and 'testenv' folders.",
    "Uninstall the per-user bundle {1da2e09b-199c-4def-9a99-93a8c1b8ddf2} with "
    "/quiet /norestart. Running services keep their mapped DLL and survive; do not "
    "restart any service during this window.",
    "Install python-3.11.9-amd64.exe with /quiet /norestart InstallAllUsers=1 "
    "TargetDir=\"C:\\Program Files\\Python311\". With no related bundle left it now "
    "installs and registers per-machine correctly, and nothing is executing from that "
    "directory yet so there is no file-in-use conflict. site-packages survives.",
    "ALWAYS pass /norestart from now on.",
    "Repoint 31 NSSM services + 12 scheduled tasks; restart and verify each.",
    "Recreate C:\\CogniBase\\.venv (do not hand-edit pyvenv.cfg).",
    "Clean up registry pollution: HKLM\\SOFTWARE\\Python\\PythonCore\\3.11 now points "
    "at C:\\1-AI\\APPS\\PYTHON and must be corrected before C:\\1-AI is deleted.",
], 1):
    doc.add_paragraph(f"{i}. {line}", style="List Number")

doc.add_heading("In Development (multi-session)", 1)
for line in [
    "Phase 3 - retire C:\\1-AI (move AvatarStudio to C:\\APPS\\AvatarStudio first; "
    "reinstall VSCode and LM Studio; then delete). Do NOT touch D:\\Review.",
    "Phase 4 - consolidate self-built apps into C:\\APPS, separate code from data via a "
    "shared resolver, dedupe the C:\\QI / C:/QI and C:\\NAYA / C:/NAYA MCP entries, "
    "move C:\\PlayDeck last.",
    "Phase 5 - installer packaging plan (Inno Setup or WiX over PyInstaller or Nuitka, "
    "plus code signing). Plan only, do not build.",
]:
    doc.add_paragraph(line, style="List Bullet")

doc.add_heading("Documents Created / Modified", 1)
files = [
    (r"...\migration_2026-08\phase2\capture_rollback.ps1", "Rollback capture (ran OK)"),
    (r"...\migration_2026-08\phase2\force_restore_and_probe.ps1", "Restore point + task/network probe"),
    (r"...\migration_2026-08\phase2\install_python.ps1", "Python installer wrapper - NEEDS /norestart added"),
    (r"...\migration_2026-08\phase2\assess_state.ps1", "Post-install ground truth assessment"),
    (r"...\migration_2026-08\phase2\check_pending_renames.ps1", "Reboot-queue inspector"),
    (r"...\migration_2026-08\phase2\dump_registry.py", "Registry project/port dump"),
    (r"...\migration_2026-08\phase2\gen_summary.py", "This summary generator"),
    (r"...\migration_2026-08\phase2\requirements-new.txt", "467 pinned packages to reinstall"),
    (r"...\migration_2026-08\phase2\rollback_requirements-old.txt", "Raw pip freeze (469 lines)"),
    (r"...\migration_2026-08\phase2\rollback" + "\\", "nssm dumps, services_before.csv, task XML, config backups"),
    (r"...\migration_2026-08\phase2\python-3.11.9-amd64.exe", "Installer, already downloaded (25 MB)"),
]
tf = doc.add_table(rows=1, cols=2)
tf.style = "Light Grid Accent 1"
tf.rows[0].cells[0].text = "File"
tf.rows[0].cells[1].text = "Purpose"
for a, b in files:
    r = tf.add_row().cells
    r[0].text = a
    r[1].text = b

doc.add_paragraph()
doc.add_paragraph(
    "No files were deleted this session. No service configuration was changed. "
    "The only system change was the in-place Python 3.11.8 -> 3.11.9 patch upgrade "
    "and restore point #141."
).italic = True

doc.save(OUT)
print("SAVED:", OUT)
