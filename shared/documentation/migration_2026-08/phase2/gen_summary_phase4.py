"""Session summary for the Phase 4 app consolidation, 2026-08-09."""
import sys
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")

OUT = (r"C:\QIH\shared\documentation\session_summaries"
       r"\QIH_Migration_Phase4_Summary_2026-08-09_1500.docx")

d = Document()
d.add_heading("QI Migration - Phase 4: app consolidation to C:\\APPS", 0)

p = d.add_paragraph()
p.add_run("Date: ").bold = True
p.add_run("2026-08-09\n")
p.add_run("Status: ").bold = True
p.add_run("PAUSED at Renne's request until next weekend. All 25 apps are moved "
          "and every service is healthy, but nothing will be deleted until "
          "Renne has tested the apps himself. That is the correct call - see "
          "the LM Studio finding below.")

d.add_heading("Completed", 1)
for t in [
    "All 25 self-built apps moved from the C: root to C:\\APPS. Every move "
    "verified: file counts matched, services repointed and restarted, "
    "C:\\<App> removed, original retained at D:\\_PREMOVE_2026-08-09\\<App>.",
    "No junction or pointer folder left behind, per Renne's requirement.",
    "Nothing deleted. AutoPDF_Portable preserved as AutoPDF_Portable_Dupe; "
    "empty QIB retired as 'QIB_for deletion'.",
    "C:\\1-AI retired; C:\\CLAUDE retired to C:\\CLAUDE.RETIRED_2026-08-09.",
    "Services steady at 47 running / 5 stopped throughout - exact parity with "
    "the pre-migration baseline. All health endpoints answering, ComfyUI "
    "included.",
    "C:\\QIH confirmed as a PERMANENT exception - it stays at the root.",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("THE BLOCKER - do not delete 1-AI.RETIRED yet", 1)
d.add_paragraph(
    "Renne asked directly whether AvatarStudio, LM Studio, Python and VSCode "
    "had all migrated successfully, and said he would delete the retired tree "
    "on that answer. Verification found three safe and one not:")

t = d.add_table(rows=1, cols=3)
t.style = "Light Grid Accent 1"
h = t.rows[0].cells
h[0].text = "Component"
h[1].text = "Live copy outside the retired tree"
h[2].text = "Verdict"
for a, b, c in [
    ("AvatarStudio", "C:\\APPS\\AvatarStudio", "SAFE"),
    ("Python 3.11", "C:\\Program Files\\Python311", "SAFE"),
    ("VSCode", "C:\\Program Files\\Microsoft VS Code", "SAFE"),
    ("LM Studio", "NONE - only copy is inside C:\\1-AI.RETIRED_2026-08-09",
     "NOT SAFE"),
]:
    cells = t.add_row().cells
    cells[0].text = a
    cells[1].text = b
    cells[2].text = c

d.add_paragraph()
d.add_paragraph(
    "LM Studio was never reinstalled - only its winget id was identified. "
    "Models and settings in C:\\Users\\renne\\.lmstudio survive regardless, so "
    "it would be a reinstall rather than data loss, but the tree cannot be "
    "deleted until a real copy exists elsewhere. Command:")
d.add_paragraph(
    "winget install --id ElementLabs.LMStudio --scope machine --silent "
    "--accept-package-agreements --accept-source-agreements",
    style="Intense Quote")
d.add_paragraph(
    "Then re-run phase2\\verify_before_delete.ps1 - it flips LM Studio to SAFE "
    "automatically once a real copy exists.")

d.add_heading("Problems found and fixed during the moves", 1)
t2 = d.add_table(rows=1, cols=2)
t2.style = "Light Grid Accent 1"
hh = t2.rows[0].cells
hh[0].text = "Problem"
hh[1].text = "Resolution"
for a, b in [
    ("move_playdeck.ps1 exited 255 with no output",
     "It killed itself. The filter matched any process whose command line "
     "contained 'PlayDeck' without 'APPS' - which describes the script's own "
     "path, ...\\phase2\\move_playdeck.ps1. Fixed by anchoring on the app "
     "directory and excluding the process and its ancestors."),
    ("PlayDeck holds a live SQLite database",
     "The generic mover CORRECTLY refused it on a file-count mismatch: "
     "playdeck.db-wal and -shm could not be copied. Copying a .db while its "
     "WAL is live risks losing recent commits. Stopping the service made both "
     "sidecars disappear - the clean-checkpoint signal - and the copied "
     "database matched SHA256 byte-for-byte."),
    ("The venv console-script fixer gave a FALSE ALL-CLEAR",
     "It reported 0 of 49 headroom launchers stale. Bash had eaten the "
     "backslashes in --old \"C:\\CLAUDE\", so it compared against 'C:CLAUDE' "
     "and matched nothing. Passing the arguments from PowerShell instead "
     "revealed 44 of 49 were stale. All 44 fixed and verified; headroom.exe "
     "now runs from C:\\APPS\\CLAUDE on port 9020. Trusting the first result "
     "would have left QI_Headroom quietly broken."),
    ("Several folders survived their move holding their own nssm.exe",
     "C:\\OC, C:\\NEXUS and C:\\PlayDeck each kept the service-host binary "
     "locked. Fixed by repointing each service's ImagePath to the copy at "
     "C:\\APPS\\<App> - no service was removed, per the no-delete rule."),
    ("A folder can linger empty because a process holds it as its cwd",
     "C:\\NEXUS did this while Renne had it open; C:\\CLAUDE is doing it now. "
     "Harmless - both clear on reboot."),
]:
    cells = t2.add_row().cells
    cells[0].text = a
    cells[1].text = b

d.add_heading("C: root as it stands", 1)
for t3 in [
    "APPS - all 25 apps",
    "QIH - the Hive (permanent exception)",
    "TEMP, tmp - keep",
    "1-AI - junction stub -> C:\\Program Files\\Python311, needed until the 11 "
    "remaining venvs are rebuilt",
    "1-AI.RETIRED_2026-08-09 - 17.28 GB, BLOCKED on LM Studio",
    "CLAUDE - empty shell, clears on reboot",
    "CLAUDE.RETIRED_2026-08-09 - 28 duplicate .pyd files",
    "ARCHIVE - 361.7 GB, to be reviewed with Renne over time",
    "GOOSE, Plex - third-party",
]:
    d.add_paragraph(t3, style="List Bullet")

d.add_heading("Next session (next weekend)", 1)
for i, t4 in enumerate([
    "Renne tests the apps. Nothing is deleted before that.",
    "Reinstall LM Studio, then re-run verify_before_delete.ps1.",
    "Once all four read SAFE, delete 1-AI.RETIRED (17.28 GB) and "
    "CLAUDE.RETIRED.",
    "Rebuild the 11 remaining venvs so the C:\\1-AI junction can go.",
    "Remove the duplicate tunnel services NEXUSTunnel and NayaTunnel.",
    "Rename OC-Keepalive-Service and ClaudeManager to the QI_ convention and "
    "register them in QI_Service_Registry.md.",
    "Dedupe the Brain's project_state - 53 rows for 32 projects.",
    "Begin ARCHIVE review together.",
    "Then the code/data separation half of Phase 4: adopt qi_paths.py per app. "
    "NOT started - the moved apps still write logs and DBs beside their code.",
], 1):
    d.add_paragraph("%d. %s" % (i, t4), style="List Number")

d.add_heading("Rollback assets", 1)
for t5 in [
    "D:\\_PREMOVE_2026-08-09\\<App> - the original of every moved app.",
    "C:\\1-AI.RETIRED_2026-08-09 and C:\\CLAUDE.RETIRED_2026-08-09.",
    "*.bak-move alongside every rewritten config and doc.",
    "*.bak-venvmove alongside all 44 rewritten launchers.",
    "Restore point #143, registry .reg exports, PATH_*_before.txt.",
]:
    d.add_paragraph(t5, style="List Bullet")

d.save(OUT)
print("Saved: " + OUT)
