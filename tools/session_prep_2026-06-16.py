# -*- coding: utf-8 -*-
"""End-of-session prep: write the session summary .docx + log to Brain."""
import sys, sqlite3
from datetime import datetime
from pathlib import Path
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt

NOW = datetime.now()
STAMP = NOW.strftime("%Y-%m-%d_%H%M")
SHARED = Path(r"C:\QIH\shared\documentation\session_summaries")
SHARED.mkdir(parents=True, exist_ok=True)
OUT = SHARED / f"QIHive_Summary_{STAMP}.docx"

doc = Document()
doc.add_heading("QI Hive — Session Summary", 0)
doc.add_paragraph(f"Date: {NOW.strftime('%Y-%m-%d %H:%M')}  ·  Project: QI Hive / Claude Manager")

def H(t): doc.add_heading(t, level=1)
def B(items):
    for i in items:
        doc.add_paragraph(i, style="List Bullet")

H("✅ Completed This Session")
B([
 "Reviewed QI Hive: verified all NSSM services + tunnels up; produced the live Cloudflare address list (8 public tunnels, all HTTP 200).",
 "Stood up persistent tunnels for LotteryWiz and CypherMiner (NSSM services); deleted a dangerous conflicting CypherMiner install_service.bat.",
 "Registered previously-missing projects into registry + Brain: lotterywiz, cypherminer, digitization, avatarstudio, tubescout (Brain backfill), claude_manager (registry).",
 "Built a brand-new compliant app: Fidelity Portfolio Analyzer (C:\\FidelityAnalyzer) via qi_new_project.py — analyzer engine + FastAPI :8504 + Gradio :7844, sample data, 19/19 compliance, full QI doc set, first git commit.",
 "Per-app independence: generated install_service.bat/install_tunnel.bat (DEMAND_START) for 5 apps + a one-click master installer C:\\QIH\\tools\\install_all_qi_app_services.bat. Public tunnels for m2v + tubescout; sensitive apps service-only.",
 "Resolved 2 registry port conflicts (8650 universal.launcher, 9011 qi_hive.brain_api).",
 "DEEP AUDIT of all 18 dashboard tabs (5 parallel agents) + comprehensive fixes.",
])

H("🔧 Dashboard Audit Fixes (all 18 tabs)")
B([
 "status.json regenerated from Brain (15->22 projects, lowercase ids, dedup).",
 "Registry: removed orphan digicost (22==Brain), backfilled paths.logs for 12 projects.",
 "Registry-driven refactors: health_check.py (11->22), test_smoke.py (4->25), project_status.py (12->22), usage_stats.py (labels + fable pricing), dashboard QI_PROJECTS.",
 "One-line bugs: brain/api.py sqlite3.Row.get 500 (Agent Profile), _collect_tasks 'QI-'->'QI_' (6 tasks), _brain_get UTF-8 (Headlines), render_project case-insensitive + fallback (Project detail).",
 "Render quality: Dashboard cards (names/ids), Hive Recent Sessions (Brain query), War Room dispatch field mapping + agent ids, Dispatch badge maps + flood cap, Logs multi-extension glob.",
 "Guide regenerated: 4 -> 18 tabs, all 22 projects.",
 "Verified: all 18 tabs + project detail + agent profile return 200.",
])

H("🔄 Next Up (immediate)")
B([
 "Run C:\\QIH\\tools\\install_all_qi_app_services.bat as admin (UAC) to make the 5 apps persistent (DEMAND_START).",
 "Rotate the AvatarStudio D-ID API key and move it from studio_config.json into secrets/.",
 "Migrate the Digitization Cost Tool out of Downloads into a proper C:\\ project folder + git init.",
])

H("🚀 In Development / Deferred")
B([
 "Brain tab: add a dedicated 'Sessions' sub-tab.",
 "Config tab: replace the obsolete gsudo card with a QI_Elevate broker card; surface QI_* services.",
 "agent_growth_log is empty + tasks.json deleted -> agent 'tasks logged' counts are 0 (agents must call log_growth).",
 "3 Headlines entries have source-corrupted text stored in the Brain DB.",
 "Fix nightly_reconcile.py status.json output path (writes C:\\QIH\\status.json; dashboard reads C:\\QIH\\data\\status.json).",
])

H("📁 Key Files Changed")
B([
 "C:\\QIH\\engine\\hive\\dashboard\\server.py (render fixes, KNOWN_TUNNELS, QI_PROJECTS)",
 "C:\\QIH\\engine\\brain\\api.py (Row.get 500 fix)",
 "C:\\QIH\\engine\\hive\\health_check.py, dashboard\\project_status.py, engine\\common\\usage_stats.py (registry-driven)",
 "C:\\Claude\\Tests\\test_smoke.py (25 targets)",
 "C:\\QIH\\ecosystem\\qi_registry.json, C:\\QIH\\data\\status.json (22 projects)",
 "C:\\QIH\\ecosystem\\QI_Claude_Manager_Guide.md (regenerated)",
 "C:\\FidelityAnalyzer\\* (new app), per-app install scripts, C:\\QIH\\tools\\*2026-06-1*.py",
])

doc.save(str(OUT))
print("Saved:", OUT)

# ── Log to Brain ──
DB = r"C:\QIH\data\qi_brain.db"
con = sqlite3.connect(DB); cur = con.cursor()
ts = NOW.strftime("%Y-%m-%d %H:%M:%S")
title = "QI Hive — registrations, Fidelity Analyzer build, per-app independence, 18-tab dashboard audit"
summary = ("Verified services/tunnels; registered missing projects; built Fidelity Portfolio Analyzer; "
           "generated per-app DEMAND_START service+tunnel installers; deep-audited and fixed all 18 dashboard tabs "
           "(status.json regen, registry-driven module refactors, one-line 500/typo/utf-8/case bugs, render quality, Guide regen).")
cur.execute("INSERT INTO session_log (project_id,agent_id,session_title,summary,decisions_made,features_logged,files_changed,next_steps,model_used,started_at,ended_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
    ("qi_hive","claude",title,summary,0,0,"[]",
     "Run install_all_qi_app_services.bat; rotate AvatarStudio key; migrate Digitization; Brain Sessions tab; Config->QI_Elevate.",
     "opus-4.8",ts,ts))
cur.execute("INSERT INTO project_state (project_id,agent_id,phase,status,summary,blockers,next_steps,recorded_at) VALUES (?,?,?,?,?,?,?,?)",
    ("qi_hive","claude","Phase 3 — agentic operation","active",
     "Dashboard fully reconciled to 22 projects across all 18 tabs; ecosystem registry==Brain; per-app service+tunnel installers ready.",
     None,
     "User to run install_all_qi_app_services.bat (UAC). Then: Brain Sessions tab, Config QI_Elevate card, fix reconcile status.json path.",
     ts))
con.commit(); con.close()
print("Brain: session_log + project_state logged for qi_hive")
