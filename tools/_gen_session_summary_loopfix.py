# -*- coding: utf-8 -*-
import sys
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")

OUT = Path(r"C:\QIH\shared\documentation\session_summaries")
OUT.mkdir(parents=True, exist_ok=True)
stamp = datetime.now().strftime("%Y-%m-%d_%H%M")
path = OUT / f"QIHive_Summary_{stamp}.docx"

d = Document()
d.add_heading("QI Hive — Session Summary", 0)
d.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
d.add_paragraph("Title: Fix broken/stale 'distilled brain memory' (qi_hive feedback loop)")

d.add_heading("✅ Completed This Session", 1)
for t in [
    "Root-caused the broken 'distilled brain memory' under The Hive: qi_hive's "
    "state file IS C:\\QIH\\data\\status.json — the same file the reconciler "
    "writes Brain summaries back into. The poller stamped '[auto:state_file] ' "
    "onto each summary; the reconciler wrote it back to status.json; the next "
    "poll re-ingested and re-prefixed it, compounding to "
    "'[auto:state_file] [auto:state_file] ...' until the real text was pushed "
    "past the 300-char cap. Only qi_hive looped because every other project "
    "reads its own separate state file.",
    "poller.py: added _strip_auto_prefix(); strips [auto:...] markers on read "
    "and compares stripped values in change-detection (root-cause break).",
    "nightly_reconcile.py: strips markers before writing status.json + "
    "LATEST.md so the displayed summary stays clean.",
    "Repaired corrupted data: deleted 586 compounded project_state rows, "
    "inserted 1 clean current row, cleaned status.json (DB backed up to "
    "qi_brain.db.bak-loopfix-20260623).",
    "Restarted QI_HiveIngest (normal) and QI_BrainAPI (via taskkill + NSSM "
    "respawn, to avoid taking down its 4 dependents: NEXUS, NayaBot, MaiaBot, "
    "Dashboard).",
    "Verified: 3 consecutive poll cycles wrote 0 changes / 0 new rows; "
    "status.json and LATEST.md qi_hive summaries are clean and stable.",
    "Committed fix (ad317cf) — 3 files, scoped to the fix only.",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("🔄 Next Up", 1)
for t in [
    "Optional: have the poller also capture next_steps from state files "
    "(currently LATEST.md shows 'Next: None' for poller-written rows).",
    "Optional: archive the remaining single-prefix historical qi_hive rows via "
    "the distiller for a fully tidy history.",
    "Monitor the dashboard /hive + LATEST.md over the next day to confirm the "
    "summary stays clean through the nightly reconcile.",
]:
    d.add_paragraph(t, style="List Bullet")

d.add_heading("📁 Documents Updated", 1)
for t in [
    "engine/brain/poller.py",
    "tools/nightly_reconcile.py",
    "tools/repair_qi_hive_loop_20260623.py (new)",
    "C:\\QIH\\data\\status.json (qi_hive summary cleaned)",
    "C:\\QIH\\LATEST.md (regenerated clean)",
    "C:\\QIH\\data\\qi_brain.db (586 rows pruned; backup created)",
]:
    d.add_paragraph(t, style="List Bullet")

d.save(path)
print("SAVED:", path)
