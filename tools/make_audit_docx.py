# -*- coding: utf-8 -*-
"""Generate the QI session-summary .docx for the 2026-06-16 security audit."""
import sys
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

OUT_DIR = Path(r"C:\QIH\shared\documentation\session_summaries")
OUT_DIR.mkdir(parents=True, exist_ok=True)
now = datetime.now()
fn = OUT_DIR / f"QIH_Summary_{now.strftime('%Y-%m-%d_%H%M')}_SecurityAudit.docx"

doc = Document()
t = doc.add_heading("QI Ecosystem — Secret-Leak & Git-Hygiene Audit", 0)
p = doc.add_paragraph()
r = p.add_run(f"Session date: {now.strftime('%Y-%m-%d %H:%M')}  ·  Scope: all 17 QI git "
              "repos + both nightly sync scripts")
r.italic = True


def h1(x): doc.add_heading(x, 1)
def h2(x): doc.add_heading(x, 2)
def para(x): doc.add_paragraph(x)
def bullet(x): doc.add_paragraph(x, style="List Bullet")


h1("✅ Completed This Session")
for x in [
    "Built a deterministic ecosystem secret scanner (gitleaks unavailable) covering "
    "working tree + FULL git history + remote-URL credentials + repo visibility.",
    "Scanned all 17 QI git repos; triaged every hit by REAL exposure "
    "(pushed / tracked-tip / local-history / gitignored / false-positive).",
    "REMEDIATED PersonalSong (only truly-pushed leak): purged the live Plex token from "
    "all history via git-filter-repo, untracked + gitignored the config, added a "
    "sanitized template, and force-pushed the private repo. Token verified gone.",
    "Scrubbed Maia's plaintext GitHub PAT from .git/config (remote URL → credential-less).",
    "Hardened BOTH nightly sync scripts with a fail-closed secret gate that aborts "
    "commit/push if any real secret is staged (smoke-tested).",
    "Authored shared QI_baseline.gitignore policy and a consolidated severity report.",
]:
    bullet(x)

h1("🔎 Key Findings")
h2("Genuinely exposed on a remote")
bullet("PersonalSong — Plex token tracked tip + pushed history (private). FIXED.")
h2("Plaintext on disk (local only)")
bullet("Maia — GitHub PAT in .git/config. FIXED (scrubbed). Must be rotated.")
h2("Real secrets in LOCAL-only history / tracked (never pushed)")
bullet("Maia — 4 Telegram bot tokens + 1 hex app secret (commit 631fe769).")
bullet("Naya — 2 Telegram bot tokens (commit baffed22).")
bullet("OpenClaw — Telegram token (tracked) + Anthropic key (untracked doc); repo has no remote.")
h2("Correctly contained (gitignored, never committed)")
bullet("NEXUS — OpenRouter + OpenAI + 2× Google keys in secrets/ and *API Key*.txt.")
bullet("EasyFlow — Chrome extension .pem in keys/.")
bullet("TubeScout — OpenAI-style key only in gitignored tubescout.db-wal (verify).")
h2("False positives")
bullet("cognibase test fixture, mapsnap HTML field literal, tubescout path-references.")

h1("🛠️ Root Cause & Hardening")
para("Root cause: nightly_git_sync.py used `git add -A` with only .gitignore as the guard. "
     "Maia's script used an allowlist but had no secret gate.")
for x in [
    "secret_gate.py — importable pre-commit/staged-diff secret scanner.",
    "nightly_git_sync.py + maia_nightly_sync.py — gated; abort + reset on any staged secret; "
    "fail closed if the gate is unavailable.",
    "QI_baseline.gitignore — shared policy (secrets/, .env, *config*.json, *.pem, ...).",
    "secret_audit.py / triage_secrets.py — re-runnable scanners for future passes.",
]:
    bullet(x)

h1("🔑 Owner Rotation Checklist (REQUIRED — Claude cannot rotate)")
h2("MANDATORY")
bullet("Plex token tMU7Gz1URMdPsdXw52it — was live on GitHub. Rotate in Plex + update config.")
bullet("GitHub PAT ghp_2MrTrx7… — was plaintext in Maia .git/config. Revoke+regen on GitHub; "
       "set up Git Credential Manager so nightly push still works.")
h2("STRONGLY RECOMMENDED")
bullet("Revoke 5 Telegram bot tokens via @BotFather /revoke (8610731233, 8271413034, "
       "8760042185, 8667837858, 8374766319) and update secrets/*.env.")
bullet("Maia app secret 4c3da67ba1144b3a782996562488a6b2 — rotate at source if still used.")
bullet("OpenClaw Anthropic key sk-ant-api03-iu1Hx… — rotate in Anthropic console if active.")
h2("VERIFY / OPTIONAL (not exposed in git)")
bullet("NEXUS keys (gitignored) — rotate only if you choose; no git exposure.")
bullet("TubeScout tubescout.db-wal key — confirm intended.")

h1("🔄 Next Up")
for x in [
    "Owner: complete the rotation checklist above (Plex + GitHub PAT first).",
    "Re-auth Maia git push (Credential Manager) after PAT rotation.",
    "Apply QI_baseline.gitignore to remaining repos in a routine pass.",
    "Optional: strip ~90 MB vendored ffmpeg DLLs from PersonalSong history.",
    "Optional: install real gitleaks for CI-grade scanning.",
]:
    bullet(x)

h1("📁 Documents / Files Updated")
for x in [
    r"C:\QIH\tools\secret_audit.py (new)",
    r"C:\QIH\tools\triage_secrets.py (new)",
    r"C:\QIH\tools\secret_gate.py (new)",
    r"C:\QIH\tools\personalsong_remediate.py (new)",
    r"C:\QIH\tools\nightly_git_sync.py (patched — secret gate)",
    r"C:\QI\TOOLS\maia_nightly_sync.py (patched — secret gate)",
    r"C:\QIH\ecosystem\QI_baseline.gitignore (new)",
    r"C:\QIH\logs\secret_audit\SECURITY_AUDIT_2026-06-16.md (report)",
    r"C:\QIH\logs\secret_audit\audit_latest.json / triage_latest.json (data)",
    r"C:\PersonalSong\.gitignore + config\app_config.template.json (remediation)",
    r"C:\QI\.git\config (PAT scrubbed)",
]:
    bullet(x)

doc.save(str(fn))
print("SAVED:", fn)
