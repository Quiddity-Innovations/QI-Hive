# -*- coding: utf-8 -*-
"""Generate the QI Hive session summary .docx for the War Room work (2026-06-18)."""
import sys
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")

OUT_DIR = Path(r"C:\QIH\shared\documentation\session_summaries")
stamp   = datetime.now().strftime("%Y-%m-%d_%H%M")
out     = OUT_DIR / f"Hive_Summary_{stamp}.docx"

doc = Document()
doc.add_heading("QI Hive — Session Summary", level=0)
p = doc.add_paragraph()
p.add_run("Date: ").bold = True
p.add_run("2026-06-18    ")
p.add_run("Title: ").bold = True
p.add_run("War Room reclaimed — Phase N Stage 0 (multi-agent chat) shipped")

doc.add_heading("✅ Completed This Session", level=1)
for t in [
    "Investigated what happened to the War Room idea — found /warroom had been built as a read-only monitoring board, NOT the 'all agents chat together' vision.",
    "Task 1 — Renamed the monitoring board to 'Mission Control' at /mission-control (render_mission_control in dashboard server.py); freed the War Room name.",
    "Task 3 — Shipped Phase N Stage 0: real multi-agent TEXT CHAT at /warroom (render_warroom_chat). Per-agent avatar bubbles, 4s auto-poll, Renne posts via input box.",
    "Created warroom_messages table in qi_brain.db (migration 2026_06_18_warroom_chat.sql, applied).",
    "Added Brain API endpoints POST/GET /api/warroom/message[s] (deployed in code, pending coordinated Brain restart).",
    "Added same-origin dashboard proxy routes /warroom/messages + /warroom/send (works through the public tunnel).",
    "Created shared poster engine/common/qi_warroom.py — post()/recent(); Brain HTTP with direct-DB fallback so any agent can join today.",
    "Task 2 — Drafted full Phase N spec (staged roadmap Stage 0 text → 1 avatars → 2 voice → 3 lip-sync video → 4 live) via hive-architect.",
    "Verified end-to-end: both pages return 200; Claude Code, Renne, and Architect all posted and appear in the room.",
    "Updated memory (project_phase_n_avatar_voice.md + index) to record Stage 0 shipped.",
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("⚠️ Operational Note / Caveat", level=1)
doc.add_paragraph(
    "QI_BrainAPI is a hard dependency ROOT — QI_NEXUS, QI_NayaBot, QI_MaiaBot and "
    "QI_Dashboard all DependOnService=QI_BrainAPI. An SCM stop cascade-stops the bots, "
    "so the new Brain HTTP endpoint cannot go live via a lone 'nssm restart'. It activates "
    "on the next coordinated Brain restart. The chat itself is fully live now via the "
    "direct-DB fallback. The QI_Elevate broker whitelist denies taskkill and 'nssm get'.",
    style="List Bullet")

doc.add_heading("🔄 Next Up", level=1)
for t in [
    "Stage 1: qi_avatars table + a static portrait per participant (SDXL-local recommended).",
    "Stage 2: per-agent TTS voice (edge-tts baseline, 9 distinct voices, free/no-GPU).",
    "Hand the spec to hive-inspector (Five/Six Laws + schema/naming) before hive-builder picks up Stage 1.",
    "On next maintenance window: coordinated Brain restart to activate the /api/warroom/message endpoint.",
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("🚀 In Development", level=1)
doc.add_paragraph("Phase N War Room — Stages 1-4 (avatars, voice, async selfie-video, live mode).", style="List Bullet")

doc.add_heading("🌅 Future Enhancements", level=1)
for t in [
    "Stage 3: async selfie-video messages via ComfyUI graph (LivePortrait default) — new QI_WarRoomRender service; gated on dispatch integration.",
    "Stage 4: live Teams/Zoom-feel real-time room — gated on dispatch + Claude Work handshake.",
]:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("📁 Documents / Files Updated", level=1)
for f in [
    r"C:\QIH\engine\hive\dashboard\server.py — Mission Control rename + War Room chat + proxy routes + RW helper",
    r"C:\QIH\engine\brain\api.py — /api/warroom/message[s] endpoints",
    r"C:\QIH\engine\brain\migrations\2026_06_18_warroom_chat.sql — NEW (warroom_messages table)",
    r"C:\QIH\engine\common\qi_warroom.py — NEW (shared multi-agent poster)",
    r"C:\QIH\shared\documentation\Phase_N_War_Room_Spec_2026-06-18.md — NEW (full staged spec)",
    r"C:\Users\renne\.claude\projects\C--CLAUDE\memory\project_phase_n_avatar_voice.md — Stage 0 update",
    r"C:\Users\renne\.claude\projects\C--CLAUDE\memory\MEMORY.md — index line refreshed",
]:
    doc.add_paragraph(f, style="List Bullet")

OUT_DIR.mkdir(parents=True, exist_ok=True)
doc.save(str(out))
print(out)
