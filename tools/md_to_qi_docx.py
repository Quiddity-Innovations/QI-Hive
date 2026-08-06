# -*- coding: utf-8 -*-
"""
md_to_qi_docx.py — render a Markdown file as a Word document in the
"QI Report" template v2 look (same template as the Project Status library).

Usage:
    python md_to_qi_docx.py <in.md> <out.docx> [--subtitle "Feature Guide"]

Reuses the palette/building blocks from
shared\documentation\project_library\_build_project_docs.py so the whole
library stays visually consistent. Supported Markdown: #/##/###/#### headings,
- / * bullets, 1. numbered lists, tables, ``` fenced code, **bold**, `code`,
[text](url) links (rendered as text + muted url), *italic-only* caption lines.
"""
from __future__ import annotations
import re, sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
sys.path.insert(0, str(Path(r"C:\QIH\shared\documentation\project_library")))
import _build_project_docs as T

from docx import Document
from docx.shared import Pt, Inches


def _inline(par, text):
    """**bold**, `code`, [text](url) — link urls rendered muted in parens."""
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    T.add_inline(par, text)


def cover(doc, title, subtitle, src: Path):
    T.accent_bar(doc, sz=48)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
    r = p.add_run("Q U I D D I T Y   I N N O V A T I O N S")
    r.font.name = T.HEAD_FONT; r.font.size = Pt(10); r.bold = True
    r.font.color.rgb = T._rgb(T.MUTED)
    p = doc.add_paragraph()
    r = p.add_run("QI Hive · Project Documentation Library")
    r.font.name = T.HEAD_FONT; r.font.size = Pt(10); r.font.color.rgb = T._rgb(T.ACCENT)
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.font.name = T.HEAD_FONT; r.font.size = Pt(32); r.bold = True
    r.font.color.rgb = T._rgb(T.INK)
    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    r.font.name = T.HEAD_FONT; r.font.size = Pt(16); r.font.color.rgb = T._rgb(T.ACCENT)
    T.para_border(p, "bottom", sz=10, color=T.BORDER, space=10)
    for _ in range(3):
        doc.add_paragraph()
    meta = [("Generated", T.DATE), ("Source", str(src)),
            ("Classification", "Internal — QI Ecosystem")]
    t = doc.add_table(rows=0, cols=2)
    for k, v in meta:
        c = t.add_row().cells
        r = c[0].paragraphs[0].add_run(k.upper())
        r.font.name = T.HEAD_FONT; r.font.size = Pt(8); r.bold = True
        r.font.color.rgb = T._rgb(T.MUTED)
        r = c[1].paragraphs[0].add_run(v)
        r.font.size = Pt(9.5)
        if k == "Source":
            r.font.name = T.MONO_FONT; r.font.size = Pt(8.5)
    t.columns[0].width = Inches(1.5); t.columns[1].width = Inches(5.1)
    for row in t.rows:
        row.cells[0].width = Inches(1.5); row.cells[1].width = Inches(5.1)
    T.table_borders(t, {"insideH": (4, T.BORDER)})
    T.table_cell_margins(t, top=60, bottom=60, left=0, right=90)
    doc.add_page_break()


def render_md(doc, lines):
    tbl, code, in_code = [], [], False
    first_h1_skipped = False
    for raw in lines:
        s = raw.rstrip("\n")
        if s.strip().startswith("```"):
            if in_code:
                T.add_code_block(doc, "\n".join(code)); code = []
            in_code = not in_code
            continue
        if in_code:
            code.append(s); continue
        if s.lstrip().startswith("|"):
            tbl.append(s); continue
        if tbl:
            T.flush_md_table(doc, tbl); tbl = []
        if s.startswith("# "):
            if not first_h1_skipped:      # title already on the cover
                first_h1_skipped = True
                continue
            T.h1(doc, s[2:].strip())
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
        elif s.startswith("#### "):
            doc.add_heading(s[5:].strip(), level=3)
        elif re.match(r"^\s*[-*]\s+", s):
            p = doc.add_paragraph(style="List Bullet")
            _inline(p, re.sub(r"^\s*[-*]\s+", "", s))
        elif re.match(r"^\s*\d+\.\s+", s):
            p = doc.add_paragraph(style="List Number")
            _inline(p, re.sub(r"^\s*\d+\.\s+", "", s))
        elif s.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            T.para_border(p, "bottom", sz=6, color=T.BORDER, space=2)
        elif (s.strip().startswith("*") and s.strip().endswith("*")
              and not s.strip().startswith("**")):
            p = doc.add_paragraph()
            r = p.add_run(s.strip().strip("*")); r.italic = True
            r.font.size = Pt(8.5); r.font.color.rgb = T._rgb(T.MUTED)
        elif s.strip():
            p = doc.add_paragraph(); _inline(p, s)
    if tbl:
        T.flush_md_table(doc, tbl)


def convert(md_path: Path, out_path: Path, subtitle: str):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    title = next((l[2:].strip() for l in lines if l.startswith("# ")), md_path.stem)
    doc = Document()
    T.set_base_styles(doc)
    T.setup_page(doc, title, header_text=title)
    cover(doc, title, subtitle, md_path)
    render_md(doc, lines)
    doc.save(str(out_path))
    return title


if __name__ == "__main__":
    args = [a for a in sys.argv[1:] if not a.startswith("--")]
    sub = "Reference Document"
    if "--subtitle" in sys.argv:
        sub = sys.argv[sys.argv.index("--subtitle") + 1]
    src, dst = Path(args[0]), Path(args[1])
    t = convert(src, dst, sub)
    print(f"OK  {t}: {dst}")
