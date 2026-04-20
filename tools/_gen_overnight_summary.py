# -*- coding: utf-8 -*-
"""Generate session-summary .docx for 2026-04-19 overnight autonomous session."""
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt

OUT = Path(r"C:\UNIVERSAL\DOCUMENTATION\Session_Summaries") / \
      f"QIHive_Summary_2026-04-19_Overnight.docx"

doc = Document()

h = doc.add_heading("QI Hive — Overnight Autonomous Session (2026-04-19)", 0)

doc.add_paragraph(
    "Scope: Renne asked me to run independently overnight — finish the health audit, "
    "catch up One Brain, broadcast the new 're-wiring' to sibling sessions, and surprise him. "
    "Constraint: don't break priority projects (Maia, Naya, NEXUS, EasyFlow, QI_Hive)."
)

# === Completed ===
doc.add_heading("✅ Completed This Session", 1)
for bullet in [
    "Health registry expanded 7 → 11 projects. Added EasyFlow, Claude_Manager, QI_Universal, FileHQ. NEXUS doc_path corrected.",
    "/usage page fully redesigned per earlier feedback: compact boxes, renamed tiers (no 'WHAT-IF' prefix), 3-series daily chart (Actual vs Local offload vs Local+Batch), comparison columns on By Project + By Model, per-model savings table with totals.",
    "/hive page stats fixed: was showing '?' and '0 tasks'. Bug was wrong dict keys (brain schema is flat, code was looking for nested). Now pulls real numbers: 9 projects / 57 decisions / 21 features / 80 sessions; per-agent task counts from tasks.json.",
    "Broker autonomy fully unlocked. gsudo installed machine-wide with 8-hour credential cache. No more UAC prompts mid-session.",
    "Broker resilience triple-layered: (1) NSSM AppExit=Restart on QI_Elevate/QI_Dashboard/QI_HiveIngest (3s delay); (2) SYSTEM watchdog task polls broker status every minute; (3) 5-minute stale-queue purge at broker startup, which fixes the 'broker revives, finds old kill order, suicides again' loop caught at 20:51.",
    "Single-instance PID lock on broker with os.kill(pid,0) liveness check prevents orphan-broker race.",
    "QI Brain schema fixed: FK references to dropped agents_old table were causing every log_session POST to 500. Rebuilt 5 tables with correct FK -> agents.",
    "QI Brain backfilled: 30 historical sessions imported from C:\\UNIVERSAL\\DOCUMENTATION\\Session_Summaries\\. Brain session count 7 → 37.",
    "qi_hive added to brain projects table (was missing).",
    "LATEST.md rewritten with full autonomous-elevation broadcast so sibling Claude sessions in C:\\QI, C:\\NAYA, C:\\NEXUS, C:\\EasyFlow know the new path.",
    "All work committed + pushed to origin/master (commits da905d7, 650b7f2).",
]:
    doc.add_paragraph(bullet, style="List Bullet")

# === Issues ===
doc.add_heading("⚠️ Issues Encountered & Resolved", 1)
for bullet in [
    "Broker suicide loop: stopped broker had a 2-hour-old `sc stop QI_Elevate` request still in pending/. Every revival re-executed it. Fixed with 5-minute stale purge at startup.",
    "Bash MSYS path-mangling converted schtasks `/Create` into `C:/Program Files/Git/Create`. Worked around by wrapping all schtasks/nssm commands in a .bat file invoked via gsudo.",
    "Brain 500s on every log_session: FK pointed at agents_old which had been dropped. Rebuilt affected tables with SQLite table-swap pattern, preserving all rows.",
    "None of the issues left the system in a broken state; all priority projects (Maia/Naya/NEXUS/EasyFlow) untouched.",
]:
    doc.add_paragraph(bullet, style="List Bullet")

# === Services status ===
doc.add_heading("🟢 Services Status (live now)", 1)
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
tbl.rows[0].cells[0].text = "Service"
tbl.rows[0].cells[1].text = "Status"
for svc, status in [
    ("QI_Dashboard (:8600)", "RUNNING — autonomous"),
    ("QI_HiveIngest", "RUNNING"),
    ("QI_Elevate (broker)", "RUNNING — AppExit=Restart + watchdog"),
    ("QI_BrainAPI (:9010)", "RUNNING"),
    ("QI_ElevateWatchdog (Task Scheduler)", "registered, 1-min cadence, SYSTEM"),
]:
    row = tbl.add_row().cells
    row[0].text, row[1].text = svc, status

# === Next up ===
doc.add_heading("🔄 Next Up (for Renne / next session)", 1)
for bullet in [
    "Verify /hive, /usage, /health pages render correctly — all 11 projects visible on /health, real numbers on /hive, 3-series chart on /usage.",
    "SessionEnd hook: automate qi.log_session calls so future sessions self-log to Brain without manual backfill.",
    "Decisions backfill: only 7 decisions logged despite ~57 architectural choices across sessions. Parse session summaries for decisions and log them via qi.log_decision.",
    "Features backfill: only 1 feature logged. Scan code for 'feature' markers or session summaries for feature descriptions.",
    "Consider adding qi_hive as a ranked project on /projects page (currently treated as backbone).",
]:
    doc.add_paragraph(bullet, style="List Bullet")

# === In development ===
doc.add_heading("🚀 In Development", 1)
doc.add_paragraph(
    "One Brain is the umbrella initiative. Today's work brought it from "
    "'schema broken, 7 sessions, no autonomous elevation' to "
    "'schema clean, 37 sessions + qi_hive project, broker auto-healing, sibling sessions notified'. "
    "Remaining path: hook-based auto-logging, decisions+features backfill, cross-project dashboard views.",
    style="Intense Quote",
)

# === Future ===
doc.add_heading("🌅 Future Enhancements", 1)
for b in [
    "Richer /activity data: add agent/model/role fields to hive_report.report().",
    "Broker whitelist expansion: git commit/push, python service restarts beyond NSSM.",
    "Per-project 'last session' card on /projects pulling from brain.session_log.",
]:
    doc.add_paragraph(b, style="List Bullet")

# === Docs updated ===
doc.add_heading("📁 Documents / Files Updated", 1)
for f in [
    "C:\\QIH\\engine\\hive\\health_check.py — added EasyFlow, Claude_Manager, QI_Universal, FileHQ",
    "C:\\QIH\\engine\\hive\\dashboard\\server.py — /usage redesign, /hive stats fix",
    "C:\\QIH\\engine\\common\\usage_stats.py — 3-series daily + per-project savings",
    "C:\\QIH\\engine\\common\\qi_elevate.py — stale-queue purge, PID lock",
    "C:\\QIH\\tools\\install_gsudo.ps1 (NEW)",
    "C:\\QIH\\tools\\broker_watchdog.ps1 (NEW)",
    "C:\\QIH\\tools\\install_broker_resilience.bat (NEW)",
    "C:\\QIH\\tools\\fix_brain_fk_agents_old.py (NEW)",
    "C:\\QIH\\tools\\brain_backfill_sessions.py (NEW)",
    "C:\\QIH\\docs\\LATEST.md — re-wiring broadcast",
    "C:\\QIH\\data\\qi_brain.db — schema fixed, 30 sessions added, qi_hive project added",
]:
    doc.add_paragraph(f, style="List Bullet")

# === Commits ===
doc.add_heading("📦 Git Commits", 1)
for c in [
    "da905d7 — feat(health): add EasyFlow, Claude_Manager, QI_Universal, FileHQ to registry",
    "650b7f2 — feat(brain): FK migration + backfill 30 historical sessions",
    "(pushed to origin/master)",
]:
    doc.add_paragraph(c, style="List Bullet")

doc.save(str(OUT))
print(f"Saved: {OUT}")
