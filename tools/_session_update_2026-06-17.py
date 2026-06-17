# -*- coding: utf-8 -*-
"""Update-all documentation pass for the 2026-06-17 CoWork Dispatch cleanup session."""
import sys
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")

DOC_DIR = Path(r"C:\QIH\hive\DOCUMENTATION")
SUMMARY_DIR = Path(r"C:\QIH\shared\documentation\session_summaries")
TODAY = "2026-06-17"
STAMP = datetime.now().strftime("%Y-%m-%d %H:%M")
TITLE = "CoWork Dispatch — Pending queue cleanup & Claude approval authority"

COMPLETED = [
    "Diagnosed the never-emptying 'Pending — Awaiting Review' box on the CoWork Dispatch board: "
    "2,844 pending dispatches, all hive_inspector compliance findings collapsing to ~29 distinct issues.",
    "Root cause: inspector file_dispatch() inserted a brand-new pending dispatch for the same "
    "(project_id, check_id) on every 4-hour run with no dedup, and never resolved old ones.",
    "Confirmed Renne's genuine CoWork approvals had succeeded — zero real CoWork/Renne/Claude items were stuck.",
    "Fix 1: made inspector.file_dispatch() idempotent — one open dispatch per project+check, "
    "refreshed in place instead of duplicated. Verified by unit test (same id returned, single row, payload refreshed).",
    "Fix 2: bulk-resolved all 2,844 stale inspector dispatches (status='resolved', NOT 'approved', "
    "so no apply-pipeline runs were queued). Pending count now 0.",
    "Fix 3: dashboard render_dispatch() now excludes inspector/compliance from the CoWork human-review "
    "queue; they are routed to the /compliance board with an inline pointer. Board shows 'showing 0 of 0'.",
    "Fix 4: created tools/dispatch_admin.py so Claude can approve/decline/resolve dispatches on Renne's "
    "behalf via the Brain API (approve triggers the real apply pipeline).",
    "Restarted QI_Dashboard via the QI_Elevate broker (nssm restart required elevation).",
]
NEXT_UP = [
    "Watch the next scheduled inspector run to confirm it re-raises exactly ONE dispatch per live finding (no flood).",
    "Decide real disposition for the recurring findings (gitignore_secrets / session_freshness / docs_no_decoys) per project.",
    "Consider auto-resolving inspector dispatches whose underlying condition is fixed (close-the-loop on /compliance).",
    "Optionally add a dashboard badge linking CoWork Dispatch ↔ /compliance counts.",
]
IN_DEV = [
    "Inspector Inbox Auto-Drain (deterministic auto-verdict worker) — design exists; separate from this dispatch-queue fix.",
]
FILES = [
    r"C:\QIH\engine\hive\inspector\inspector.py — idempotent file_dispatch()",
    r"C:\QIH\engine\hive\dashboard\server.py — render_dispatch() excludes inspector/compliance from human queue",
    r"C:\QIH\tools\dispatch_admin.py — NEW: Claude/Renne dispatch approval CLI",
    r"C:\QIH\data\qi_brain.db — 2,844 stale inspector dispatches resolved",
]


def append_log(path: Path, heading: str, bullets, version_note=None):
    if not path.exists():
        print(f"  ! missing, skipping: {path.name}")
        return
    doc = Document(str(path))
    doc.add_heading(f"{heading} — {STAMP}", level=1)
    p = doc.add_paragraph()
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(12)
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")
    if version_note:
        doc.add_paragraph(version_note)
    doc.save(str(path))
    print(f"  ✓ updated {path.name}")


# 1-3: project logs
append_log(DOC_DIR / "Claude_Manager_Implementation_Log.docx", "Implementation Log", COMPLETED)
append_log(DOC_DIR / "Claude_Manager_Meeting_Minutes.docx", "Meeting Minutes",
           ["Decision: inspector compliance findings are a separate channel (/compliance) and must NOT "
            "pollute the CoWork Dispatch human-review queue.",
            "Decision: stale inspector dispatches resolved (not approved) to avoid mass apply-pipeline runs.",
            "Decision: Claude granted dispatch approval authority via tools/dispatch_admin.py.",
            "Next steps: " + "; ".join(NEXT_UP)])
append_log(DOC_DIR / "Claude_Manager_Version_History.docx", "Version History",
           ["inspector.py: file_dispatch() now idempotent (dedup by project_id+check_id while open).",
            "dashboard server.py: render_dispatch() filters out source=hive_inspector / type=compliance.",
            "tools/dispatch_admin.py added (v1)."])

# 4: session summary
SUMMARY_DIR.mkdir(parents=True, exist_ok=True)
sdoc = Document()
sdoc.add_heading(f"QI Hive — {TITLE}", level=0)
sdoc.add_paragraph(f"Date: {TODAY}  |  Saved: {STAMP}")

sdoc.add_heading("✅ Completed This Session", level=1)
for b in COMPLETED:
    sdoc.add_paragraph(b, style="List Bullet")

sdoc.add_heading("🔄 Next Up", level=1)
for b in NEXT_UP:
    sdoc.add_paragraph(b, style="List Bullet")

sdoc.add_heading("🚀 In Development", level=1)
for b in IN_DEV:
    sdoc.add_paragraph(b, style="List Bullet")

sdoc.add_heading("🌅 Future Enhancements", level=1)
sdoc.add_paragraph("Auto-close inspector dispatches when the underlying condition is remediated.", style="List Bullet")
sdoc.add_paragraph("Unify CoWork Dispatch and /compliance into one operator triage view.", style="List Bullet")

sdoc.add_heading("📁 Documents Updated", level=1)
for f in FILES:
    sdoc.add_paragraph(f, style="List Bullet")

out = SUMMARY_DIR / f"QIHive_Summary_{TODAY}_{datetime.now().strftime('%H%M')}.docx"
sdoc.save(str(out))
print(f"\nSESSION SUMMARY SAVED: {out}")
