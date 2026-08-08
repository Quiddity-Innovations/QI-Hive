# -*- coding: utf-8 -*-
"""
Build the BU Hive Starter package — the portable documentation + prompt that
seeds (and levels-up) the BU Hive development control plane on the BU laptop.

Outputs into C:\\QIH\\docs\\BU_Hive_Starter\\:
  - BU_Hive_Build_Guide.docx   (formatted spec, with QI Hive visual reference shots)
  - BU_Hive_Build_Guide.md     (markdown companion)
  - BU_Hive_Master_Build_Prompt.md  (standalone, paste-ready prompt)
  - README.md                  (how to use this package on the BU laptop)
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT   = Path(r"C:\QIH\docs\BU_Hive_Starter")
SHOTS = OUT / "reference_shots"
TODAY = "2026-06-19"

RED   = RGBColor(0xB3, 0x1B, 0x1B)   # BU crimson-ish accent
INK   = RGBColor(0x22, 0x26, 0x2B)
GREY  = RGBColor(0x6B, 0x72, 0x80)

# ============================================================ content ========
VISION = (
    "BU Hive is your personal development control plane at BU — a single, self-contained "
    "web application that becomes the home base for your development work there. Its first "
    "mission is to take CogniBase from its current MOC (proof-of-concept) stage to a real, "
    "production-grade application; from there it serves as the hub for your other BU "
    "projects. It is modelled on the QI Hive (the sibling control plane on your personal "
    "machine), but it stands completely on its own: it runs only on the BU machine, depends "
    "on nothing outside its own folder, and hardcodes no personal paths — so it can be moved "
    "to its rightful home and stay independent of your personal PC."
)

PILLARS = [
    ("Feature parity with QI Hive",
     "Not a single static page — the full control-plane page set below (dashboard, launcher, "
     "projects, services, health, logs, tests, board, activity, library, brain, config, guide)."),
    ("Visual & UX polish",
     "An AdminLTE shell with Dark/Light/System themes, charts, and a consistent sidebar + top "
     "bar + footer — a professional look that matches the QI Hive reference screenshots in this guide."),
    ("Live data integrations",
     "Wired to real systems — SQL Server, OnBase, file shares and (optionally) a local LLM — each "
     "with a health indicator. Nothing hardcoded; every connection comes from config."),
    ("Knowledge / AI layer",
     "A local knowledge substrate (bu_brain.db) that logs decisions, features and sessions; "
     "semantic-or-full-text search across your docs; and an assistant chat scoped to your BU projects."),
]

PORTABILITY = [
    "One folder holds everything — copy it to any machine and it runs.",
    "A single config.json (or a BU_HIVE_HOME env var) defines the root, port and data paths; no hardcoded absolute paths anywhere.",
    "Its own bu_registry.json and bu_brain.db — no dependency on the QI ecosystem or your personal PC.",
    "Vendored static assets (no CDN) so it works on a locked-down network.",
    "Pin dependencies in requirements.txt and document the exact Python version.",
]

SECURITY = [
    "Bind to 127.0.0.1 only — never expose it; no Cloudflare / ngrok / public tunnels.",
    "Respect BU IT policy for services: if NSSM/Windows services aren't permitted, run via a Startup-folder launcher or Task Scheduler instead.",
    "Secrets (DB connection strings, OnBase credentials) live in a gitignored .env — never in code or the registry.",
    "All reporting data access is read-only by default; honour BU data governance.",
    "Works fully offline; no telemetry. If embeddings/Ollama can't be installed, fall back to SQLite FTS5 full-text search.",
]

CORE_PAGES = [
    ("Dashboard", "/", "At-a-glance home: project cards with status, recent activity, quick links and today's focus."),
    ("Launcher", "/launcher", "Click-to-open every BU project's UI/API, each with a live status dot."),
    ("Projects", "/projects", "Per-project detail pages (CogniBase + others), rendered from each project's INTRO files."),
    ("Services", "/services", "Local run/process inventory with start/stop/status (NSSM, or a policy-friendly launcher fallback)."),
    ("Health", "/health", "On-demand scan of each project — process up, port listening, git clean, data source reachable; serves JSON for monitors."),
    ("Logs", "/logs", "Centralised log viewer: project/file picker, large-file tail, level colouring."),
    ("Tests", "/tests", "Run smoke/API/UI suites; pass/fail tiles; failures auto-file board tasks."),
    ("Task Board", "/board", "Drag-and-drop Kanban (Backlog → In Progress → Review → Done), persisted to the DB."),
    ("Activity", "/activity", "Session and audit log of your dev work — what changed, when, and notable events."),
    ("Library", "/library", "Documentation Brain: search across your BU docs/specs with filters and a relationship graph."),
    ("BU Brain", "/brain", "Knowledge substrate UI: Overview, Decisions, Features, Sessions, Search."),
    ("Config", "/config", "Data-source connections, log levels, theme and paths."),
    ("Guide", "/guide", "The operator manual, rendered from Markdown with a raw download."),
]

BU_PAGES = [
    ("CogniBase Workbench", "/cognibase",
     "The flagship. A dedicated multi-panel page to develop CogniBase from MOC to a real tool — detailed below."),
    ("Data Sources", "/data",
     "Registry + live health of every configured connection (SQL Server, OnBase, file shares, LLM endpoint)."),
    ("Assistant", "/assistant",
     "A local AI chat scoped to your BU project context (uses a local or configured LLM; degrades to 'offline' if none)."),
    ("LLM Usage", "/usage",
     "Optional — token and cost tracking if/when you use cloud or local LLMs for CogniBase RAG."),
]

COGNIBASE_PANELS = [
    "OnBase connection status + a document-type browser + a retrieval test.",
    "Ingestion pipeline view: documents in → chunk → embed → index, with live counts and a re-index button.",
    "Index health: vector index size / last build / coverage — or FTS index stats if embeddings aren't available.",
    "RAG query tester: type a question → see the retrieved chunks + an answer, with the keyword/SQL fallback path shown.",
    "Ad-hoc SQL reporting console: pick a saved query or write read-only SQL against a configured BU source; results in a table with CSV export.",
    "Everything parameterised from config/.env — no live credentials in code, read-only by default.",
]

OPTIONAL_PAGES = [
    ("Tunnels", "Disabled on the BU network for security — BU Hive is local-only."),
    ("Headlines / News", "External feeds; enable only if outbound access is permitted and useful."),
    ("Mission Control / War Room / CoWork Dispatch", "Multi-agent orchestration from QI Hive — collapse into Assistant + Activity until you actually need it."),
    ("Scheduled Tasks", "Add when you have recurring BU jobs (re-index, nightly reports)."),
    ("Compliance", "Repurpose later for BU coding/standards checks if useful."),
]

CONVENTIONS = [
    "Service / process names: BU_<Project><Role> (e.g. BU_Dashboard, BU_CogniBaseAPI); always set a description and working directory (or the launcher equivalent).",
    "Port discipline: pick one free block for BU projects; default the dashboard to 8700 (verify it's free and allowed); read every port from the registry — never hardcode.",
    "UTF-8 everywhere; read-only pages never mutate state; vendored assets only.",
    "Write scripts to files (no fragile shell heredocs); pin and document dependencies.",
]

ROADMAP = [
    ("Phase 1 — The standing control plane",
     "FastAPI + AdminLTE shell (sidebar/topbar/themes), config.json, bu_registry.json, and the Dashboard, Launcher, Projects, Services, Health, Logs and Guide pages. Goal: the hub stands up and shows your projects."),
    ("Phase 2 — The dev workflow + knowledge base",
     "Task Board, Tests, Activity, BU Brain and Library. Goal: you can plan, test, and accumulate knowledge as you work."),
    ("Phase 3 — CogniBase Workbench",
     "OnBase panel, ingestion pipeline, index health, RAG query tester and the SQL reporting console. Goal: CogniBase development happens inside BU Hive."),
    ("Phase 4 — AI layer + polish",
     "The Assistant chat, semantic/FTS search across everything, charts and visual polish, and LLM-usage tracking if relevant. Goal: the sophistication targets are met."),
]

REF_SHOTS = [
    ("ref_dashboard.png", "QI Hive — Dashboard. The look & feel to match: bento tiles, a status table, calm dark theme."),
    ("ref_hive.png", "QI Hive — agent roster + ecosystem stats. The model for BU Hive's project/agent panels."),
    ("ref_brain.png", "QI Hive — Brain (knowledge substrate). The model for BU Brain (decisions / features / sessions / search)."),
    ("ref_library.png", "QI Hive — Library / Documentation Brain. The model for BU Hive's doc search + knowledge graph."),
]

# ----- the standalone master build prompt (also written to its own .md) ------
PROMPT = """You are a senior full-stack engineer. Build BU Hive - a personal development control
plane and project hub for my work at BU (a corporate / enterprise environment). Its first
mission is to take CogniBase (an OnBase vector-search + ad-hoc SQL-reporting tool) from its
current MOC (proof-of-concept) stage to a real, production-grade application, and to serve as
the home base for my other BU development projects.

It is modelled on a sibling system ("QI Hive") but must be fully SELF-CONTAINED and PORTABLE:
it runs only on this machine, depends on nothing outside its own folder, and hardcodes no
personal paths.

GOAL
One self-contained Python web application that renders my BU projects - their status,
services, logs, tests, tasks, docs and a shared knowledge base - onto server-rendered pages,
plus a dedicated CogniBase workbench, reading live from a local registry and a local knowledge
DB, and degrading gracefully when a data source is offline.

ENVIRONMENT & SECURITY (this is a corporate machine - treat as strict)
- Bind to 127.0.0.1 only. No public tunnels, no external exposure, no telemetry.
- Respect BU IT policy: if Windows services aren't permitted, run via a Startup-folder
  launcher or Task Scheduler instead of NSSM.
- Secrets (DB connection strings, OnBase creds) live in a gitignored .env, never in code.
- All reporting data access is read-only by default; honour BU data governance.
- Everything works offline. If embeddings / Ollama can't be installed, fall back to SQLite
  FTS5 full-text search.

STACK & SHELL
- Python 3.11+, FastAPI, Uvicorn. One app serves all routes, server-rendered HTML (no SPA).
- UI: AdminLTE 4 + Bootstrap 5 + Bootstrap Icons, all vendored under /static (no CDN).
  Dark default + Light + System themes.
- Heavy interactive tools may be embedded as Streamlit/Gradio sub-apps behind the same nav
  when that is faster than hand-rolling them.
- Shared layout: fixed left sidebar (nav + status legend), top bar (clock, theme switch,
  health shortcut), footer.
- A single config (config.json or env BU_HIVE_HOME) sets the root, port (default 8700 -
  verify it is free and allowed) and data paths. No hardcoded absolute paths.

DATA SOURCES (read live; never crash if missing)
1. bu_registry.json - my BU projects: identity, paths, ports, run-commands, data connections.
2. bu_brain.db (SQLite; optional local vector index) - sessions, decisions, features, project
   state, a document catalogue, and searchable memory (FTS5, or vector if available).

CORE PAGES (parity with the sibling control plane)
1.  Dashboard (/) - at-a-glance: project cards w/ status, recent activity, quick links, today's work.
2.  Launcher (/launcher) - click-to-open every BU project's UI/API with a live status dot.
3.  Projects (/projects) - per-project detail pages (CogniBase + others) from each project's INTRO files.
4.  Services (/services) - local run/process inventory with start/stop/status (NSSM or fallback launcher).
5.  Health (/health) - on-demand scan of each project (process up, port listening, git clean,
    data-source reachable); content-negotiated JSON for probes.
6.  Logs (/logs) - centralised log viewer with project/file picker, large-file tail, level colouring.
7.  Tests (/tests) - run smoke/API/UI suites; pass/fail tiles; failures file board tasks.
8.  Task Board (/board) - drag-drop Kanban (Backlog->In Progress->Review->Done), persisted.
9.  Activity (/activity) - session/audit log of dev work (turns, duration, notable events).
10. Library (/library) - "Documentation Brain": search across my BU docs/specs with filters +
    a relationship graph; index built locally.
11. BU Brain (/brain) - knowledge substrate UI: Overview, Decisions, Features, Sessions, Search.
12. Config (/config) - data-source connections, log levels, theme, paths.
13. Guide (/guide) - the operator manual rendered from Markdown.

COGNIBASE WORKBENCH (the flagship - /cognibase)
A dedicated multi-panel page to develop CogniBase from MOC to real:
- OnBase connection status + document-type browser + retrieval test.
- Ingestion pipeline view: documents in -> chunked -> embedded -> indexed, with counts and a re-index button.
- Index health: vector index size / last build / coverage (or FTS index stats if no embeddings).
- RAG query tester: ask a question -> see retrieved chunks + an answer, with the keyword/SQL fallback shown.
- Ad-hoc SQL reporting console: a saved query or read-only SQL against a configured BU source;
  results table with CSV export.
- Everything parameterised by config; no live credentials in code.

KNOWLEDGE / AI LAYER
- A local assistant/chat panel scoped to my BU project context (local or configured LLM;
  degrade to "offline" if none).
- Semantic or full-text search across docs, decisions, features and sessions.
- Log every dev session, decision and feature to bu_brain.db so the hub gets smarter over time.

CONVENTIONS
- Service/process names: BU_<Project><Role>; always set a description and working directory
  (or the launcher equivalent).
- One port block for BU projects; never hardcode - read the registry; default dashboard 8700.
- UTF-8 everywhere; read-only pages never mutate state; vendored assets only.

BUILD IN PHASES (ship each before the next)
1. Shell + registry + Dashboard/Launcher/Projects/Services/Health/Logs/Guide (the standing control plane).
2. Task Board + Tests + Activity + BU Brain + Library (the dev workflow + knowledge base).
3. CogniBase Workbench (OnBase panel, ingestion, index, RAG tester, SQL reporting).
4. AI assistant + semantic search + visual polish (charts, theming) + LLM usage if relevant.

DELIVERABLE
A runnable FastAPI app + vendored static assets + a config.json + a start script (and a
service/launcher install appropriate to BU policy), self-contained in one folder, faithful to
the pages and constraints above. Use the accompanying BU Hive Build Guide as the detailed spec
and the QI Hive screenshots as the visual target."""

# ============================================================ helpers ========
def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_prompt_block(doc, text):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]; set_cell_bg(cell, "FBF0F0")  # pale crimson
    cell.paragraphs[0].text = ""
    first = True
    for line in text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph(); first = False
        pf = p.paragraph_format; pf.space_after = Pt(0); pf.space_before = Pt(0); pf.line_spacing = 1.0
        r = p.add_run(line if line else ""); r.font.name = 'Consolas'; r.font.size = Pt(8.5); r.font.color.rgb = INK

def two_col_table(doc, header, rows, widths=(2.0, 4.4)):
    tbl = doc.add_table(rows=1, cols=2); tbl.style = 'Light Grid Accent 2'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = tbl.rows[0].cells[i]; c.paragraphs[0].add_run(h).bold = True
        set_cell_bg(c, "B31B1B"); c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for a, b in rows:
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].add_run(a).bold = True
        cells[1].paragraphs[0].add_run(b)
    return tbl

def three_col_table(doc, header, rows):
    tbl = doc.add_table(rows=1, cols=3); tbl.style = 'Light Grid Accent 2'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        c = tbl.rows[0].cells[i]; c.paragraphs[0].add_run(h).bold = True
        set_cell_bg(c, "B31B1B"); c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for a, b, cc in rows:
        cells = tbl.add_row().cells
        rp = cells[0].paragraphs[0]; rp.add_run(a).bold = True
        cells[1].paragraphs[0].add_run(b).italic = True
        cells[2].paragraphs[0].add_run(cc)
    return tbl

# ============================================================ docx ===========
def build_docx():
    doc = Document()
    st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)

    # cover
    for _ in range(3): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BU Hive"); r.bold = True; r.font.size = Pt(42); r.font.color.rgb = RED
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Build Guide & Master Prompt"); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = INK
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("A portable spec for a self-contained development control plane at BU")
    r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREY
    for _ in range(2): doc.add_paragraph()
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in ["Renne Santiago", "Flagship: CogniBase (MOC → production)  ·  modelled on QI Hive",
                 f"Generated {TODAY}  ·  carry this package to the BU laptop"]:
        rr = meta.add_run(line + "\n"); rr.font.size = Pt(11); rr.font.color.rgb = INK
    doc.add_page_break()

    # vision
    doc.add_heading("1. Purpose & vision", level=1)
    doc.add_paragraph(VISION)
    doc.add_paragraph()
    note = doc.add_paragraph()
    nr = note.add_run("Independence is the point. ")
    nr.bold = True; nr.font.color.rgb = RED
    note.add_run("This package is the last piece needed to send BU Hive to its rightful home — "
                 "a clean-room build kit that needs nothing from your personal PC.")

    # pillars
    doc.add_heading("2. The four sophistication pillars", level=1)
    doc.add_paragraph("Every decision below serves these four goals (all in scope):")
    for name, desc in PILLARS:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(name + " — ").bold = True; p.add_run(desc)

    # architecture
    doc.add_heading("3. Architecture & stack", level=1)
    for line in [
        "One self-contained Python web app: FastAPI + Uvicorn, server-rendered HTML (no SPA framework).",
        "UI: AdminLTE 4 + Bootstrap 5 + Bootstrap Icons, vendored under /static (no CDN). Dark default + Light + System themes.",
        "Adaptable: heavy interactive tools (e.g. a CogniBase report builder) may be embedded as Streamlit/Gradio sub-apps behind the same nav when that is faster than hand-rolling — the shell stays FastAPI.",
        "Two live data sources read on every request: bu_registry.json (projects/ports/services/connections) and bu_brain.db (knowledge substrate).",
        "Shared layout on every page: fixed left sidebar (nav + status legend), top bar (clock, theme switch, health shortcut), footer.",
    ]:
        doc.add_paragraph(line, style='List Bullet')

    # portability
    doc.add_heading("4. Portability & independence rules", level=1)
    for line in PORTABILITY: doc.add_paragraph(line, style='List Bullet')

    # security
    doc.add_heading("5. BU environment & security", level=1)
    doc.add_paragraph("BU Hive runs on a corporate machine — treat the constraints as strict, not optional:")
    for line in SECURITY: doc.add_paragraph(line, style='List Bullet')
    doc.add_page_break()

    # page catalog
    doc.add_heading("6. Page catalogue", level=1)
    doc.add_heading("6.1  Core pages (parity with QI Hive)", level=2)
    three_col_table(doc, ["Page", "Route", "What it does"],
                    [(n, r, d) for n, r, d in CORE_PAGES])
    doc.add_paragraph()
    doc.add_heading("6.2  BU-specific pages (the differentiators)", level=2)
    three_col_table(doc, ["Page", "Route", "What it does"],
                    [(n, r, d) for n, r, d in BU_PAGES])
    doc.add_paragraph()
    doc.add_heading("6.3  Optional / off by default", level=2)
    two_col_table(doc, ["Page", "Why it's optional"], OPTIONAL_PAGES)
    doc.add_page_break()

    # cognibase deep dive
    doc.add_heading("7. CogniBase Workbench (the flagship)", level=1)
    doc.add_paragraph(
        "CogniBase is the reason BU Hive exists first. The /cognibase page is a dedicated, "
        "multi-panel workbench to drive it from MOC to a real tool. Panels:")
    for line in COGNIBASE_PANELS: doc.add_paragraph(line, style='List Bullet')
    doc.add_paragraph()
    tip = doc.add_paragraph()
    tip.add_run("Fallback by design: ").bold = True
    tip.add_run("if the BU machine can't run embeddings/Ollama, the workbench uses SQLite FTS5 "
                "keyword search and shows that path explicitly — CogniBase still works, just without vectors.")

    # data integrations
    doc.add_heading("8. Live data integrations", level=1)
    two_col_table(doc, ["Source", "How BU Hive uses it"], [
        ("SQL Server (BU)", "Read-only ad-hoc reporting + saved queries; connection from .env; results to table/CSV."),
        ("OnBase (Hyland)", "Document-type browse, retrieval test, and the ingestion source for CogniBase."),
        ("File shares", "Document intake for indexing; the Library doc catalogue."),
        ("LLM endpoint (optional)", "RAG answers + the Assistant chat; local model preferred; degrade to offline."),
    ], widths=(2.2, 4.2))
    doc.add_paragraph("Each connection appears on the Data Sources page with a live health indicator. "
                      "No credentials in code or registry — only in the gitignored .env.")

    # knowledge layer
    doc.add_heading("9. Knowledge / AI layer (BU Brain)", level=1)
    for line in [
        "bu_brain.db (SQLite) logs every dev session, decision and feature — the hub gets smarter the more you use it.",
        "Search across docs, decisions, features and sessions — semantic if embeddings are available, otherwise SQLite FTS5.",
        "An Assistant chat panel scoped to your BU project context for day-to-day development help.",
        "Mirrors the QI Brain model (see the reference screenshot) but is local to the BU machine.",
    ]:
        doc.add_paragraph(line, style='List Bullet')

    # conventions
    doc.add_heading("10. Conventions", level=1)
    for line in CONVENTIONS: doc.add_paragraph(line, style='List Bullet')

    # roadmap
    doc.add_heading("11. Build roadmap", level=1)
    doc.add_paragraph("Ship each phase before starting the next — each leaves BU Hive usable:")
    for name, desc in ROADMAP:
        p = doc.add_paragraph(style='List Number')
        p.add_run(name + " — ").bold = True; p.add_run(desc)
    doc.add_page_break()

    # visual reference
    doc.add_heading("12. Visual reference (the look to match)", level=1)
    doc.add_paragraph("These are QI Hive — BU Hive's sibling — included so the builder has an exact visual target. "
                      "Match this calm, professional AdminLTE aesthetic; the content becomes BU's.")
    for fn, cap in REF_SHOTS:
        fp = SHOTS / fn
        if fp.exists():
            doc.add_picture(str(fp), width=Inches(6.2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = c.add_run(cap); cr.italic = True; cr.font.size = Pt(8.5); cr.font.color.rgb = GREY
    doc.add_page_break()

    # how to use
    doc.add_heading("13. How to use this package on the BU laptop", level=1)
    for i, line in enumerate([
        "Copy the BU_Hive_Starter folder to the BU machine.",
        "Choose the BU Hive working folder (e.g. C:\\BUHive) and initialise a git repo there.",
        "Open Claude Code (or your coding agent) in that folder.",
        "Paste BU_Hive_Master_Build_Prompt.md as the first instruction; let it scaffold Phase 1.",
        "Use this Build Guide as the detailed spec and reference_shots as the visual target.",
        "Configure data sources (SQL / OnBase) in config + .env — never commit credentials.",
        "Iterate phase by phase; commit as you go; keep it local-only.",
    ], 1):
        p = doc.add_paragraph(style='List Number'); p.add_run(line)

    # appendix: prompt
    doc.add_heading("Appendix — Master Build Prompt", level=1)
    doc.add_paragraph("The same prompt ships as a standalone file (BU_Hive_Master_Build_Prompt.md) for easy copy. "
                      "Copy everything inside the box:")
    add_prompt_block(doc, PROMPT)

    doc.add_paragraph()
    foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Quiddity Innovations  ·  BU Hive — a self-contained development control plane")
    fr.italic = True; fr.font.color.rgb = GREY; fr.font.size = Pt(9)

    p = OUT / "BU_Hive_Build_Guide.docx"
    doc.save(str(p)); print("DOCX   ->", p)

# ============================================================ markdown =======
def build_md():
    L = []
    L.append("# BU Hive — Build Guide & Master Prompt\n")
    L.append("**Renne Santiago**  ·  Flagship: **CogniBase (MOC → production)**  ·  modelled on QI Hive  ")
    L.append(f"*Generated {TODAY} — carry this package to the BU laptop*\n")
    L.append("## 1. Purpose & vision\n")
    L.append(VISION + "\n")
    L.append("> **Independence is the point.** This package is the last piece needed to send BU Hive "
             "to its rightful home — a clean-room build kit that needs nothing from your personal PC.\n")
    L.append("## 2. The four sophistication pillars\n")
    for n, d in PILLARS: L.append(f"- **{n}** — {d}")
    L.append("\n## 3. Architecture & stack\n")
    for line in [
        "One self-contained Python web app: FastAPI + Uvicorn, server-rendered HTML (no SPA).",
        "UI: AdminLTE 4 + Bootstrap 5 + Bootstrap Icons, vendored under `/static` (no CDN). Dark default + Light + System.",
        "Adaptable: embed Streamlit/Gradio sub-apps behind the same nav for heavy tools; the shell stays FastAPI.",
        "Two live data sources per request: `bu_registry.json` and `bu_brain.db`.",
        "Shared layout: fixed sidebar (nav + status legend), top bar (clock, theme, health), footer.",
    ]: L.append(f"- {line}")
    L.append("\n## 4. Portability & independence rules\n")
    for line in PORTABILITY: L.append(f"- {line}")
    L.append("\n## 5. BU environment & security\n")
    L.append("BU Hive runs on a corporate machine — treat these as strict:\n")
    for line in SECURITY: L.append(f"- {line}")
    L.append("\n## 6. Page catalogue\n")
    L.append("### 6.1 Core pages (parity with QI Hive)\n")
    L.append("| Page | Route | What it does |\n|---|---|---|")
    for n, r, d in CORE_PAGES: L.append(f"| **{n}** | `{r}` | {d} |")
    L.append("\n### 6.2 BU-specific pages (the differentiators)\n")
    L.append("| Page | Route | What it does |\n|---|---|---|")
    for n, r, d in BU_PAGES: L.append(f"| **{n}** | `{r}` | {d} |")
    L.append("\n### 6.3 Optional / off by default\n")
    L.append("| Page | Why it's optional |\n|---|---|")
    for n, d in OPTIONAL_PAGES: L.append(f"| **{n}** | {d} |")
    L.append("\n## 7. CogniBase Workbench (the flagship)\n")
    L.append("CogniBase is the reason BU Hive exists first. The `/cognibase` page is a dedicated, multi-panel workbench:\n")
    for line in COGNIBASE_PANELS: L.append(f"- {line}")
    L.append("\n*Fallback by design:* if the BU machine can't run embeddings/Ollama, the workbench uses SQLite FTS5 keyword search and shows that path explicitly.\n")
    L.append("## 8. Live data integrations\n")
    L.append("| Source | How BU Hive uses it |\n|---|---|")
    for a, b in [
        ("SQL Server (BU)", "Read-only ad-hoc reporting + saved queries; connection from `.env`; results to table/CSV."),
        ("OnBase (Hyland)", "Document-type browse, retrieval test, and the ingestion source for CogniBase."),
        ("File shares", "Document intake for indexing; the Library doc catalogue."),
        ("LLM endpoint (optional)", "RAG answers + the Assistant chat; local model preferred; degrade to offline."),
    ]: L.append(f"| **{a}** | {b} |")
    L.append("\nEach connection shows a live health indicator on **Data Sources**. Credentials live only in the gitignored `.env`.\n")
    L.append("## 9. Knowledge / AI layer (BU Brain)\n")
    for line in [
        "`bu_brain.db` (SQLite) logs every dev session, decision and feature.",
        "Search across docs/decisions/features/sessions — semantic if embeddings exist, else SQLite FTS5.",
        "An Assistant chat scoped to your BU project context.",
        "Mirrors the QI Brain model, but local to the BU machine.",
    ]: L.append(f"- {line}")
    L.append("\n## 10. Conventions\n")
    for line in CONVENTIONS: L.append(f"- {line}")
    L.append("\n## 11. Build roadmap\n")
    for i, (n, d) in enumerate(ROADMAP, 1): L.append(f"{i}. **{n}** — {d}")
    L.append("\n## 12. Visual reference (the look to match)\n")
    L.append("QI Hive screenshots included as the exact visual target — match the aesthetic, swap in BU content.\n")
    for fn, cap in REF_SHOTS:
        L.append(f"![{cap}](reference_shots/{fn})\n")
        L.append(f"*{cap}*\n")
    L.append("## 13. How to use this package on the BU laptop\n")
    for i, line in enumerate([
        "Copy the `BU_Hive_Starter` folder to the BU machine.",
        "Choose the working folder (e.g. `C:\\BUHive`) and `git init` there.",
        "Open Claude Code (or your agent) in that folder.",
        "Paste `BU_Hive_Master_Build_Prompt.md` as the first instruction; let it scaffold Phase 1.",
        "Use this Build Guide as the spec and `reference_shots/` as the visual target.",
        "Configure data sources in config + `.env` — never commit credentials.",
        "Iterate phase by phase; commit as you go; keep it local-only.",
    ], 1): L.append(f"{i}. {line}")
    L.append("\n## Appendix — Master Build Prompt\n")
    L.append("Also available as the standalone `BU_Hive_Master_Build_Prompt.md`.\n")
    L.append("```text")
    L.append(PROMPT)
    L.append("```\n")
    L.append("*Quiddity Innovations · BU Hive — a self-contained development control plane*")
    p = OUT / "BU_Hive_Build_Guide.md"
    p.write_text("\n".join(L), encoding="utf-8"); print("MD     ->", p)

# ============================================================ prompt + readme =
def build_prompt_and_readme():
    pp = OUT / "BU_Hive_Master_Build_Prompt.md"
    pp.write_text("# BU Hive — Master Build Prompt\n\n"
                  "Paste everything below into your coding agent on the BU laptop as the first "
                  "instruction. Pair it with `BU_Hive_Build_Guide` (detailed spec) and "
                  "`reference_shots/` (visual target).\n\n```text\n" + PROMPT + "\n```\n",
                  encoding="utf-8")
    print("PROMPT ->", pp)

    rd = OUT / "README.md"
    rd.write_text(
        "# BU Hive Starter\n\n"
        "The portable kit to build **BU Hive** — your self-contained development control plane "
        "at BU, with **CogniBase (MOC → production)** as its first mission. Everything here is "
        "designed to be moved to the BU laptop and run independently of your personal PC.\n\n"
        "## What's in this folder\n\n"
        "| File | Purpose |\n|---|---|\n"
        "| `BU_Hive_Master_Build_Prompt.md` | The paste-ready prompt — start here on the BU laptop. |\n"
        "| `BU_Hive_Build_Guide.docx` | The full, formatted spec (pages, data model, security, roadmap) + visual reference. |\n"
        "| `BU_Hive_Build_Guide.md` | Markdown version of the same spec. |\n"
        "| `reference_shots/` | QI Hive screenshots — the visual target to match. |\n\n"
        "## Steps on the BU laptop\n\n"
        "1. Copy this whole `BU_Hive_Starter` folder over.\n"
        "2. Make the working folder (e.g. `C:\\BUHive`) and `git init`.\n"
        "3. Open Claude Code (or your agent) in that folder.\n"
        "4. Paste `BU_Hive_Master_Build_Prompt.md`; let it scaffold **Phase 1**.\n"
        "5. Drive the rest with the Build Guide, phase by phase.\n"
        "6. Put SQL/OnBase credentials in a gitignored `.env` — never in code.\n"
        "7. Keep it **local-only** (bind 127.0.0.1, no tunnels) per BU IT policy.\n\n"
        "## The four goals\n\n"
        "Feature parity with QI Hive · Visual & UX polish · Live data integrations · Knowledge/AI layer.\n\n"
        "*Generated " + TODAY + " — Quiddity Innovations.*\n",
        encoding="utf-8")
    print("README ->", rd)

if __name__ == "__main__":
    build_docx()
    build_md()
    build_prompt_and_readme()
    print("OK — BU Hive Starter package built.")
