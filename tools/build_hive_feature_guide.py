# -*- coding: utf-8 -*-
"""
Build the QI Hive — Dashboard Feature Guide.
Produces a polished Word .docx (screenshots embedded) + a Markdown companion,
both written to C:\\QIH\\docs.
"""
import sys, os
sys.stdout.reconfigure(encoding="utf-8")
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SHOTS = Path(r"C:\QIH\docs\assets\hive_screens")
OUT_DOCX = Path(r"C:\QIH\docs\QI_Hive_Dashboard_Feature_Guide.docx")
OUT_MD   = Path(r"C:\QIH\docs\QI_Hive_Dashboard_Feature_Guide.md")
TODAY = "2026-06-19"

PURPLE = RGBColor(0x6A, 0x4C, 0x93)
INK    = RGBColor(0x22, 0x26, 0x2B)
GREY   = RGBColor(0x6B, 0x72, 0x80)

# ---------------------------------------------------------------- content ----
INTRO = (
    "The QI Hive is the unified control plane for every Quiddity Innovations project. "
    "It is a single FastAPI web application (the Dashboard, served on port 8600 as the "
    "Windows service QI_Dashboard) that reads two live sources of truth — the ecosystem "
    "registry (qi_registry.json) and the shared knowledge substrate (QI Brain: SQLite + "
    "ChromaDB on port 9011) — and renders the whole ecosystem onto one set of pages. "
    "From here you can see project health, run tests, manage a task board, browse logs, "
    "track LLM spend, coordinate the seven specialist Hive agents, and search ~900 "
    "documents across nine drives. This guide documents every page of the Hive, what it "
    "does, and what you see on it."
)

ANATOMY = [
    ("Left sidebar", "Fixed navigation listing all pages under the “QI HIVE” header, plus a colour-coded Status Legend (Complete / In Progress / Backlog / New / Pre-POC / Retired)."),
    ("Top bar", "Live clock, a theme switcher (Dark / Light / System — all three are consistent; Dark is the default), and a heart-pulse shortcut that jumps straight to the Health Check."),
    ("Content area", "The page itself. Most pages open with an “About this page” expander and read live data on every load."),
    ("Footer", "“Quiddity Innovations” on the left and “QI Hive v3.0 — Powered by QI Brain” on the right, present on every page."),
]

# ------------------------------------------------------- master build prompt -
# A single, self-contained prompt distilled from this guide. Hand it to a capable
# coding agent to reconstruct the QI Hive dashboard faithfully.
PROMPT = """You are a senior full-stack engineer. Build QI Hive - a single-pane operations
dashboard and control plane for a portfolio of ~22 independent software projects
(the "QI ecosystem") that share one Windows machine, port allocations and Git.

GOAL
Produce one self-contained Python web application that renders the entire ecosystem -
health, services, tasks, logs, cost, agents and a shared knowledge base - onto a set of
server-rendered pages, reading live from two sources of truth and degrading gracefully
when either is offline.

STACK & SHELL
- Python 3.11+, FastAPI, Uvicorn. One module serves all routes and returns
  server-rendered HTML (no SPA framework).
- UI: AdminLTE 4 + Bootstrap 5 + Bootstrap Icons, all vendored under /static (no CDN).
  Serve on port 8600; run as a Windows NSSM service named QI_Dashboard.
- Shared layout on every page: a fixed left sidebar (nav links + a colour-coded status
  legend), a top bar (live clock, a Dark/Light/System theme switcher with Dark as the
  default, and a heart-pulse shortcut to Health), and a footer
  ("Quiddity Innovations" / "QI Hive v3.0 - Powered by QI Brain").

DATA SOURCES (read live on every request; never crash if unavailable)
1. qi_registry.json - the ecosystem registry: each project's identity, paths, ports,
   NSSM services, integrations and family tier.
2. qi_brain.db (SQLite + ChromaDB), reached through a separate "QI Brain" API on port
   9011 - session logs, a decisions registry, a features registry, per-project state, a
   documentation catalogue + knowledge graph, and full-text/semantic memory.

PAGES (each is its own route)
1.  Dashboard (/) - at-a-glance home: a Today->YTD spend ladder, a status table for all
    projects, the agent roster, recent sessions, and per-project local-LLM usage.
2.  Launcher (/launcher) - Launchpad grid of every app by category, live status dots,
    open-local vs open-tunnel buttons, grid<->columns toggle; URLs built live from the
    registry + a port probe + a tunnel resolver.
3.  Tunnels (/tunnels) - one card per live Cloudflare tunnel with clickable URL, copy
    button and an offline QR code (segno); a "not running" list naming the
    QI_<App>Tunnel service to start.
4.  The Hive (/hive) - roster of seven specialist agents (Architect, Builder, Scout,
    Scribe, Inspector, Tester, Ops) with Brain-logged activity counts; ecosystem stat
    tiles; recent sessions; a Brain poller + a memory-distillation control.
5.  Health Check (/health) - on-demand scan of every project (service up/down, port
    listening, git clean, docs, state) with an "Action Needed" remediation list;
    content-negotiated so monitors/validators receive JSON.
6.  Task Board (/board) - drag-and-drop Kanban (Backlog->In Progress->Review->Done);
    add/edit/delete cards with project, agent and priority; persists to DB; failed tests
    auto-file cards.
7.  Tests (/tests) - Smoke/API/UI (Playwright) suites with a Run-All; pass/fail/skip
    tiles and per-test timings; failures become board tasks.
8.  Project Status (/projects/status) - an index linking to detailed per-project pages
    rendered from each project's INTRO files.
9.  Services (/services) - NSSM service inventory: every QI_* service with status, app
    directory, port and description.
10. Scheduled Tasks (/tasks) - Windows Task Scheduler jobs: schedule, last run/result,
    next run.
11. LLM Usage (/usage) - token/cost analytics by project and model, framed as
    Claude-API-vs-run-it-local savings; spend ladder + daily chart.
12. Headlines (/news) - a chronological feed of everything happening across the Hive
    (sessions, decisions, features, dispatch, compliance) with type filter chips.
13. Activity (/activity) - event log/audit trail merging Hive reports with a per-session
    table (turns, duration, tokens, cost) drawn from agent transcripts.
14. CoWork Dispatch (/dispatch) - review queue for AI-proposed edits/tasks: pending vs
    resolved, the suggested JSON, and Approve/Apply actions.
15. QI Brain (/brain) - tabbed UI over the knowledge substrate:
    Overview, Decisions, Features, Archive, Distillation, Inbox, Search.
16. Mission Control (/mission-control) - single-pane ops board: a live agent strip
    (model + last activity), a project index, a Brain snapshot and the dispatch queue.
17. War Room (/warroom) - multi-way @-addressable chat with every agent; replies
    generated by a local LLM; auto-refresh; an inbound bridge endpoint for external
    channels.
18. Logs (/logs) - centralised log viewer: project/file picker, large-file tail,
    colour-coded levels; plus a per-service log-level config card.
19. Config (/config) - elevation (gsudo) presets + per-project profiles, and per-service
    log-level configuration persisted to file.
20. Library (/library) - the "Documentation Brain": full-text + semantic search over
    ~900 docs, filters, open/reveal, and a TheBrain-style knowledge graph ("Plex");
    index built nightly, storage stays federated.
21. Guide (/guide) - renders the operator manual (Markdown->HTML) with a Raw .md
    download.
22. Compliance (/compliance) - per-project standards-compliance cards + a findings log;
    triggers scans via the Brain.

CONVENTIONS
- Windows services use NSSM and the name QI_<Project><Role>; always set Description and
  AppDirectory.
- Respect each project's allocated port block; never hardcode ports - read the registry.
- UTF-8 everywhere; read-only pages must never mutate state.
- Keep it one app, dependency-light, and resilient when the Brain or a project is down.

DELIVERABLE
A runnable FastAPI app plus its vendored static assets and a service-install script,
faithful to the pages and behaviours above."""

# Each: (n, label, route, icon, screenshot, what, [features], [subtabs])
PAGES = [
    (1, "Dashboard", "/", "speedometer2", "01_dashboard.png",
     "The home landing page and at-a-glance command centre. It folds the whole ecosystem onto one screen: a token / cost consumption ladder across time windows, a status table for every project, the live agent roster, recent session summaries, and a per-project local-LLM usage breakdown.",
     ["Consumption ladder — spend for Today, This Week, 30 days, QTD and YTD",
      "Projects table — status, open tasks and quick links for all ~22 projects",
      "Agent team panel — which Hive agents are idle vs active and their current model",
      "Recent sessions feed — what was last worked on, per project",
      "Local LLM-by-project usage table"], None),

    (2, "Launcher", "/launcher", "grid-3x3-gap", "02_launcher.png",
     "A Launchpad-style click-to-open grid of every QI app, grouped by category (Core Product, Backbone, Assistants, Cousins…). Cards are built live from the registry plus a port probe and the tunnel resolver, so the URLs — including rotating Cloudflare quick-tunnel URLs — are always current. Opened over a tunnel, it prefers public URLs.",
     ["Categorised cards for every project",
      "Live status dot — green healthy, yellow warning, red down",
      "Open the local URL or the public (tunnel) URL in one click",
      "Grid ⇄ columns layout toggle",
      "URLs always reflect the registry + live probe"], None),

    (3, "Tunnels", "/tunnels", "globe2", "03_tunnels.png",
     "A human-readable view of every live Cloudflare tunnel. Each running tunnel gets a card with its clickable public URL, a copy button, and an offline QR code so you can open the app on a phone by scanning. Tunnels that are down are listed separately with the service to start.",
     ["One card per live tunnel with its public https URL",
      "Scannable offline QR code (segno) per tunnel — open on your phone",
      "Copy and “Open ↗” buttons",
      "“Not running” list naming the QI_<App>Tunnel service to start",
      "Live count of running tunnels"], None),

    (4, "The Hive", "/hive", "hexagon", "04_hive.png",
     "The agent operations hub. It shows the seven specialist Hive agents — Architect, Builder, Scout, Scribe, Inspector, Tester, Ops — alongside Claude Code / Claude Work and the project assistants, each with its Brain-logged activity count and a profile link. It also surfaces ecosystem stat tiles, recent Brain sessions, the Brain Poller control, and a memory-distillation panel.",
     ["Agent roster cards — role, Brain-logged counts and a Profile link",
      "Ecosystem stat tiles — active projects, decisions, sessions, features",
      "Recent sessions table (project, summary, date)",
      "Brain Poller — running indicator + poll history",
      "“Distill Brain Memory” — compress long-term memory on demand"], None),

    (5, "Health Check", "/health", "heart-pulse", "05_health.png",
     "A live, on-demand scan of the whole ecosystem that runs every time you open it. For each project on disk it checks service status, port listening, git cleanliness, docs and a state summary, then surfaces an “Action Needed” list of concrete issues with remediation. The same route is content-negotiated — monitors and the QI validator receive JSON.",
     ["“Action Needed” remediation list at the top",
      "Per-project rows: service, port, git, docs, summary, health badge",
      "Colour-coded health — ok / warning / attention",
      "Re-check button; doubles as the JSON health probe at /health"], None),

    (6, "Task Board", "/board", "kanban", "06_board.png",
     "A drag-and-drop Kanban board for work across all projects, with four columns — Backlog → In Progress → Review → Done. Cards carry a project, an assigned Hive agent and a priority colour; every move and edit persists to the database. Failed tests file cards here automatically.",
     ["Four columns; drag cards between them (auto-saves)",
      "Add Task modal — title, description, project, agent, priority",
      "Filter by project; per-card delete; multi-select mode",
      "Priority colour stripe (high / medium / low) + agent icon",
      "Test failures appear as new cards automatically"], None),

    (7, "Tests", "/tests", "bug", "07_tests.png",
     "A test runner and results panel for the whole ecosystem. It runs Smoke (~15 s health pings), API (~60 s endpoint coverage) and UI (~90 s Playwright) suites, shows pass / fail / skip counts with a pass-rate gauge and per-test timings, and includes the EasyFlow Chrome-extension test launcher. Failures auto-create board tasks.",
     ["Smoke / API / UI / Run-All buttons",
      "Summary tiles — passed, failed, skipped, pass-rate",
      "Per-test results table with durations",
      "EasyFlow extension test card",
      "Failures become Kanban tasks"], None),

    (8, "Project Status", "/projects/status", "clipboard-data", "08_projects.png",
     "An index of Maia-style status pages for every project. Each row links to a detailed per-project page rendered from that project’s INTRO folder (status_intro.md plus features / techstack / future JSON). The index shows each project’s readiness (ready / empty) and its INTRO path.",
     ["One row per project → a detailed status page",
      "Ready / empty badge per project",
      "Content sourced from each project’s INTRO/ files",
      "Edit those files and Refresh to update the page"], None),

    (9, "Services", "/services", "gear-wide-connected", "09_services.png",
     "The complete Windows NSSM service inventory. It lists every QI_* service (and known legacy ones) with run status, app directory, port and a description, read straight from the registry — the single place to see what is installed and what is running.",
     ["All QI_* services with Running / Stopped status",
      "App directory, port and description columns",
      "Refresh; sourced from qi_registry.json + live status",
      "Mirrors the QI_<Project><Role> naming standard"], None),

    (10, "Scheduled Tasks", "/tasks", "calendar-event", "10_tasks.png",
     "A view of the QI-relevant Windows Task Scheduler jobs — nightly sync, reconcilers, TubeScout AM/PM, daily broadcasts and more. For each task it shows the schedule, last run and result, and the next run, with indicators for hidden vs visible consoles and killed runs.",
     ["Task name, state (Ready / Disabled) and command",
      "Schedule (“Every”), Last Run + Result, Next Run",
      "Indicators for hidden-window, visible-window and killed runs",
      "Refresh"], None),

    (11, "LLM Usage", "/usage", "graph-up-arrow", "11_usage.png",
     "Token and cost analytics across every project and model, framed as Claude-API-vs-run-it-local savings. It shows the same Today→YTD consumption ladder, a daily-spend chart, and by-project and by-model tables comparing actual cost, what it would cost locally, and a combined total plus total savings.",
     ["Consumption ladder — Today, Week, 30d, QTD, YTD",
      "Daily spend bar chart (last N days)",
      "By-project and by-model cost tables",
      "“Savings” — Claude API vs local Ollama comparison",
      "Estimate caveats noted inline"], None),

    (12, "Headlines", "/news", "newspaper", "12_news.png",
     "A Twitter/X-style chronological feed of everything happening across the Hive — sessions, decisions, features, dispatches, compliance findings and state changes, newest first. Filter chips narrow the feed by type. It is backed by the Brain’s event stream (and surfaces the NEXUS Scout AI-news digest).",
     ["Chronological “Latest Across the Hive” feed",
      "Type filter chips — sessions, decisions, features, dispatch, compliance…",
      "Relative timestamps and source per item",
      "Backed by Brain events + NEXUS Scout"], None),

    (13, "Activity", "/activity", "activity", "13_activity.png",
     "The event log and audit trail. It combines Hive reports (session start/end, decisions, errors, hook events) with a per-session table drawn from the actual Claude Code transcripts — turns, duration, tokens and cost per session. Tiles at the top show sessions, assistant turns and spend for the period.",
     ["Tiles — sessions, assistant turns, spend (period)",
      "Hive Reports feed — session / decision / error events with host",
      "Per-session table from Claude Code transcripts (turns, dur, tokens, cost)",
      "Two data sources clearly distinguished"], None),

    (14, "CoWork Dispatch", "/dispatch", "send-check", "14_dispatch.png",
     "The review queue for the Claude Code ⇄ Hive “CoWork” loop. Claude Work / Claude Code propose edits or tasks; this page shows the items pending review and the resolved ones, each with its suggested JSON payload and Approve / Apply-label / Apply-edit controls. Compliance can flag risky human-on-purpose changes.",
     ["Pending (“Awaiting Review”) vs Resolved columns",
      "Suggested edit/task JSON for each item",
      "Approve / Apply-label / Apply-edit actions",
      "Compliance flagging on dispatched changes"], None),

    (15, "QI Brain", "/brain", "cpu", "15_brain.png",
     "The dashboard front-end for the QI Brain knowledge substrate (Brain API on :9011 — SQLite + ChromaDB). A tabbed view over the shared memory, decisions, features, sessions and ecosystem state of every project, with a Brain-poller status indicator. The sub-tabs are documented below.",
     ["Shared memory, decisions, features and sessions in one place",
      "Brain API status + poller indicator",
      "Tabbed: Overview · Decisions · Features · Archive · Distillation · Inbox · Search"],
     [("Overview", "15_brain.png", "“Projects in the Brain” — a card per project with its phase, status, last-updated date and logged counts."),
      ("Decisions", "15b_brain_decisions.png", "The decision registry: project, title, rationale and timestamp for every recorded design decision."),
      ("Features", "15c_brain_features.png", "The feature registry: project, name, domain and description — the cross-project catalogue used to reuse work instead of rebuilding it."),
      ("Archive", "15d_brain_archive.png", "Archived and superseded records, kept for history without cluttering the live views."),
      ("Distillation", "15e_brain_distill.png", "Outputs of memory distillation — compressed long-term memory the Hive can re-load cheaply."),
      ("Inbox", "15f_brain_inbox.png", "The Brain inbox log. Drop JSON messages in engine\\brain\\inbox or POST to /api/inbox and they appear here with status, source, kind and received time."),
      ("Search", "15g_brain_search.png", "Full-text / semantic memory search across all logged context, usable locally and over the tunnel.")]),

    (16, "Mission Control", "/mission-control", "broadcast-pin", "16_mission.png",
     "A single-pane-of-glass operations board. It combines the live agent strip (Claude Code, Claude Work, CoWork and the Hive agents, each with current model and last activity), a project index with phase / status / last-active for every project, a Brain snapshot, and the dispatch queue — the busiest “what is everything doing right now” view.",
     ["Active-agents strip — model + last-active per agent",
      "Project index — phase, status and last-active for all projects",
      "Brain snapshot panel",
      "Dispatch queue table"], None),

    (17, "War Room", "/warroom", "chat-dots", "17_warroom.png",
     "A live text chat between Renne and every QI agent — Claude Code, Claude Work, CoWork and the seven Hive agents. Address a specialist with @architect, @builder, @inspector, @ops, @scout, @scribe or @tester; the Hive host replies by default. Replies are generated by the local NEXUS LLM, the feed auto-refreshes every few seconds, and inbound bridges (e.g. Telegram via Tasuke) can post in.",
     ["Multi-way, @-addressable agent chat",
      "Replies generated by the local NEXUS LLM; 4 s auto-refresh",
      "Posts as “renne”; Hive host is the default responder",
      "Inbound bridge endpoint for external channels (Telegram, etc.)"], None),

    (18, "Logs", "/logs", "journal-text", "18_logs.png",
     "A centralised log browser across every project. Pick a project and a file, tail the last N lines with optional auto-refresh, and read colour-coded entries. Files are discovered from each project’s registered log root, and an efficient reverse-block tail handles files over 100 MB. The per-service Log Level card sits below.",
     ["Project + file picker, tail size and auto-refresh toggle",
      "Reverse-block tail — fast even on very large files",
      "File list with size and modified time",
      "Log Level Configuration card (set level per service)"], None),

    (19, "Config", "/config", "sliders", "19_config.png",
     "Operational configuration for the Hive host. Two cards: gsudo (elevation) configuration — quick presets (Loose / Normal / Strict / Locked), credential-cache mode and duration, and security toggles, plus per-project gsudo profiles — and Log Level configuration, which sets verbosity per service and persists to config/logging.json.",
     ["gsudo presets + credential cache + UAC / security toggles",
      "Per-project gsudo profiles",
      "Log Level per service — persists immediately for the dashboard",
      "Changes are written to config files (the UI is the editor)"], None),

    (20, "Library", "/library", "journals", "20_library.png",
     "The Documentation Brain — a searchable index and knowledge graph (“Plex”) over 900+ documents across all project drives. Search by keyword or meaning, filter by project and type, open or reveal any document, and explore TheBrain-style relationships. Storage stays federated; only the index is centralised, rebuilt nightly by the doc harvester.",
     ["Full-text + semantic search over ~900 cataloged docs",
      "Filters by project and type; live counts (docs / embedded / stale)",
      "Open a doc or reveal it in Explorer",
      "Knowledge-graph neighbourhood (“Plex”)",
      "Backed by the qi_brain.db docs table + qi_docs Chroma collection"], None),

    (21, "Guide", "/guide", "book", "21_guide.png",
     "The built-in operator manual. It renders QI_Claude_Manager_Guide.md as HTML — Quick Start, every dashboard tab explained, the 22-project catalogue, services & elevation, key files, the golden integration rules, ports, troubleshooting and the convergence vision — with a Raw .md download.",
     ["The full operator cheatsheet, rendered in-app",
      "Quick-start “I want to… / where to go” table",
      "Project, service and port reference",
      "Golden rules + troubleshooting playbook",
      "Raw .md download"], None),

    (22, "Compliance", "/compliance", "shield-check", "22_compliance.png",
     "The standards-compliance panel (a utility page reached at /compliance). Per-project compliance cards sit above a recent-activity log of checks — severity, status, action and message — talking to the Brain’s /api/compliance/* endpoints. You can trigger a scan and watch findings (missing /health, naming, registry gaps…) accumulate.",
     ["Per-project compliance status cards",
      "Recent findings log — severity / status / action / message",
      "Trigger a scan; proxied to the Brain",
      "Surfaces QI standards violations across the ecosystem"], None),
]

# ---------------------------------------------------------------- helpers ----
def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'C9C9C9')
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def img_for(doc, name, width=6.4):
    fp = SHOTS / name
    if fp.exists():
        doc.add_picture(str(fp), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_prompt_block(doc, text):
    """Render `text` as a shaded, monospace copy block inside a 1-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, "F4F1FA")  # pale QI purple
    cell.paragraphs[0].text = ""
    first = True
    for line in text.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        pf = p.paragraph_format
        pf.space_after = Pt(0); pf.space_before = Pt(0); pf.line_spacing = 1.0
        r = p.add_run(line if line else "")
        r.font.name = 'Consolas'; r.font.size = Pt(8.5); r.font.color.rgb = INK

# ---------------------------------------------------------------- docx -------
def build_docx():
    doc = Document()
    # base font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(10.5)

    # ---- cover ----
    for _ in range(3): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("QI Hive"); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = PURPLE
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Dashboard Feature Guide"); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = INK
    s2 = doc.add_paragraph(); s2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s2.add_run("Every page of the unified control plane — what it does, with screenshots")
    r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREY
    for _ in range(2): doc.add_paragraph()
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in ["Quiddity Innovations", "Service: QI_Dashboard  ·  http://localhost:8600  ·  QI Hive v3.0",
                 f"Generated {TODAY}"]:
        rr = meta.add_run(line + "\n"); rr.font.size = Pt(11); rr.font.color.rgb = INK
    doc.add_page_break()

    # ---- overview ----
    doc.add_heading("Overview", level=1)
    doc.add_paragraph(INTRO)

    doc.add_heading("Anatomy of every page", level=2)
    doc.add_paragraph("The same chrome wraps every page in the Hive:")
    for label, desc in ANATOMY:
        p = doc.add_paragraph(style='List Bullet')
        rr = p.add_run(label + " — "); rr.bold = True
        p.add_run(desc)

    # ---- navigation map table ----
    doc.add_heading("Navigation map", level=2)
    doc.add_paragraph("The left sidebar exposes 21 pages; a 22nd (Compliance) is a utility page reached directly.")
    tbl = doc.add_table(rows=1, cols=3); tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["#", "Page (route)", "What it is for"]):
        hdr[i].paragraphs[0].add_run(h).bold = True
        set_cell_bg(hdr[i], "6A4C93")
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    one_liners = {
        1:"At-a-glance command centre", 2:"Click-to-open every app", 3:"Live Cloudflare tunnels + QR",
        4:"The seven Hive agents", 5:"Live ecosystem health scan", 6:"Kanban task board",
        7:"Test runner + results", 8:"Per-project status pages", 9:"NSSM service inventory",
        10:"Windows scheduled jobs", 11:"Token + cost analytics", 12:"Cross-Hive activity feed",
        13:"Event log / audit trail", 14:"Claude Code review queue", 15:"Knowledge substrate UI",
        16:"Single-pane ops board", 17:"Chat with every agent", 18:"Centralised log viewer",
        19:"Elevation + log config", 20:"Documentation Brain search", 21:"Built-in operator manual",
        22:"Standards compliance",
    }
    for n, label, route, *_ in PAGES:
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].add_run(str(n))
        rp = cells[1].paragraphs[0]; rp.add_run(label).bold = True
        rp.add_run(f"  ({route})").italic = True
        cells[2].paragraphs[0].add_run(one_liners.get(n, ""))
    doc.add_page_break()

    # ---- per-page sections ----
    for n, label, route, icon, shot, what, feats, subtabs in PAGES:
        doc.add_heading(f"{n}. {label}", level=1)
        meta = doc.add_paragraph()
        rr = meta.add_run(f"Route: {route}"); rr.bold = True; rr.font.color.rgb = PURPLE
        meta.add_run(f"      Sidebar icon: {icon}").font.color.rgb = GREY
        img_for(doc, shot)
        cap = doc.paragraphs[-1]
        # caption line under image
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = c.add_run(f"The {label} page (http://localhost:8600{route})")
        cr.italic = True; cr.font.size = Pt(8.5); cr.font.color.rgb = GREY

        doc.add_heading("What it does", level=3)
        doc.add_paragraph(what)
        doc.add_heading("Key features", level=3)
        for f in feats:
            doc.add_paragraph(f, style='List Bullet')

        if subtabs:
            doc.add_heading("Sub-tabs", level=3)
            doc.add_paragraph("QI Brain is itself tabbed. Each sub-tab:")
            for st_name, st_shot, st_desc in subtabs:
                h = doc.add_paragraph()
                hr = h.add_run(st_name); hr.bold = True; hr.font.color.rgb = PURPLE
                doc.add_paragraph(st_desc)
                img_for(doc, st_shot, width=6.0)
        doc.add_page_break()

    # ---- appendix ----
    doc.add_heading("Appendix A — Data sources", level=1)
    doc.add_paragraph("Every page reads from two live sources of truth:")
    for label, desc in [
        ("qi_registry.json", "The ecosystem registry — project identity, ports, NSSM services, integrations and family tiers. Drives Launcher, Services, Health, Project Status and more."),
        ("qi_brain.db (SQLite + ChromaDB)", "The QI Brain knowledge substrate on :9011 — session logs, the decision and feature registries, project state, the documentation catalogue/graph and searchable memory. Drives The Hive, QI Brain, Activity, Headlines, Mission Control and Library."),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(label + " — ").bold = True
        p.add_run(desc)

    doc.add_heading("Appendix B — Hosting", level=1)
    for line in [
        "Service: QI_Dashboard (NSSM) — AppDirectory C:\\QIH, port 8600.",
        "Public access: QI_DashboardTunnel (Cloudflare, on-demand) — see the Tunnels page.",
        "Code: C:\\QIH\\engine\\hive\\dashboard\\server.py (single FastAPI app).",
        "Companion service: QI_BrainAPI on :9011 supplies the knowledge layer.",
        "Stack: FastAPI + AdminLTE / Bootstrap 5, Bootstrap Icons; three themes (Dark default).",
    ]:
        doc.add_paragraph(line, style='List Bullet')

    # ---- Appendix C: master build prompt ----
    doc.add_page_break()
    doc.add_heading("Appendix C — Master Build Prompt", level=1)
    doc.add_paragraph(
        "The prompt below distils this entire guide into a single, self-contained "
        "specification. Hand it to a capable coding agent (e.g. Claude) to reconstruct the "
        "QI Hive dashboard faithfully, or use it as the brief for a clean-room rebuild. "
        "Copy everything inside the box.")
    add_prompt_block(doc, PROMPT)

    doc.add_paragraph()
    foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run("Quiddity Innovations  ·  QI Hive v3.0 — Powered by QI Brain")
    fr.italic = True; fr.font.color.rgb = GREY; fr.font.size = Pt(9)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print("DOCX  ->", OUT_DOCX)

# ---------------------------------------------------------------- markdown ---
def build_md():
    L = []
    L.append("# QI Hive — Dashboard Feature Guide\n")
    L.append("**Quiddity Innovations**  ·  Service `QI_Dashboard` · http://localhost:8600 · QI Hive v3.0  ")
    L.append(f"*Generated {TODAY}*\n")
    L.append("## Overview\n")
    L.append(INTRO + "\n")
    L.append("### Anatomy of every page\n")
    for label, desc in ANATOMY:
        L.append(f"- **{label}** — {desc}")
    L.append("\n### Navigation map\n")
    L.append("The left sidebar exposes 21 pages; a 22nd (Compliance) is a utility page reached directly.\n")
    L.append("| # | Page | Route | What it is for |")
    L.append("|---|---|---|---|")
    one_liners = {
        1:"At-a-glance command centre", 2:"Click-to-open every app", 3:"Live Cloudflare tunnels + QR",
        4:"The seven Hive agents", 5:"Live ecosystem health scan", 6:"Kanban task board",
        7:"Test runner + results", 8:"Per-project status pages", 9:"NSSM service inventory",
        10:"Windows scheduled jobs", 11:"Token + cost analytics", 12:"Cross-Hive activity feed",
        13:"Event log / audit trail", 14:"Claude Code review queue", 15:"Knowledge substrate UI",
        16:"Single-pane ops board", 17:"Chat with every agent", 18:"Centralised log viewer",
        19:"Elevation + log config", 20:"Documentation Brain search", 21:"Built-in operator manual",
        22:"Standards compliance",
    }
    for n, label, route, *_ in PAGES:
        L.append(f"| {n} | **{label}** | `{route}` | {one_liners.get(n,'')} |")
    L.append("\n---\n")
    for n, label, route, icon, shot, what, feats, subtabs in PAGES:
        L.append(f"## {n}. {label}\n")
        L.append(f"**Route:** `{route}`  ·  **Sidebar icon:** `{icon}`\n")
        L.append(f"![{label}](assets/hive_screens/{shot})\n")
        L.append(f"*The {label} page (http://localhost:8600{route})*\n")
        L.append("**What it does**\n")
        L.append(what + "\n")
        L.append("**Key features**\n")
        for f in feats:
            L.append(f"- {f}")
        L.append("")
        if subtabs:
            L.append("**Sub-tabs**\n")
            for st_name, st_shot, st_desc in subtabs:
                L.append(f"#### {st_name}\n")
                L.append(st_desc + "\n")
                L.append(f"![QI Brain — {st_name}](assets/hive_screens/{st_shot})\n")
        L.append("\n---\n")
    L.append("## Appendix A — Data sources\n")
    L.append("Every page reads from two live sources of truth:\n")
    L.append("- **qi_registry.json** — the ecosystem registry (project identity, ports, NSSM services, integrations, family tiers).")
    L.append("- **qi_brain.db (SQLite + ChromaDB)** — the QI Brain knowledge substrate on :9011 (sessions, decisions, features, project state, doc index, searchable memory).\n")
    L.append("## Appendix B — Hosting\n")
    L.append("- Service: `QI_Dashboard` (NSSM) — AppDirectory `C:\\QIH`, port 8600.")
    L.append("- Public access: `QI_DashboardTunnel` (Cloudflare, on-demand) — see the Tunnels page.")
    L.append("- Code: `C:\\QIH\\engine\\hive\\dashboard\\server.py` (single FastAPI app).")
    L.append("- Companion service: `QI_BrainAPI` on :9011 supplies the knowledge layer.")
    L.append("- Stack: FastAPI + AdminLTE / Bootstrap 5, Bootstrap Icons; three themes (Dark default).\n")
    L.append("## Appendix C — Master Build Prompt\n")
    L.append("The prompt below distils this entire guide into a single, self-contained "
             "specification. Hand it to a capable coding agent (e.g. Claude) to reconstruct "
             "the QI Hive dashboard faithfully, or use it as the brief for a clean-room rebuild.\n")
    L.append("```text")
    L.append(PROMPT)
    L.append("```\n")
    L.append("*Quiddity Innovations · QI Hive v3.0 — Powered by QI Brain*")
    OUT_MD.write_text("\n".join(L), encoding="utf-8")
    print("MD    ->", OUT_MD)

if __name__ == "__main__":
    build_docx()
    build_md()
    print("OK — pages documented:", len(PAGES))
