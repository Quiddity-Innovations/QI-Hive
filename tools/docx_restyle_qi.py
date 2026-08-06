# -*- coding: utf-8 -*-
"""
docx_restyle_qi.py — re-render an existing hand-authored .docx in the
"QI Report" template v2 look (same template as the Project Status library).

Walks the source document body in order (paragraphs + tables), maps built-in
styles (Title / Heading 1-3 / List Bullet / List Number / Normal) onto the QI
template equivalents, preserves run-level bold/italic/monospace, and restyles
tables with the accent header + banded rows. Source files must contain no
images (checked; aborts if any are found).

Usage:
    python docx_restyle_qi.py <in.docx> <out.docx> [--subtitle "IT Components"]
"""
from __future__ import annotations
import sys
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
sys.path.insert(0, str(Path(r"C:\QIH\shared\documentation\project_library")))
import _build_project_docs as T
from md_to_qi_docx import cover  # same generic cover block

sys.path.insert(0, str(Path(r"C:\QIH\tools")))

from docx import Document
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


def iter_body(doc):
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def copy_runs(src_par, dst_par, size=None):
    for r in src_par.runs:
        nr = dst_par.add_run(r.text)
        nr.bold = r.bold
        nr.italic = r.italic
        if r.font.name and r.font.name.lower() in ("consolas", "courier new", "cascadia code"):
            nr.font.name = T.MONO_FONT
            nr.font.size = Pt(9.5)
        if size:
            nr.font.size = Pt(size)


def restyle(src: Path, dst: Path, subtitle: str, title: str | None = None):
    old = Document(str(src))
    if len(old.inline_shapes):
        raise SystemExit(f"ABORT: {src.name} contains images — restyle manually.")

    # title: explicit override, else first Title/Heading 1 paragraph
    consume_title = title is None
    if title is None:
        title = src.stem
        for el in iter_body(old):
            if isinstance(el, Paragraph) and el.style.name in ("Title", "Heading 1") and el.text.strip():
                title = el.text.strip()
                break

    doc = Document()
    T.set_base_styles(doc)
    T.setup_page(doc, title, header_text=title)
    cover(doc, title, subtitle, src)

    title_used = False
    for el in iter_body(old):
        if isinstance(el, Table):
            rows = [[("\n".join(p.text for p in c.paragraphs)).strip() for c in r.cells]
                    for r in el.rows]
            if not rows:
                continue
            t = doc.add_table(rows=0, cols=max(len(r) for r in rows))
            for r in rows:
                cells = t.add_row().cells
                for j, val in enumerate(r):
                    cells[j].paragraphs[0].add_run(val)
            T.style_data_table(t)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        st = el.style.name
        txt = el.text.strip()
        if st in ("Title", "Heading 1"):
            if consume_title and not title_used and txt == title:
                title_used = True   # already on the cover
                continue
            if txt:
                T.h1(doc, txt)
        elif st.startswith("Heading 2"):
            if txt: doc.add_heading(txt, level=2)
        elif st.startswith("Heading 3") or st.startswith("Heading 4"):
            if txt: doc.add_heading(txt, level=3)
        elif st.startswith("List Bullet"):
            copy_runs(el, doc.add_paragraph(style="List Bullet"))
        elif st.startswith("List Number"):
            copy_runs(el, doc.add_paragraph(style="List Number"))
        elif txt:
            copy_runs(el, doc.add_paragraph())

    doc.save(str(dst))
    return title


if __name__ == "__main__":
    argv = sys.argv[1:]
    sub, ttl = "Reference Document", None
    if "--subtitle" in argv:
        i = argv.index("--subtitle"); sub = argv[i + 1]; del argv[i:i + 2]
    if "--title" in argv:
        i = argv.index("--title"); ttl = argv[i + 1]; del argv[i:i + 2]
    t = restyle(Path(argv[0]), Path(argv[1]), sub, ttl)
    print(f"OK  {t}: {argv[1]}")
