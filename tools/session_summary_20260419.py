# -*- coding: utf-8 -*-
"""Session summary .docx per CLAUDE.md auto-save rule."""
from docx import Document
from docx.shared import Pt
from pathlib import Path
from datetime import datetime

OUT_DIR = Path(r"C:\UNIVERSAL\DOCUMENTATION\Session_Summaries")
OUT = OUT_DIR / f"QIHive_Summary_2026-04-19_{datetime.now().strftime('%H%M')}.docx"

d = Document()
d.add_heading("QI Hive — Session Summary (2026-04-19 evening)", level=1)
d.add_paragraph("Autonomous session: dashboard observability panels + Claude usage tracking + autonomous service control via sc.exe.")

d.add_heading("COMPLETED THIS SESSION", level=2)
for item in [
    "Dashboard pages /services, /tasks, /usage, /activity live on port 8600",
    "Project status colors corrected: Maia/OpenClaw red Backlog, FileHQ gray Retired, MQ light-green New, EasyFlow/QI_Hive/NEXUS orange In Progress",
    "Status color legend added at bottom of left sidebar",
    "Claude usage tracking built from ~/.claude/projects/**/*.jsonl — Today $334, 30d $6,151, ~53.2% what-if savings",
    "Three tiers on /usage: actual spend + local-LLM offload + batch-API scheduling, with by-project / by-model / savings-by-model tables and totals",
    "Activity page shows 32 sessions over 7d plus hive_reports log",
    "Attribution fixes: Gmail_Beyond, QI_Universal, C:\\Users\\* paths now map correctly",
    "Autonomous service control: sc.exe resolver added to broker whitelist; qi_service.py wraps sc for stop/start/restart/status",
    "Orphan-broker race diagnosed, stale PID killed, single-instance lock added to qi_elevate.py",
    "FileHQ path corrected in status.json to C:\\NAYA\\filehq (retired)",
    "Committed c98ce9a and pushed to origin/master",
]:
    d.add_paragraph(item, style="List Bullet")

d.add_heading("NEXT UP (Immediate Priority)", level=2)
for item in [
    "Renne to run `nssm start QI_Elevate` once on return — broker stopped after circular self-restart attempt",
    "Verify single-instance lock activates (check C:\\QIH\\logs\\elevation\\broker.lock after restart)",
    "Add broker auto-recovery / watchdog to eliminate circular-restart dead-end",
    "Extend hive_report.report() with agent/model/role fields for richer /activity data",
    "Start populating /tasks with current in-flight work across projects",
]:
    d.add_paragraph(item, style="List Number")

d.add_heading("IN DEVELOPMENT", level=2)
for item in [
    "QI Hive orchestration layer — dashboard observability is now mature; next is active delegation",
    "Yubin v2 multi-account Gmail agent (Maia project) — schema + write-guard middleware pending",
]:
    d.add_paragraph(item, style="List Bullet")

d.add_heading("FUTURE ENHANCEMENTS", level=2)
for item in [
    "Per-agent scorecard on /usage (who ran what, cost per agent)",
    "Local LLM execution path — actually route offloadable tasks to Ollama instead of just estimating",
    "Batch API scheduler — shift overnight work to batch pricing automatically",
]:
    d.add_paragraph(item, style="List Bullet")

d.add_heading("DOCUMENTS UPDATED", level=2)
for f in [
    "C:\\QIH\\engine\\hive\\dashboard\\server.py — /usage, /activity, sidebar legend, color map",
    "C:\\QIH\\engine\\common\\usage_stats.py — NEW (~340 lines): JSONL parser + pricing + what-if models",
    "C:\\QIH\\engine\\common\\qi_service.py — NEW: sc.exe service control wrapper",
    "C:\\QIH\\engine\\common\\qi_elevate.py — sc resolver + single-instance lock",
    "C:\\QIH\\commands\\whitelist.json — sc_service_control rule",
    "C:\\QIH\\data\\status.json — project statuses + FileHQ path",
    "C:\\QIH\\.gitignore — ignore commands/{pending,completed,archive} + shared/reports",
    "C:\\QIH\\docs\\LATEST.md — session note with broker restart instructions",
]:
    d.add_paragraph(f, style="List Bullet")

d.save(str(OUT))
print(f"Saved: {OUT}")
