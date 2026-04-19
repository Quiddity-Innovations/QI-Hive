# -*- coding: utf-8 -*-
"""Generate QIHive Session 02 summary .docx in the shared QI location."""
import sys, os
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

sys.stdout.reconfigure(encoding="utf-8")

OUT_DIR = Path(r"C:\UNIVERSAL\DOCUMENTATION\Session_Summaries")
OUT_DIR.mkdir(parents=True, exist_ok=True)
stamp = datetime.now().strftime("%Y-%m-%d_%H%M")
out_path = OUT_DIR / f"QIHive_Summary_{stamp}.docx"

doc = Document()

# Title
t = doc.add_heading("QI Hive — Session 02 Summary", level=0)
doc.add_paragraph(f"Date: 2026-04-19    |    Status: READY FOR FINAL ADMIN STEP")

# Completed
doc.add_heading("✅ Completed This Session", level=1)
items = [
    "Defined the QI Project Standard — canonical 7-folder layout for every QI project",
    "Reorganized C:\\QIH into engine/, config/, data/, logs/, tests/, docs/, tools/, ecosystem/, shared/",
    "Moved brain code to engine/brain/ (api.py, mcp.py, core/, bootstrap.py, feature_engine.py, tools/)",
    "Moved dashboard code to engine/hive/dashboard/ (server.py, qi_brain_client.py, static/)",
    "Moved qi_brain.db to data/ (DB_PATH now configurable via QI_BRAIN_DB env var)",
    "Built engine/common/qi_logger.py — centralized logging factory with runtime level control",
    "Created config/logging.json — per-service log levels (brain_api, dashboard, etc.)",
    "Added /config dashboard page — drop-down per service, persists to logging.json, hot-applies",
    "Mirrored C:\\UNIVERSAL\\ECOSYSTEM → C:\\QIH\\ecosystem (registry, standards, validators)",
    "Distributed project-local nssm.exe to C:\\QIH\\engine\\bin\\ (one per project, not shared)",
    "Wrote C:\\QIH\\docs\\QI_Project_Standard.md — 10-section canonical spec",
    "Wrote C:\\QIH\\README.md — project overview",
    "Generated C:\\QIH\\tools\\finalize_migration.bat — admin bat to rewire services + remove dupes",
    "Updated QI_Service_Registry.md with new engine/ paths",
    "Identified the real port 8600 orphan — it is the ClaudeManager legacy service, not a stray python",
]
for it in items:
    doc.add_paragraph(it, style="List Bullet")

# Next Up
doc.add_heading("🔄 Next Up (Session 03)", level=1)
next_items = [
    "Run C:\\QIH\\tools\\finalize_migration.bat as admin — removes ClaudeManager + MaiaBot, repoints QI_BrainAPI + QI_Dashboard to engine/ paths",
    "Verify http://localhost:8600/hive and http://localhost:8600/config",
    "Switch services to project-local nssm.exe (remove + reinstall)",
    "Update qi_brain_mcp.py to use C:\\QIH\\data\\qi_brain.db",
    "Build tools/port_audit.py to flag services on wrong ports",
    "Delete C:\\UNIVERSAL\\qi_brain and archive C:\\CLAUDE",
    "Begin Maia migration to C:\\QIP\\Maia under QI Project Standard",
]
for it in next_items:
    doc.add_paragraph(it, style="List Bullet")

# In Development
doc.add_heading("🚀 In Development", level=1)
for it in [
    "Agent growth loop — first real agent growth log entry (endpoints ready, awaiting real task)",
    "Growth MCP tools for qi_brain_mcp.py so Claude can call qi.log_growth() directly",
    "qi_validator extension to check 7-folder standard compliance",
]:
    doc.add_paragraph(it, style="List Bullet")

# Future
doc.add_heading("🌅 Future Enhancements", level=1)
for it in [
    "Project migration pipeline — move Maia, NEXUS, Naya, OC, EasyFlow, FileHQ one at a time to C:\\QIP",
    "Rename C:\\QI → C:\\QIB (QI Business) after all projects moved",
    "Bee Dance — cross-agent pattern sharing via qi_brain.features table",
    "Inter-project event bus for hive-wide coordination",
]:
    doc.add_paragraph(it, style="List Bullet")

# Documents Updated
doc.add_heading("📁 Documents Updated / Created", level=1)
for it in [
    "C:\\QIH\\docs\\QI_Project_Standard.md  (NEW — canonical standard)",
    "C:\\QIH\\README.md  (NEW — project overview)",
    "C:\\QIH\\config\\logging.json  (NEW — per-service log levels)",
    "C:\\QIH\\engine\\common\\qi_logger.py  (NEW — logging factory)",
    "C:\\QIH\\engine\\brain\\*  (MOVED from brain/)",
    "C:\\QIH\\engine\\hive\\dashboard\\*  (MOVED from hive/Dashboard/)",
    "C:\\QIH\\engine\\brain\\core\\db.py  (MODIFIED — DB_PATH now points to data/)",
    "C:\\QIH\\engine\\hive\\dashboard\\server.py  (MODIFIED — /config page + qi_logger wired in)",
    "C:\\QIH\\data\\qi_brain.db  (MOVED from brain/)",
    "C:\\QIH\\data\\status.json  (MOVED from hive/)",
    "C:\\QIH\\ecosystem\\*  (MIRRORED from C:\\UNIVERSAL\\ECOSYSTEM)",
    "C:\\QIH\\engine\\bin\\nssm.exe  (NEW — project-local NSSM binary)",
    "C:\\QIH\\tools\\finalize_migration.bat  (NEW — admin service migration bat)",
    "C:\\QIH\\shared\\sessions\\LATEST.md  (NEW location — was hive/Session Summaries/)",
    "C:\\UNIVERSAL\\ECOSYSTEM\\QI_Service_Registry.md  (MODIFIED — new engine/ paths)",
]:
    doc.add_paragraph(it, style="List Bullet")

doc.save(out_path)
print(f"Saved: {out_path}")
