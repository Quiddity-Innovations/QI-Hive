# -*- coding: utf-8 -*-
"""Generate the session summary .docx for the usage-ledger recovery work."""
from __future__ import annotations

import sqlite3
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT_DIR = Path(r"C:\QIH\shared\documentation\session_summaries")
DB = r"C:\QIH\data\qi_brain.db"


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8")
    con = sqlite3.connect(DB)
    rows = con.execute(
        """SELECT substr(day,1,7), ROUND(SUM(cost_usd),2), SUM(tokens), SUM(turns),
                  ROUND(100.0*SUM(CASE WHEN source='measured' THEN cost_usd ELSE 0 END)
                        /NULLIF(SUM(cost_usd),0),0)
             FROM usage_daily GROUP BY 1 ORDER BY 1"""
    ).fetchall()
    total = con.execute("SELECT ROUND(SUM(cost_usd),2) FROM usage_daily").fetchone()[0]
    con.close()

    d = Document()
    d.add_heading("QI Hive — LLM Usage History: Loss Diagnosis & Reconstruction", 0)
    p = d.add_paragraph()
    p.add_run("Date: ").bold = True
    p.add_run("2026-08-05    ")
    p.add_run("Project: ").bold = True
    p.add_run("QI Hive (QIH)")

    d.add_heading("Root Cause", 1)
    for t in [
        "No counter was reset — there was never a stored counter. usage_stats.py is "
        "stateless and re-parses ~/.claude/projects/**/*.jsonl on every page load.",
        "Claude Code deletes transcripts older than cleanupPeriodDays. That key was "
        "unset, so the 30-day default applied and history decayed continuously.",
        "A cleanup pass ran 2026-08-05T04:31:23Z (~/.claude/.last-cleanup), leaving "
        "only 2026-06-26 onward on disk — this is the drop that was noticed.",
        "'Year to date' was never year-to-date: on the 2026-06-19 screenshot YTD "
        "($19,554) equalled Q2-to-date, and 88% of it fell inside the trailing 30 days.",
        "The figures are Anthropic API list-price equivalents, not money paid. The "
        "plan is Claude MAX 5x (flat subscription), so no billing record was lost.",
    ]:
        d.add_paragraph(t, style="List Bullet")

    d.add_heading("Recovery Attempted", 1)
    t = d.add_table(rows=1, cols=2)
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(("Source", "Result")):
        t.rows[0].cells[i].text = h
    for src, res in [
        ("Volume Shadow Copies", "None exist"),
        ("System Restore points", "None exist"),
        ("Recycle Bin", "Not present"),
        ("OneDrive / backups of ~/.claude", "None"),
        ("qi_brain.db session_log", "No token or cost columns"),
        ("QIH git history", "Never tracked a usage snapshot"),
        ("Dashboard screenshot 2026-06-19", "RECOVERED — full 30d totals + breakdowns"),
        ("Hive report archive (21,596 files)", "RECOVERED — 240 sessions, 198 deleted"),
    ]:
        c = t.add_row().cells
        c[0].text, c[1].text = src, res

    d.add_heading("Reconstructed Calendar (2026)", 1)
    d.add_paragraph(
        "Unit rates were derived from two independent measured windows that agree to "
        "within ~5% ($0.7301/turn vs $0.6968/turn; 11,049 vs 10,673 tokens/turn). "
        "Estimates are compressed and capped at measured peaks so bursty proxies "
        "(April's 348-commit folder reorganisation) cannot inflate them."
    )
    t2 = d.add_table(rows=1, cols=5)
    t2.style = "Light Grid Accent 1"
    for i, h in enumerate(("Month", "API-equiv cost", "Tokens", "Turns", "Basis")):
        t2.rows[0].cells[i].text = h
    for m, cost, tok, turns, pct in rows:
        c = t2.add_row().cells
        c[0].text = m
        c[1].text = f"${cost:,.2f}"
        c[2].text = f"{tok:,}"
        c[3].text = f"{turns:,}"
        c[4].text = "no account" if cost == 0 else f"{int(pct or 0)}% measured"
    c = t2.add_row().cells
    c[0].text, c[1].text = "TOTAL", f"${total:,.2f}"

    d.add_heading("Fixes Shipped", 1)
    for t in [
        "cleanupPeriodDays set to 3650 in ~/.claude/settings.json — transcripts no "
        "longer auto-delete.",
        "usage_ledger.py — durable per-day store (usage_daily table in qi_brain.db) "
        "with source/confidence provenance on every row. Measured rows can never be "
        "overwritten by estimates.",
        "usage_reconstruct.py / usage_backfill.py — activity-proxy model and one-shot "
        "historical backfill (2026-01-01 onward).",
        "usage_snapshot_task.py wired into the Claude Code SessionEnd hook — every "
        "session now persists its day while the transcript still exists.",
        "server.py — QTD/YTD tiles read the ledger (with live-parse fallback) and "
        "display a '% measured' provenance badge.",
        "101 MB of surviving transcripts archived to "
        "C:\\QIH\\data\\usage_archive\\jsonl_snapshot_2026-08-05\\.",
    ]:
        d.add_paragraph(t, style="List Bullet")

    d.add_heading("Caveats", 1)
    for t in [
        "83% of the YTD figure is reconstructed, not measured. It is an informed "
        "model, not a record — treat it as an order-of-magnitude account.",
        "Jan 1 – Feb 17 is a hard zero: the subscription was created "
        "2026-02-17T22:56:38Z and Claude Code first ran 2026-02-18T01:50:57Z.",
        "Feb 18 – May 20 has no surviving measurement of any kind and rests entirely "
        "on activity proxies.",
        "Anthropic's own usage records remain the only authoritative history and were "
        "not consulted (requires selecting a connected browser).",
    ]:
        d.add_paragraph(t, style="List Bullet")

    d.add_heading("Next Up", 1)
    for t in [
        "Optionally cross-check against Anthropic's usage page in a logged-in browser.",
        "Consider backfilling per-project / per-model splits from the 2026-06-19 "
        "screenshot tables (currently held in usage_reconstruct.ANCHOR but not yet "
        "written to a per-project ledger).",
        "Add a ledger panel to the LLM Usage tab showing measured vs reconstructed days.",
    ]:
        d.add_paragraph(t, style="List Bullet")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUT_DIR / f"QIHive_Summary_{datetime.now():%Y-%m-%d_%H%M}.docx"
    d.save(out)
    print(out)


if __name__ == "__main__":
    main()
