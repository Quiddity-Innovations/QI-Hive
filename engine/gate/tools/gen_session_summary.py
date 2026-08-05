# -*- coding: utf-8 -*-
"""Generate the QI Gate session summary .docx (per standing CLAUDE.md rule)."""

import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding="utf-8")

OUT_DIR = Path(r"C:\QIH\shared\documentation\session_summaries")
OUT_DIR.mkdir(parents=True, exist_ok=True)
STAMP = datetime.now().strftime("%Y-%m-%d_%H%M")
OUT = OUT_DIR / f"QIHive_Summary_{STAMP}.docx"

d = Document()
for s, sz in (("Normal", 10.5),):
    d.styles[s].font.size = Pt(sz)
    d.styles[s].font.name = "Segoe UI"

t = d.add_heading("QI Hive — Emergency Security Hardening", 0)
sub = d.add_paragraph("Authentication for every internet-exposed application")
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
sub.runs[0].italic = True
p = d.add_paragraph()
r = p.add_run(f"Session date: 2026-08-05  ·  Generated {datetime.now():%Y-%m-%d %H:%M}")
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def h(text, lvl=1):
    d.add_heading(text, lvl)


def bullets(items):
    for i in items:
        d.add_paragraph(i, style="List Bullet")


def table(headers, rows):
    tb = d.add_table(rows=1, cols=len(headers))
    tb.style = "Light Grid Accent 1"
    for i, hd in enumerate(headers):
        c = tb.rows[0].cells[i]
        c.text = hd
        for run in c.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = tb.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    d.add_paragraph()


# ── The problem ──────────────────────────────────────────────────────────────
h("The problem found")
d.add_paragraph(
    "Renne reported a suspicion that someone was accessing his internet-exposed "
    "applications. An audit confirmed the estate published 22 hostnames through 16 "
    "Cloudflare tunnels, and every single tunnel pointed straight at an application "
    "port with no authentication of any kind.")
bullets([
    "hive.quiddityinnovations.com exposed the ecosystem snapshot, service inventory, "
    "QI Brain data, session history and the War Room to any anonymous caller.",
    "maia.quiddityinnovations.com published /panel, /memory, /history, /export and "
    "/admin/* — conversation history, contact data and a full export endpoint — "
    "on the same host that receives LINE traffic.",
    "Tunnel logs showed 1,580 hostile probes across 783 distinct attack paths "
    "(wp-admin, .env, xmlrpc, webshell drops) in the first days of August.",
    "The QI-Hive GitHub repository is PUBLIC and publishes the full hostname "
    "inventory — attackers did not need to guess anything.",
])

d.add_paragraph()
pw = d.add_paragraph()
pwr = pw.add_run(
    "Important limitation: we cannot prove whether anyone got in. Cloudflare logs "
    "only FAILED origin connections; successful requests to running apps were never "
    "logged anywhere. That absence of evidence is itself the key finding — and the "
    "new audit log is the fix.")
pwr.bold = True

# ── What was built ───────────────────────────────────────────────────────────
h("What was built — QI Gate")
d.add_paragraph(
    "A single login wall in front of the whole estate. No application code was "
    "modified, which is what made a 22-host rollout safe in one session.")
d.add_paragraph("internet → Cloudflare tunnel → Caddy :9040 → app port")
d.add_paragraph("                                    └ forward_auth → QI Gate :9041")
bullets([
    "Caddy does the proxying — websockets, SSE and Gradio queues pass through "
    "untouched. A Python proxy would have broken every Gradio UI.",
    "QI Gate answers one question: is this caller signed in? It serves the login "
    "page and writes the audit log.",
    "Policy is data: config/gate.json is the single source of truth and generates "
    "the Caddy config, so the two cannot drift.",
])

h("Security properties", 2)
table(["Control", "Implementation"], [
    ["Passwords", "pbkdf2-sha256, 200,000 iterations, per-user salt"],
    ["Sessions", "256-bit opaque tokens in SQLite; 12h, or 30d with 'remember'"],
    ["Cookie", "HttpOnly, Secure, SameSite=Lax, scoped .quiddityinnovations.com"],
    ["Brute force", "5 failures per user+IP → 15 min lockout"],
    ["User enumeration", "absent users burn identical hashing time"],
    ["Open redirect", "?rd= validated against known hosts, forced to https"],
    ["First-run setup", "LAN-only, or a single-use token that self-deletes"],
    ["Audit", "every allow/deny/login as JSON with real client IP + country"],
])

h("Failure behaviour (verified)", 2)
table(["Situation", "Result"], [
    ["QI_Gate stopped", "Protected hosts return 502 — fails CLOSED, never open"],
    ["QI_Gate stopped", "Webhooks keep working — LINE/Telegram bots survive"],
    ["QI_Gate restarted", "Wall returns; sessions survive (SQLite, not memory)"],
    ["Both services", "Set to Automatic start"],
])

# ── Completed ────────────────────────────────────────────────────────────────
h("Completed this session")
bullets([
    "Audited all 22 public hostnames and documented the true exposure — the "
    "ecosystem map had been wrong, listing publicly-tunnelled apps as 'LAN only'.",
    "Built QI Gate (identity service, login UI, audit log, admin CLI).",
    "Installed QI_Gate as an auto-start NSSM service via the QI_Elevate broker.",
    "Repointed 13 tunnels (18 hostnames) through the gate; backed up originals.",
    "Verified all 18 hosts gate anonymous callers and open for a signed-in user.",
    "Confirmed LINE/Telegram webhooks still reach Maia and Naya.",
    "Fixed a mixed-content regression that broke NEXUS through the tunnel.",
    "Browser-verified all 11 running apps render correctly through the gate.",
    "Wrote the runbook, the audit record, and updated the ecosystem registries.",
    "Committed locally — deliberately NOT pushed, because the repo is public.",
])

h("Regression found and fixed — NEXUS", 2)
d.add_paragraph(
    "Renne reported mid-rollout that NEXUS rendered unstyled through the tunnel with "
    "'Connection to the server was lost'. Cause: the cloudflared→Caddy hop is plain "
    "HTTP, so Caddy told Gradio the client was on http. Gradio then emitted "
    "http:// asset URLs on an https:// page, which browsers block as mixed content — "
    "killing both the CSS preload and the heartbeat stream. Fixed by forcing "
    "X-Forwarded-Proto: https on every upstream. This was latent for all four Gradio "
    "apps. Note an HTTP-level check would NOT have caught it — the HTML was "
    "byte-identical; it only appeared at runtime in a browser.")

# ── Next up ──────────────────────────────────────────────────────────────────
h("Next up — needs Renne")
table(["Priority", "Action"], [
    ["HIGH", "Decide: make QI-Hive repo private (recommended), or scrub topology "
             "from it. The commit is held locally until then."],
    ["HIGH", "Rotate the gate password — it was set over a chat session. "
             "gate_admin.py passwd Admin <new>"],
    ["HIGH", "Enumerate routes for Claude Voice, OpenClaw gateway and MQ API so "
             "they can move from 'open' to 'mixed'."],
    ["MED",  "Review 'gate_admin.py suspects 24' for a few days — first real "
             "visibility into who is knocking."],
    ["MED",  "Add Cloudflare Access in front of the tunnels for MFA + edge blocking."],
    ["MED",  "Rotate any credential that was reachable from an exposed endpoint."],
    ["LOW",  "Add QI_Gate/QI_Caddy to the Hive health check and monthly self-audit."],
    ["LOW",  "Add a password-change screen so rotation doesn't need the CLI."],
])

h("Still open / known gaps")
bullets([
    "Four hosts remain unauthenticated: connector (permanent, has its own bearer "
    "auth), plus claudevoice, oc-line and api.quiddam — left untouched on purpose "
    "because their routes were not enumerated and breaking a live bot overnight was "
    "the worse risk. They are no worse than before, but not yet fixed.",
    "Single factor — one password now fronts the estate.",
    "Caddy is now a single point of failure for 18 hostnames.",
    "M2V, AutoPDF, quiddam.com and dev.quiddam.com return 502 because those apps "
    "were already stopped before this work began — not a regression.",
])

h("Documents created or updated")
table(["File", "What"], [
    [r"C:\QIH\engine\gate\ (10 files)", "The gate: service, auth, policy, tools"],
    [r"C:\QIH\engine\gate\README.md", "Runbook + troubleshooting"],
    [r"...\documentation\security\QI_Public_Exposure_Hardening_2026-08-05.md",
     "Full audit record — evidence, before/after, verification"],
    [r"C:\QIH\ecosystem\QI_Service_Registry.md", "QI Gate section + symptom lookup"],
    [r"C:\QIH\ecosystem\qi_registry.json", "Registered gate + ports 9040/9041"],
    [r"C:\QIH\ecosystem\QI_Ecosystem_Map.md", "Corrected the wrong 'Internet' column"],
    [r"C:\QIH\engine\proxy\Caddyfile", "Imports the generated gate config"],
    ["13 tunnel configs", "Repointed to :9040 (originals backed up)"],
])

d.save(OUT)
print(f"[OK] {OUT}")
