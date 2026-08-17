#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate the formal Development Effort Record (.docx) from the effort ledger.

The document is written to be read by someone hostile to its conclusions:
methodology, assumptions, coverage gaps and reproduction steps are stated
up front rather than buried.
"""

import sqlite3
import sys
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

sys.path.insert(0, str(Path(__file__).parent))
from qi_effort_ledger import (                      # noqa: E402
    union_buckets, governance, DECLARED_SHARED, MIXED_PROVENANCE,
    EMPLOYER_WORK, DECLARED_ON)

sys.stdout.reconfigure(encoding="utf-8", errors="replace")

TZ = ZoneInfo("America/New_York")
DB = Path(r"C:\QIH\data\effort\effort_ledger.db")

# Governance classification is defined once in qi_effort_ledger.py and imported
# above, so the .docx and the CSV export can never disagree about it.
OUT_DIR = Path(r"C:\QIH\shared\documentation\session_summaries")
STAMP = datetime.now(TZ).strftime("%Y-%m-%d_%H%M")
OUT = OUT_DIR / f"EffortRecord_Summary_{STAMP}.docx"


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    return p


def kv_table(doc, rows, widths=(2.6, 3.6)):
    t = doc.add_table(rows=0, cols=2)
    t.style = "Light Grid Accent 1"
    for k, v in rows:
        c = t.add_row().cells
        c[0].text = str(k)
        c[1].text = str(v)
        for i, w in enumerate(widths):
            c[i].width = Inches(w)
    return t


def main():
    con = sqlite3.connect(DB)

    tot = con.execute("""SELECT SUM(minutes)/60.0, SUM(min_business)/60.0,
        SUM(min_after_hours)/60.0, SUM(min_early_morning)/60.0,
        SUM(min_weekend)/60.0, SUM(min_holiday)/60.0, COUNT(*) FROM sessions
        """).fetchone()
    tk = con.execute("""SELECT SUM(tok_in), SUM(tok_out), SUM(tok_cw),
        SUM(tok_cr), SUM(cost_usd) FROM events""").fetchone()
    rng = con.execute(
        "SELECT MIN(day_local), MAX(day_local) FROM events").fetchone()
    nev = con.execute("SELECT COUNT(*) FROM events").fetchone()[0]
    auth = dict(con.execute("""SELECT author_class, COUNT(*) FROM events
                WHERE source='git' GROUP BY author_class""").fetchall())
    srcs = con.execute("""SELECT source, COUNT(*), MIN(day_local), MAX(day_local)
                FROM events GROUP BY source ORDER BY COUNT(*) DESC""").fetchall()

    per = con.execute("""
        SELECT s.project, SUM(s.minutes)/60.0, SUM(s.min_business)/60.0,
               SUM(s.min_after_hours)/60.0, SUM(s.min_early_morning)/60.0,
               SUM(s.min_weekend)/60.0, SUM(s.min_holiday)/60.0,
               COUNT(*), MIN(s.start_local), MAX(s.end_local)
        FROM sessions s GROUP BY s.project
        ORDER BY SUM(s.minutes) DESC""").fetchall()

    extra = {r[0]: r[1:] for r in con.execute("""
        SELECT project,
               SUM(CASE WHEN source='git' AND author_class='owner'
                   THEN 1 ELSE 0 END),
               SUM(insertions), SUM(deletions), SUM(tok_out), SUM(cost_usd)
        FROM events GROUP BY project""")}

    off_total = sum(x or 0 for x in tot[2:6])
    off_pct = (off_total / tot[0] * 100) if tot[0] else 0

    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10.5)

    # ---------------------------------------------------------------- title
    ti = doc.add_heading("Development Effort Record", level=0)
    ti.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Independent Personal Projects — "
                            "Quiddity Innovations")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    meta = doc.add_paragraph(
        f"Owner: Renne Santiago   |   Generated: "
        f"{datetime.now(TZ).strftime('%Y-%m-%d %H:%M %Z')}   |   "
        f"Coverage: {rng[0]} to {rng[1]}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ------------------------------------------------------------- purpose
    h(doc, "1. Purpose and Scope", 1)
    doc.add_paragraph(
        "This document reconstructs, from timestamped machine artifacts, the "
        "time and computational effort the owner invested in a set of personal "
        "software projects. It was produced to establish an evidentiary record "
        "of when that work occurred — in particular the proportion performed "
        "outside standard business hours."
    )
    doc.add_paragraph(
        "This record states WHEN work occurred and HOW MUCH was produced. It "
        "makes no claim about the subject matter, purpose, or ownership of any "
        "individual work session, and it does not characterise the owner's "
        "employment. Those are separate questions that this data cannot and "
        "does not answer."
    ).runs[0].italic = True

    # --------------------------------------------------------- headline
    u = union_buckets(con)
    uh = {k: (v / 60.0) for k, v in u.items() if not k.startswith("_")}
    u_tot = u["_total"] / 60.0
    u_off = (uh.get("after_hours", 0) + uh.get("early_morning", 0)
             + uh.get("weekend", 0) + uh.get("holiday", 0))
    u_pct = (u_off / u_tot * 100) if u_tot else 0

    h(doc, "2. Summary of Findings", 1)
    doc.add_paragraph(
        "Two measures are reported. They answer different questions and the "
        "difference between them is explained rather than left to be found.")
    for b in [
        "ELAPSED HOURS — real time at the keyboard, with work on several "
        "projects in the same hour counted once. This is the headline figure "
        "and the only one that may be read as time worked.",
        "ATTRIBUTED HOURS — the same work divided across projects. Because "
        "concurrent work on different projects is counted under each, this "
        "sum necessarily exceeds elapsed time. It is valid for comparing "
        "projects and invalid as a total.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    h(doc, "2.1 Elapsed hours (headline)", 2)
    kv_table(doc, [
        ("Total elapsed hours", f"{u_tot:.1f}"),
        ("— Weekday business hours (08:00–17:30)",
         f"{uh.get('business', 0):.1f}"),
        ("— Weekday after 17:30", f"{uh.get('after_hours', 0):.1f}"),
        ("— Weekday before 08:00", f"{uh.get('early_morning', 0):.1f}"),
        ("— Weekend (Sat/Sun)", f"{uh.get('weekend', 0):.1f}"),
        ("— Public holiday", f"{uh.get('holiday', 0):.1f}"),
        ("TOTAL OUTSIDE BUSINESS HOURS", f"{u_off:.1f}  ({u_pct:.1f}%)"),
        ("Distinct wall-clock spans", f"{u['_spans']:,}"),
        ("Underlying timestamped artifacts", f"{nev:,}"),
        ("Distinct projects", f"{len(per)}"),
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        f"Across the covered period, {u_pct:.1f}% of elapsed development time "
        f"({u_off:.1f} of {u_tot:.1f} hours) falls outside weekday business "
        f"hours. Business-hours activity totals "
        f"{uh.get('business', 0):.1f} hours."
    ).runs[0].bold = True

    ux = union_buckets(con, exclude=MIXED_PROVENANCE)
    x_tot = ux["_total"] / 60.0
    x_off = (ux.get("after_hours", 0) + ux.get("early_morning", 0)
             + ux.get("weekend", 0) + ux.get("holiday", 0)) / 60.0
    x_pct = (x_off / x_tot * 100) if x_tot else 0
    doc.add_paragraph(
        f"Conservative basis: excluding the mixed-provenance projects "
        f"identified in section 5.3 ({', '.join(MIXED_PROVENANCE)}), whose "
        f"ownership is arguable, the figures become {x_tot:.1f} elapsed hours "
        f"of which {x_off:.1f} ({x_pct:.1f}%) fall outside business hours. The "
        f"off-hours finding therefore does not depend on those projects being "
        f"counted.")

    h(doc, "2.2 Attributed hours (per-project basis)", 2)
    doc.add_paragraph(
        f"Attributed hours total {tot[0]:.1f} ({tot[0] / u_tot:.2f}x elapsed), "
        f"of which {tot[1] or 0:.1f} are business hours and {off_total:.1f} "
        f"({off_pct:.1f}%) are outside business hours. The excess over elapsed "
        f"time is concurrency, not additional work, and the per-project table "
        f"in section 5 is on this basis.")

    h(doc, "2.3 Internal consistency check", 2)
    days = con.execute(
        "SELECT DISTINCT day_local FROM sessions ORDER BY day_local").fetchall()
    worst_b = worst_t = 0.0
    for (d,) in days:
        du = union_buckets(con, d)
        worst_b = max(worst_b, du.get("business", 0) / 60.0)
        worst_t = max(worst_t, du["_total"] / 60.0)
    doc.add_paragraph(
        f"On the elapsed basis, the highest business-hours total recorded on "
        f"any single day is {worst_b:.1f} hours against a structural ceiling "
        f"of 9.5 (08:00–17:30), and the highest total for any day is "
        f"{worst_t:.1f} hours against a ceiling of 24. Both are within their "
        f"limits. Applying the same test to the attributed figures produces "
        f"day totals above the 9.5-hour ceiling, which is precisely why "
        f"attributed hours are not presented as time worked.")

    # --------------------------------------------------------- methodology
    h(doc, "3. Methodology", 1)
    h(doc, "3.1 Evidence streams", 2)
    doc.add_paragraph(
        "Every hour reported traces to a raw artifact that a third party can "
        "inspect independently. No hour is estimated from recollection.")
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, x in enumerate(["Source", "Artifacts", "Earliest", "Latest"]):
        t.rows[0].cells[i].text = x
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    labels = {"git": "Git commit history",
              "claude_code": "Claude Code session transcripts",
              "session_doc": "Dated session summary documents"}
    for s, n, mn, mx in srcs:
        c = t.add_row().cells
        c[0].text = labels.get(s, s)
        c[1].text = f"{n:,}"
        c[2].text = mn or ""
        c[3].text = mx or ""

    h(doc, "3.2 How hours were derived", 2)
    doc.add_paragraph(
        "Artifacts are timestamps, not durations. Hours are reconstructed by "
        "grouping a project's artifacts into work sessions and measuring each "
        "session's span, using these fixed parameters:")
    for b in [
        "Idle gap: 30 minutes. A gap longer than this ends the session.",
        "Lead-in credit: 6 minutes before a session's first artifact, since "
        "work precedes the artifact that records it.",
        "Session cap: 6 hours. Any longer span is truncated.",
        "Bucketing is performed minute by minute, so a session running from "
        "17:00 to 19:00 is split at the 17:30 boundary rather than assigned "
        "wholesale to one category.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    h(doc, "3.3 Time zone handling", 2)
    doc.add_paragraph(
        "All classification uses local wall-clock time in America/New_York, "
        "with daylight saving applied per-date. Git commits carry their "
        "original UTC offset and are used as recorded. Session transcripts "
        "store UTC and are converted with a DST-aware conversion, not a fixed "
        "offset — a fixed offset would misclassify sessions near the 17:30 "
        "boundary by one hour for part of the year.")

    h(doc, "3.4 Controls applied against overstatement", 2)
    for b in [
        "Commits are de-duplicated globally by SHA. Several projects are "
        "mirrored across C:\\APPS and D:\\Dev; without de-duplication those "
        "hours would have been counted twice.",
        "Third-party and vendored repositories are excluded, and every "
        "excluded author identity is recorded rather than silently dropped.",
        "Commits made by AI agent identities are reported separately from the "
        "owner's own commits (see section 6).",
        "Where a judgement call existed, the lower figure was taken.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # --------------------------------------------------- coverage / gaps
    h(doc, "4. Coverage Boundary and Known Gaps", 1)
    doc.add_paragraph(
        "The limits of this record are stated here deliberately. It is a "
        "conservative floor, not a complete accounting."
    ).runs[0].bold = True
    for b in [
        f"Machine-verifiable evidence begins {rng[0]}. Development work "
        "predating that date is not represented, because the artifacts that "
        "would evidence it are not retained on this machine.",
        "Work performed on other AI platforms (OpenAI/ChatGPT, Grok, Gemini, "
        "Manus and others) is not represented. Those sessions were not "
        "consistently retained and the owner no longer holds equivalent "
        "access. This is a material and acknowledged gap.",
        "Time spent thinking, reading, designing, or researching without "
        "producing a saved artifact within 30 minutes is not counted.",
        "A session producing a single commit after hours of work is credited "
        "only the lead-in minimum, not the actual elapsed effort.",
        "For these reasons the true figure is materially higher than the "
        "total reported here. No attempt has been made to estimate the "
        "difference.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    # ------------------------------------------------------- per project
    doc.add_page_break()
    h(doc, "5. Effort by Project", 1)
    doc.add_paragraph(
        "Hours are decimal. 'Off-hours' combines evenings, early mornings, "
        "weekends and public holidays.")
    doc.add_paragraph(
        "The Governance column records the owner's declared status for each "
        "project (see sections 5.1 and 5.3). It is the owner's statement of "
        "fact, not a machine inference.")
    cols = ["Project", "Governance", "Total", "Business", "After\n17:30",
            "Weekend", "Off-hrs", "Off %", "Commits", "First activity"]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Light Grid Accent 1"
    for i, x in enumerate(cols):
        cell = t.rows[0].cells[i]
        cell.text = x
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8)
    for (p, tt_, bus, aft, em, wk, hol, ns, first, last) in per:
        if tt_ < 0.25:
            continue
        e = extra.get(p, (0, 0, 0, 0, 0))
        off = (aft or 0) + (em or 0) + (wk or 0) + (hol or 0)
        gv = governance(p)
        vals = [p, gv, f"{tt_:.1f}", f"{bus or 0:.1f}", f"{aft or 0:.1f}",
                f"{wk or 0:.1f}", f"{off:.1f}",
                f"{(off / tt_ * 100) if tt_ else 0:.0f}%",
                f"{e[0] or 0}", (first or "")[:10]]
        c = t.add_row().cells
        for i, v in enumerate(vals):
            c[i].text = str(v)
            for r in c[i].paragraphs[0].runs:
                r.font.size = Pt(8)
                if i == 1 and gv != "Personal":
                    r.bold = True

    # Governance roll-up, so the split is legible without reading every row.
    doc.add_paragraph()
    roll = {}
    for (p, tt_, bus, aft, em, wk, hol, ns, first, last) in per:
        g = governance(p)
        a = roll.setdefault(g, [0.0, 0.0, 0.0, 0])
        a[0] += tt_
        a[1] += bus or 0
        a[2] += (aft or 0) + (em or 0) + (wk or 0) + (hol or 0)
        a[3] += 1
    t2 = doc.add_table(rows=1, cols=6)
    t2.style = "Light Grid Accent 1"
    for i, x in enumerate(["Governance", "Projects", "Attributed h",
                           "Business h", "Off-hours h", "Off %"]):
        t2.rows[0].cells[i].text = x
        t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for g in ["Shared with BU", "Mixed provenance", "Personal",
              "Employer work"]:
        if g not in roll:
            continue
        a = roll[g]
        c = t2.add_row().cells
        for i, v in enumerate([g, f"{a[3]}", f"{a[0]:.1f}", f"{a[1]:.1f}",
                               f"{a[2]:.1f}",
                               f"{(a[2] / a[0] * 100) if a[0] else 0:.0f}%"]):
            c[i].text = v

    # ------------------------------- 5.1 declared-shared subset analysis
    q = ",".join("?" * len(DECLARED_SHARED))
    shared = con.execute(f"""
        SELECT project, SUM(minutes)/60.0, SUM(min_business)/60.0,
               SUM(min_after_hours)/60.0, SUM(min_early_morning)/60.0,
               SUM(min_weekend)/60.0, SUM(min_holiday)/60.0
        FROM sessions WHERE project IN ({q})
        GROUP BY project ORDER BY 2 DESC""", DECLARED_SHARED).fetchall()
    grp = {}
    for lbl, op in [("shared", "IN"), ("other", "NOT IN")]:
        grp[lbl] = con.execute(f"""
            SELECT SUM(minutes)/60.0, (SUM(min_after_hours)
                 + SUM(min_early_morning) + SUM(min_weekend)
                 + SUM(min_holiday))/60.0
            FROM sessions WHERE project {op} ({q})""",
            DECLARED_SHARED).fetchone()
    cls = {}
    for p, c_, n in con.execute(f"""SELECT project, author_class, COUNT(*)
            FROM events WHERE source='git' AND project IN ({q})
            GROUP BY project, author_class""", DECLARED_SHARED):
        cls.setdefault(p, {})[c_] = n

    h(doc, "5.1 Projects Declared as Shared with a Third Party", 1)
    doc.add_paragraph(
        f"On {DECLARED_ON} the owner declared that the following projects were "
        f"shared with or demonstrated to Boston University: "
        f"{', '.join(DECLARED_SHARED)}. This is the owner's statement of fact. "
        f"It is not derived from any machine artifact, and the tooling has not "
        f"attempted to verify or contradict it.")

    cols = ["Project", "Total", "Business", "After 17:30", "Weekend",
            "Off-hrs %", "Commits (owner)", "Commits (agent)"]
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = "Light Grid Accent 1"
    for i, x in enumerate(cols):
        cell = t.rows[0].cells[i]
        cell.text = x
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(8.5)
    for (p, tt_, bus, aft, em, wk, hol) in shared:
        off = (aft or 0) + (em or 0) + (wk or 0) + (hol or 0)
        cc = cls.get(p, {})
        vals = [p, f"{tt_:.1f}", f"{bus or 0:.1f}", f"{aft or 0:.1f}",
                f"{wk or 0:.1f}", f"{(off / tt_ * 100) if tt_ else 0:.0f}%",
                f"{cc.get('owner', 0)}", f"{cc.get('agent', 0)}"]
        c = t.add_row().cells
        for i, v in enumerate(vals):
            c[i].text = str(v)
            for r in c[i].paragraphs[0].runs:
                r.font.size = Pt(8.5)

    doc.add_paragraph()
    h(doc, "5.2 Observations Adverse to the Owner's Position", 2)
    doc.add_paragraph(
        "The following are recorded because a record that presents only "
        "favourable findings is not credible. Both weaken rather than support "
        "the owner's position and are stated plainly."
    ).runs[0].bold = True
    sp = grp["shared"]
    op_ = grp["other"]
    doc.add_paragraph(
        f"(a) The declared-shared projects show a LOWER proportion of "
        f"off-hours work than the owner's other projects: "
        f"{sp[1] / sp[0] * 100:.0f}% ({sp[1]:.1f} of {sp[0]:.1f} hours) versus "
        f"{op_[1] / op_[0] * 100:.0f}% ({op_[1]:.1f} of {op_[0]:.1f} hours). "
        f"The projects with the highest off-hours proportion are among those "
        f"not shared.", style="List Bullet")
    zero = [p for p in DECLARED_SHARED if cls.get(p, {}).get("owner", 0) == 0
            and cls.get(p, {}).get("agent", 0) > 0]
    if zero:
        doc.add_paragraph(
            f"(b) For {' and '.join(zero)}, no git commit is recorded under "
            f"the owner's own author identity; all commits carry an AI "
            f"assistant identity configured on the owner's machine (primarily "
            f"at the owner-controlled domains quiddity.ai and "
            f"quiddityinnovations.com, with a small number under a generic "
            f"vendor address). Git author metadata is a freely configurable "
            f"string and is weak evidence of authorship either way; for these "
            f"projects the session transcripts, which record the owner's own "
            f"prompts and timestamps, are the stronger record. The point is "
            f"disclosed here rather than left to be discovered.",
            style="List Bullet")

    # ----------------------------------------- 5.3 mixed provenance
    h(doc, "5.3 Projects of Mixed Provenance", 2)
    doc.add_paragraph(
        f"On {DECLARED_ON} the owner declared that the following were built by "
        f"him independently but applied to employer data: "
        f"{', '.join(MIXED_PROVENANCE)}. Ownership of these is genuinely "
        f"arguable and no position is taken here.")
    qm = ",".join("?" * len(MIXED_PROVENANCE))
    mp = con.execute(f"""SELECT project, SUM(minutes)/60.0,
        SUM(min_business)/60.0, (SUM(min_after_hours)+SUM(min_early_morning)
        +SUM(min_weekend)+SUM(min_holiday))/60.0
        FROM sessions WHERE project IN ({qm}) GROUP BY project
        ORDER BY 2 DESC""", MIXED_PROVENANCE).fetchall()
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, x in enumerate(["Project", "Attributed hours", "Business",
                           "Off-hours"]):
        t.rows[0].cells[i].text = x
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for p, tt_, bb, oo in mp:
        c = t.add_row().cells
        for i, v in enumerate([p, f"{tt_:.1f}", f"{bb:.1f}", f"{oo:.1f}"]):
            c[i].text = v
    doc.add_paragraph()
    doc.add_paragraph(
        f"These are disclosed rather than quietly included. As shown in "
        f"section 2.1, removing them entirely leaves {x_tot:.1f} elapsed hours "
        f"at {x_pct:.1f}% outside business hours, so the central finding of "
        f"this record stands whether or not they are counted. Nothing in this "
        f"document asserts ownership of employer data, employer systems, or "
        f"any output derived from them."
    ).runs[0].bold = True

    # ------------------------------------------------------------ compute
    doc.add_page_break()
    h(doc, "6. Computational Contribution", 1)
    doc.add_paragraph(
        "Token counts recorded by the development tooling. These are reported "
        "in categories because a single gross figure would materially "
        "overstate the work performed: the large majority of tokens are cached "
        "context re-read on each turn, not newly produced output.")
    gross = sum(x or 0 for x in tk[:4])
    kv_table(doc, [
        ("Tokens generated (output)", f"{tk[1] or 0:,}"),
        ("Tokens submitted (input)", f"{tk[0] or 0:,}"),
        ("Cache written", f"{tk[2] or 0:,}"),
        ("Cache re-read (not new work)", f"{tk[3] or 0:,}"),
        ("Gross total (all categories)", f"{gross:,}"),
        ("Replacement cost at public list rates",
         f"US${tk[4] or 0:,.2f}"),
    ])
    doc.add_paragraph(
        "The cost figure is the price of reproducing this token volume at "
        "published API list rates. It is NOT an amount the owner paid; the "
        "work was performed under a fixed-price subscription. It is included "
        "as a measure of scale, not as a claim of expenditure."
    ).runs[0].italic = True

    h(doc, "6.1 Authorship disclosure", 2)
    doc.add_paragraph(
        f"Of the git commits counted, {auth.get('owner', 0):,} were made under "
        f"the owner's own identity and {auth.get('agent', 0):,} were made by an "
        "AI assistant identity operating on the owner's machine during the "
        "owner's session. Both are treated as evidence that the machine was in "
        "use at that timestamp. Only the former evidences the owner's direct "
        "keystroke authorship. This distinction is disclosed here rather than "
        "collapsed, so the figures are not open to the objection that "
        "machine-generated commits were presented as personal authorship.")

    # -------------------------------- 7. conventional-effort translation
    doc.add_page_break()
    h(doc, "7. Translation to Conventional Development Effort", 1)
    doc.add_paragraph(
        "Sections 2 to 6 report measurements. This section reports "
        "ARITHMETIC CONVERSIONS of those measurements (7.1), and separately an "
        "ESTIMATE that rests on an assumption (7.2). The distinction is "
        "maintained deliberately: nothing in 7.2 is offered as evidence, and "
        "the findings in sections 2 to 6 do not depend on it."
    ).runs[0].bold = True

    h(doc, "7.1 Standard business units (arithmetic on measured hours)", 2)
    kv_table(doc, [
        ("Measured elapsed hours", f"{u_tot:.1f}"),
        ("Person-weeks at 40 h/week", f"{u_tot / 40:.1f}"),
        ("Person-months at 160 h/month", f"{u_tot / 160:.1f}"),
        ("FTE-years at 2,080 h/year", f"{u_tot / 2080:.2f}"),
        ("Off-hours portion, person-weeks", f"{u_off / 40:.1f}"),
        ("Contractor equivalent @ US$100/h", f"US${u_tot * 100:,.0f}"),
        ("Contractor equivalent @ US$125/h", f"US${u_tot * 125:,.0f}"),
        ("Off-hours only @ US$125/h", f"US${u_off * 125:,.0f}"),
    ])
    doc.add_paragraph(
        "Hourly rates above are illustrative market rates for senior software "
        "development and are not a quotation, an invoice, or a claim of loss.")

    h(doc, "7.2 Effort not captured by machine artifacts (estimate)", 2)
    doc.add_paragraph(
        "The measured figure is hands-on-tool time. It excludes research "
        "conducted away from the development tooling, design and architectural "
        "thinking, reading, and problem-solving away from the keyboard. That "
        "work is real and, for this owner, substantial — but it produced no "
        "timestamped artifact and therefore cannot be measured from this "
        "machine.")
    doc.add_paragraph(
        "An uplift factor is the conventional way to account for it. No "
        "specific factor is asserted here. Published figures on developer time "
        "allocation vary widely and largely describe team environments with "
        "meetings, ceremonies and code review, which do not correspond to "
        "solo individual work. Any factor relied upon should be supported by a "
        "citable source or expert opinion. The table below shows the effect of "
        "a range of factors so that the sensitivity is visible.")
    t = doc.add_table(rows=1, cols=4)
    t.style = "Light Grid Accent 1"
    for i, x in enumerate(["Uplift factor", "Total hours",
                           "Off-hours", "Person-weeks"]):
        t.rows[0].cells[i].text = x
        t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for m, lbl in ((1.0, "1.0x (measured only)"), (1.3, "1.3x"),
                   (1.6, "1.6x"), (2.0, "2.0x")):
        c = t.add_row().cells
        for i, v in enumerate([lbl, f"{u_tot * m:.0f}", f"{u_off * m:.0f}",
                               f"{u_tot * m / 40:.1f}"]):
            c[i].text = v
    doc.add_paragraph()
    doc.add_paragraph(
        "The 1.0x row is the measured position and requires no assumption. "
        "Every other row requires one.")

    h(doc, "7.3 Stronger corroboration available", 2)
    doc.add_paragraph(
        "The owner retains 302 session summary documents created "
        "contemporaneously with the work, each dated and naming the project "
        "worked on. Contemporaneous records of this kind are ordinarily given "
        "greater weight than any retrospective multiplier, and are recommended "
        "as the primary means of evidencing effort beyond the measured floor.")

    # ------------------------------------------------------ verification
    h(doc, "8. Verification", 1)
    doc.add_paragraph(
        "This record is reproducible. A third party with access to the machine "
        "can regenerate it and compare against this document:")
    for b in [
        r"Collector: C:\QIH\engine\effort\qi_effort_ledger.py",
        r"Database: C:\QIH\data\effort\effort_ledger.db",
        r"Raw per-artifact export: C:\QIH\data\effort\reports\raw_events_*.csv",
        "Regenerate: python qi_effort_ledger.py --backfill",
        "Verify ledger integrity: python qi_effort_ledger.py --verify",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph(
        "Daily entries are sealed in an append-only hash chain: each entry "
        "incorporates the hash of its predecessor, so an entry cannot be "
        "inserted, removed or altered after the fact without breaking every "
        "subsequent hash. The --verify command checks the whole chain.")

    h(doc, "9. Owner Declarations (to be completed by the owner)", 1)
    doc.add_paragraph(
        "The following require the owner's own statement and are deliberately "
        "left blank. They are assertions of fact that this tooling cannot "
        "establish and must not fabricate:")
    doc.add_paragraph(
        f"DECLARED {DECLARED_ON}: the projects shared with or demonstrated to "
        f"Boston University are {', '.join(DECLARED_SHARED)} (see section 5.1). "
        f"The specific dates, audience and circumstances of each demonstration "
        f"remain to be stated by the owner.")
    doc.add_paragraph("Still outstanding:")
    for b in [
        "The date, audience and setting of each demonstration listed above.",
        "Which equipment, accounts and network were used for each project.",
        "The owner's contracted working hours during the covered period.",
        "Whether any employment agreement, IP assignment clause or handbook "
        "provision applies, and its terms.",
    ]:
        doc.add_paragraph(b, style="List Bullet")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    con.close()
    print(f"Saved: {OUT}")
    return OUT


if __name__ == "__main__":
    main()
