#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""One-off: session summary .docx for 2026-08-13 effort-ledger build."""
import sys
from datetime import datetime
from pathlib import Path
from zoneinfo import ZoneInfo
from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8", errors="replace")
TZ = ZoneInfo("America/New_York")
OUT = (Path(r"C:\QIH\shared\documentation\session_summaries") /
       f"EffortLedger_Summary_{datetime.now(TZ).strftime('%Y-%m-%d_%H%M')}.docx")

d = Document()
d.styles["Normal"].font.size = Pt(10.5)
d.add_heading("QI Effort Ledger — Build Session", 0)
d.add_paragraph(f"{datetime.now(TZ).strftime('%Y-%m-%d %H:%M %Z')} · "
                "Project: QI Hive / Claude Manager")

SECTIONS = [
    ("✅ Completed This Session", [
        "Built qi_effort_ledger.py — forensic effort collector over 4 evidence "
        "streams (git, Claude transcripts, session docs, Brain DB).",
        "Backfilled 2026-03-23 → 2026-08-13: 1,061 commits, 55.8k transcript "
        "events, 254 session docs, 849 reconstructed sessions.",
        "Built gen_effort_report.py → 9-section Development Effort Record .docx.",
        "Registered QI_EffortLedger_Daily (23:50 daily, conhost --headless).",
        "Hash-chained append-only ledger; --verify confirms 87 entries intact.",
        "Fixed: global SHA dedup (C:\\APPS vs D:\\Dev mirrors double-counted "
        "~4 projects).",
        "Fixed: concurrency bug — per-project sessions summed to 12.7 business "
        "hours on 2026-07-02 vs a 9.5h ceiling. union_buckets() now merges "
        "overlapping intervals. Corrected total 303.6 → 220.8 elapsed hours.",
        "Fixed: agent-authored commits separated from owner authorship.",
        "Fixed: token reporting split by category (96.4% was cache re-read).",
        "Fixed: project-name normalisation (stray 'repo' and 'unknown' buckets).",
        "Renne set git user.email on C:\\QIH to his own identity.",
    ]),
    ("📊 Baseline Findings", [
        "220.8 elapsed hours · 99.5 business · 121.3 off-hours (54.9%).",
        "36.9M tokens generated; $8,139 replacement cost at list rates.",
        "Owner-declared BU-shared: CogniBase, QI Hive, NEXUS, MapSnap, AutoPDF "
        "— 107.4h attributed, 42% off-hours.",
        "ADVERSE: shared projects are 42% off-hours vs 59% unshared.",
        "ADVERSE: QI Hive (156) and AutoPDF (35) have zero owner-identity "
        "commits — all agent identities.",
        "220.8h = 5.5 person-weeks = ~$27.6k contractor-equivalent @ $125/h.",
    ]),
    ("🔄 Next Up", [
        "Set git user.email on C:\\APPS\\AutoPDF and D:\\Dev\\AutoPDF (BU-shared, "
        "still agent identity); also CLAUDE, PersonalSong, D:\\Dev\\QI\\*.",
        "Renne to complete §9 owner declarations: demo dates/audience per "
        "project, equipment used, contracted hours, IP clause terms.",
        "Optional: provenance timeline showing what existed before BU first "
        "saw each of the five (needs approximate demo dates).",
        "Optional: recover pre-2026-03 history via OpenAI/Google data export.",
        "Consider a short daily attested time log to close the unmeasured gap.",
    ]),
    ("🌅 Future Enhancements", [
        "Cross-check hours against Google Calendar entries.",
        "Filtered SLOC metric (raw 12.1M lines is unusable — vendored files).",
        "Monthly auto-generated .docx snapshot for a rolling sealed record.",
    ]),
    ("⚠️ Decisions Recorded", [
        "Browser history EXCLUDED from the evidentiary record: starts after "
        "git, dwell data unusable (121h/day), and dominated by personal "
        "browsing that would harm rather than help.",
        "No industry uplift multiplier asserted — §7.2 shows a 1.0–2.0x "
        "sensitivity range and requires a citable source before reliance.",
        "Elapsed hours (not attributed) is the only figure presentable as "
        "time worked.",
    ]),
    ("📁 Documents Updated", [
        r"C:\QIH\engine\effort\qi_effort_ledger.py (new)",
        r"C:\QIH\engine\effort\gen_effort_report.py (new)",
        r"C:\QIH\engine\effort\EffortLedger_Daily.bat (new)",
        r"C:\QIH\data\effort\effort_ledger.db (new)",
        r"C:\QIH\data\effort\reports\*.csv (new)",
        r"...\session_summaries\EffortRecord_Summary_2026-08-13_1537.docx (final)",
        r"C:\Users\renne\.claude\projects\C--CLAUDE\memory\project_effort_ledger.md",
    ]),
]
for title, items in SECTIONS:
    d.add_heading(title, 1)
    for i in items:
        d.add_paragraph(i, style="List Bullet")

OUT.parent.mkdir(parents=True, exist_ok=True)
d.save(OUT)
print(f"Saved: {OUT}")
