# -*- coding: utf-8 -*-
"""
QI Brain — Documentation Harvester  (the Librarian's index builder)

Walks every documentation source in the QI ecosystem, catalogs each file into the
`docs` SQL table, derives typed graph edges into `doc_relationships`, and (best
effort) embeds the text into the `qi_docs` Chroma collection for semantic search.

This is the engine behind the "Documentation Brain" — a TheBrain-style Plex where
docs, projects, decisions, features and sessions are nodes connected by links.

Design notes:
  - Files STAY where they live (federated storage). We index, we never move.
  - SQL catalog is ALWAYS written. Embedding is best-effort: if Ollama/nomic is
    down, the row is still cataloged (embedded=0) and picked up on the next run.
  - Idempotent: unchanged files (same content hash, already embedded) are skipped.

Usage:
    python C:\\QIH\\engine\\brain\\doc_harvester.py            # full run + embed
    python C:\\QIH\\engine\\brain\\doc_harvester.py --no-embed # catalog + graph only (fast)
    python C:\\QIH\\engine\\brain\\doc_harvester.py --stats    # print index stats and exit
"""
from __future__ import annotations
import sys, os, re, json, hashlib, argparse, asyncio
from pathlib import Path
from datetime import datetime, timezone

sys.stdout.reconfigure(encoding="utf-8")
_BRAIN_DIR = Path(__file__).parent
sys.path.insert(0, str(_BRAIN_DIR))

from core.db import open_brain_db  # noqa: E402

LOG = Path(r"C:\QIH\logs\doc_harvest.log")
LOG.parent.mkdir(parents=True, exist_ok=True)
REGISTRY = Path(r"C:\QIH\ecosystem\qi_registry.json")
MIGRATION = _BRAIN_DIR / "migrations" / "2026_06_18_doc_index.sql"

# Folders we never descend into
SKIP_DIRS = {".git", "node_modules", "__pycache__", "venv", ".venv", "env",
             "site-packages", "qi_memory", "logs", "LOGS", "worktrees",
             ".claude", "dist", "build", ".next", ".pytest_cache", ".mypy_cache",
             ".ruff_cache", "egg-info"}
# Archive/snapshot folders: their contents are frozen copies of live docs
# (superseded historical snapshots). Indexing them produces duplicate-group
# noise against the live originals, so they are excluded from active search.
# Matched against directory parts only (case-insensitive) — never filenames.
ARCHIVE_DIRS = {"maia_archive", "_legacy_archive", "_archive", "rag_archive"}
# Any directory part containing one of these substrings is also treated as an
# archive (covers dated snapshots like project_library_BACKUP_2026-08-06).
ARCHIVE_SUBSTRINGS = ("_backup", "backup_")
STALE_TYPES = {"implementation_log", "meeting_minutes", "version_history"}
STALE_DAYS = 45


def log(msg: str):
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{ts}] {msg}"
    print(line)
    with open(LOG, "a", encoding="utf-8") as f:
        f.write(line + "\n")


# ── Schema ───────────────────────────────────────────────────────────────────
def ensure_tables(conn):
    """Self-ensure the doc index tables exist (mirrors the migration file)."""
    sql = MIGRATION.read_text(encoding="utf-8")
    for stmt in [s.strip() for s in sql.split(";") if s.strip()]:
        try:
            conn.execute(stmt)
        except Exception as e:  # noqa: BLE001
            log(f"  schema stmt skipped: {type(e).__name__}: {e}")
    conn.commit()


# ── Registry / project mapping ───────────────────────────────────────────────
def load_projects() -> list[dict]:
    """Return [{id, display_name, path_norm}] from qi_registry.json."""
    out = []
    try:
        reg = json.loads(REGISTRY.read_text(encoding="utf-8"))
        projects = reg.get("projects", {})
        items = projects.items() if isinstance(projects, dict) else \
            [(p.get("id"), p) for p in projects]
        for pid, p in items:
            path = (p or {}).get("path") or (p or {}).get("root")
            out.append({
                "id": (pid or (p or {}).get("id") or "").lower(),
                "display_name": (p or {}).get("display_name") or (p or {}).get("name") or pid,
                "path_norm": _norm(path) if path else None,
            })
    except Exception as e:  # noqa: BLE001
        log(f"  registry load failed: {e}")
    return [p for p in out if p["id"]]


def _norm(p) -> str:
    return str(Path(p)).replace("\\", "/").rstrip("/").lower()


def map_project(path: Path, projects: list[dict]) -> str | None:
    """Longest-prefix match of the file path against known project roots."""
    pn = _norm(path)
    best, best_len = None, -1
    for pr in projects:
        root = pr["path_norm"]
        if root and pn.startswith(root + "/") and len(root) > best_len:
            best, best_len = pr["id"], len(root)
    return best


# Session-summary filename prefix → project id (shared store has no path context)
PREFIX_MAP = {
    "qihive": "qihive", "qiorchestrator": "qihive", "claudemanager": "qihive",
    "dashboard": "qihive", "maia": "maia", "naya": "naya", "nexus": "nexus",
    "easyflow": "easyflow", "openclaw": "openclaw", "oc": "openclaw",
    "filehq": "filehq", "tubescout": "tubescout", "personalsong": "personalsong",
    "mapsnap": "cognibase", "cognibase": "cognibase",
}


# ── Classification ───────────────────────────────────────────────────────────
def classify(name: str, path: Path) -> str:
    n = name.lower()
    if "summary" in n and re.search(r"\d{4}-\d{2}-\d{2}", n):
        return "session_summary"
    if "implementation_log" in n or "implementation log" in n:
        return "implementation_log"
    if "meeting_minutes" in n or "meeting minutes" in n:
        return "meeting_minutes"
    if "version_history" in n or "version history" in n:
        return "version_history"
    if n.startswith("readme"):
        return "readme"
    if "changelog" in n:
        return "changelog"
    if "guide" in n:
        return "guide"
    if any(k in n for k in ("standard", "principles", "registry", "ecosystem_map",
                            "service_registry")):
        return "standard"
    if any(k in n for k in ("architecture", "blueprint", "design", "plan")):
        return "architecture"
    return "other"


def prefix_project(name: str) -> str | None:
    m = re.match(r"^([A-Za-z]+)_", name)
    if not m:
        return None
    return PREFIX_MAP.get(m.group(1).lower())


# ── Text extraction ──────────────────────────────────────────────────────────
def read_text(path: Path) -> str:
    suf = path.suffix.lower()
    if suf == ".md":
        try:
            return path.read_text(encoding="utf-8", errors="replace")
        except Exception:  # noqa: BLE001
            return ""
    if suf == ".docx":
        try:
            from docx import Document
            d = Document(str(path))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            for tbl in d.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)
        except Exception as e:  # noqa: BLE001
            log(f"  docx read failed {path.name}: {e}")
            return ""
    return ""


def doc_id_for(path: Path) -> str:
    return hashlib.sha1(_norm(path).encode("utf-8")).hexdigest()[:16]


def title_of(text: str, path: Path) -> str:
    for line in text.splitlines():
        s = line.strip().lstrip("#").strip()
        if s:
            return s[:200]
    return path.stem


# ── Discovery ────────────────────────────────────────────────────────────────
def discover(projects: list[dict]) -> list[Path]:
    roots: list[Path] = [
        Path(r"C:\QIH\ecosystem"),
        Path(r"C:\QIH\docs"),
        Path(r"C:\QIH\shared\documentation"),
    ]
    for pr in projects:
        if pr["path_norm"]:
            roots.append(Path(pr["path_norm"]))
    seen: set[str] = set()
    found: list[Path] = []
    for root in roots:
        if not root.exists():
            continue
        for p in root.rglob("*"):
            if p.is_dir():
                continue
            if p.suffix.lower() not in (".md", ".docx"):
                continue
            if p.name.startswith("~$") or p.name.startswith("."):
                continue  # Word temp-lock files, dotfiles
            if any(part in SKIP_DIRS for part in p.parts):
                continue
            if any(part.lower() in ARCHIVE_DIRS
                   or any(s in part.lower() for s in ARCHIVE_SUBSTRINGS)
                   for part in p.parts):
                continue  # frozen snapshot copies — see ARCHIVE_DIRS
            key = _norm(p)
            if key in seen:
                continue
            seen.add(key)
            found.append(p)
    return found


# ── Relationship extraction ──────────────────────────────────────────────────
def add_edge(conn, src_type, src_id, edge, dst_type, dst_id, source, weight=1.0):
    conn.execute(
        """INSERT OR IGNORE INTO doc_relationships
           (src_type, src_id, edge_type, dst_type, dst_id, weight, source)
           VALUES (?,?,?,?,?,?,?)""",
        (src_type, src_id, edge, dst_type, dst_id, weight, source),
    )


def derive_edges(conn, doc_id, project_id, text, projects):
    # belongs_to: doc -> its project
    if project_id:
        add_edge(conn, "doc", doc_id, "belongs_to", "project", project_id, "path")
    # links_to: [[wikilinks]]
    for target in set(re.findall(r"\[\[([^\]]+)\]\]", text)):
        slug = target.strip().lower().replace(" ", "-")[:80]
        if slug:
            add_edge(conn, "doc", doc_id, "links_to", "doc", slug, "wikilink", 0.8)
    # mentions: other projects named in the body (capped, weighted low)
    low = text.lower()
    mentioned = 0
    for pr in projects:
        if mentioned >= 6:
            break
        if pr["id"] == project_id:
            continue
        dn = (pr["display_name"] or "").lower()
        if len(dn) >= 4 and re.search(r"\b" + re.escape(dn) + r"\b", low):
            add_edge(conn, "doc", doc_id, "mentions", "project", pr["id"], "mention", 0.4)
            mentioned += 1


# ── Prune ────────────────────────────────────────────────────────────────────
def prune(conn, valid_paths: set[str]) -> int:
    """Remove docs that are no longer discoverable — deleted from disk, moved,
    or now under an excluded/archive dir — from the catalog, the relationship
    graph, and the qi_docs vector store. Keeps the index a faithful mirror of
    the live doc tree (no stale duplicates, no phantom paths). The caller guards
    this against running on a broken/empty scan."""
    rows = conn.execute("SELECT doc_id, path FROM docs").fetchall()
    orphans = [(r["doc_id"], r["path"]) for r in rows if r["path"] not in valid_paths]
    if not orphans:
        log("  prune: index already clean (0 orphans)")
        return 0
    ids = [d for d, _ in orphans]
    conn.executemany("DELETE FROM docs WHERE doc_id=?", [(d,) for d in ids])
    conn.executemany(
        "DELETE FROM doc_relationships WHERE src_type='doc' AND src_id=?",
        [(d,) for d in ids])
    conn.commit()
    # Vector store deletion is best-effort: the DB prune is still valid if the
    # embedder/chroma is down — the next run reconciles.
    try:
        from core.memory_store import MemoryStore, COL_DOCS
        store = MemoryStore()
        store.init_collections()
        store._get_col(COL_DOCS).delete(ids=ids)
    except Exception as e:  # noqa: BLE001
        log(f"  prune: vector-store delete skipped: {type(e).__name__}: {e}")
    log(f"  pruned {len(orphans)} de-indexed/orphaned docs")
    for _, p in orphans[:10]:
        log(f"  PRUNE: {p}")
    return len(orphans)


# ── Main harvest ─────────────────────────────────────────────────────────────
def harvest(do_embed: bool):
    log("=== Doc harvest START ===")
    projects = load_projects()
    log(f"  {len(projects)} projects from registry")
    files = discover(projects)
    log(f"  {len(files)} doc files discovered")

    conn = open_brain_db()
    ensure_tables(conn)

    now = datetime.now().isoformat(timespec="seconds")
    existing = {r["path"]: dict(r) for r in conn.execute(
        "SELECT path, content_hash, embedded FROM docs")}
    by_hash: dict[str, list[str]] = {}

    new, changed, unchanged = 0, 0, 0
    to_embed: list[tuple[str, str, dict]] = []

    for p in files:
        text = read_text(p)
        chash = hashlib.sha256(text.encode("utf-8", "replace")).hexdigest()
        by_hash.setdefault(chash, []).append(_norm(p))
        path_str = _norm(p)
        did = doc_id_for(p)
        project_id = map_project(p, projects) or prefix_project(p.name)
        if project_id is None and _norm(p).startswith(_norm(Path(r"C:\QIH"))):
            project_id = "qihive"
        dtype = classify(p.name, p)
        try:
            mtime = datetime.fromtimestamp(p.stat().st_mtime).isoformat(timespec="seconds")
            size = p.stat().st_size
        except Exception:  # noqa: BLE001
            mtime, size = now, 0
        wc = len(text.split())
        title = title_of(text, p)

        prev = existing.get(path_str)
        is_changed = (prev is None) or (prev["content_hash"] != chash)
        # staleness heuristic
        stale, reason = 0, None
        if dtype in STALE_TYPES:
            try:
                age = (datetime.now() - datetime.fromisoformat(mtime)).days
                if age > STALE_DAYS:
                    stale, reason = 1, f"{dtype} untouched {age}d (> {STALE_DAYS}d)"
            except Exception:  # noqa: BLE001
                pass

        conn.execute(
            """INSERT INTO docs
                 (doc_id, path, title, project_id, doc_type, fmt, size_bytes,
                  content_hash, mtime, word_count, embedded, stale, stale_reason,
                  indexed_at, updated_at)
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
               ON CONFLICT(path) DO UPDATE SET
                  title=excluded.title, project_id=excluded.project_id,
                  doc_type=excluded.doc_type, fmt=excluded.fmt,
                  size_bytes=excluded.size_bytes, content_hash=excluded.content_hash,
                  mtime=excluded.mtime, word_count=excluded.word_count,
                  stale=excluded.stale, stale_reason=excluded.stale_reason,
                  updated_at=excluded.updated_at""",
            (did, path_str, title, project_id, dtype, p.suffix.lower().lstrip("."),
             size, chash, mtime, wc,
             1 if (prev and prev["embedded"] and not is_changed) else 0,
             stale, reason, now, now),
        )

        prev_embedded = bool(prev and prev["embedded"])
        if is_changed:
            derive_edges(conn, did, project_id, text, projects)
            new += (prev is None)
            changed += (prev is not None)
        else:
            unchanged += 1
        # Enqueue for embedding when content changed OR it was cataloged but never
        # embedded (e.g. a prior --no-embed run, or the embedder was down).
        if text.strip() and (is_changed or not prev_embedded):
            to_embed.append((did, text, {
                "path": path_str, "title": title,
                "project_id": project_id or "", "doc_type": dtype}))

    conn.commit()

    # Duplicate report (same content, different paths)
    dups = {h: ps for h, ps in by_hash.items() if len(ps) > 1}

    edges = conn.execute("SELECT COUNT(*) FROM doc_relationships").fetchone()[0]
    stale_n = conn.execute("SELECT COUNT(*) FROM docs WHERE stale=1").fetchone()[0]
    log(f"  cataloged: new={new} changed={changed} unchanged={unchanged}")
    log(f"  edges={edges} stale={stale_n} duplicate-groups={len(dups)} to_embed={len(to_embed)}")
    for h, ps in list(dups.items())[:10]:
        log(f"  DUP: {' == '.join(ps)}")

    # Prune docs that vanished from disk or are now excluded (archives, etc.).
    # Guard: only prune after a healthy scan, so a transiently-missing scan root
    # can never wipe the index.
    pruned = 0
    if len(files) >= 100:
        pruned = prune(conn, {_norm(p) for p in files})
    else:
        log(f"  prune skipped — only {len(files)} files discovered (guard tripped)")

    # Embedding pass (best effort)
    embedded_ok = 0
    if do_embed and to_embed:
        try:
            from core.memory_store import MemoryStore
            store = MemoryStore()
            store.init_collections()

            async def run():
                nonlocal embedded_ok
                consecutive = 0
                for did, text, meta in to_embed:
                    ok = False
                    for attempt in range(2):  # one retry — embeds blip occasionally
                        try:
                            # nomic-embed-text has a ~2048-token window; dense prose
                            # overflows ~8k chars and Ollama 500s. Stay well under.
                            await store.add_doc(doc_id=did, text=text[:1800], metadata=meta)
                            ok = True
                            break
                        except Exception as e:  # noqa: BLE001
                            if attempt == 0:
                                await asyncio.sleep(1.0)
                            else:
                                log(f"  embed fail {meta['title'][:40]}: {type(e).__name__}: {e}")
                    if ok:
                        conn.execute("UPDATE docs SET embedded=1 WHERE doc_id=?", (did,))
                        embedded_ok += 1
                        consecutive = 0
                        if embedded_ok % 100 == 0:
                            conn.commit()
                            log(f"  embedded {embedded_ok}/{len(to_embed)}...")
                    else:
                        consecutive += 1
                        if consecutive >= 15:  # embedder genuinely down — stop, resume next run
                            log("  15 consecutive embed failures — aborting pass, will resume next run")
                            break
            asyncio.run(run())
            conn.commit()
        except Exception as e:  # noqa: BLE001
            log(f"  embedding pass skipped: {type(e).__name__}: {e}")
        log(f"  embedded this run: {embedded_ok}/{len(to_embed)}")
    elif not do_embed:
        log("  embedding skipped (--no-embed)")

    conn.close()
    log("=== Doc harvest END ===\n")
    return {"new": new, "changed": changed, "unchanged": unchanged,
            "edges": edges, "stale": stale_n, "duplicates": len(dups),
            "pruned": pruned, "embedded": embedded_ok}


def print_stats():
    conn = open_brain_db()
    ensure_tables(conn)
    total = conn.execute("SELECT COUNT(*) FROM docs").fetchone()[0]
    emb = conn.execute("SELECT COUNT(*) FROM docs WHERE embedded=1").fetchone()[0]
    edges = conn.execute("SELECT COUNT(*) FROM doc_relationships").fetchone()[0]
    print(f"docs={total} embedded={emb} edges={edges}")
    print("\nby type:")
    for r in conn.execute("SELECT doc_type, COUNT(*) c FROM docs GROUP BY doc_type ORDER BY c DESC"):
        print(f"  {r[0]:<20} {r[1]}")
    print("\nby project:")
    for r in conn.execute("SELECT COALESCE(project_id,'(none)') p, COUNT(*) c FROM docs GROUP BY p ORDER BY c DESC"):
        print(f"  {r[0]:<20} {r[1]}")
    stale = conn.execute("SELECT path, stale_reason FROM docs WHERE stale=1 ORDER BY project_id").fetchall()
    if stale:
        print(f"\nstale ({len(stale)}):")
        for r in stale[:30]:
            print(f"  {r[0]}  — {r[1]}")
    conn.close()


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--no-embed", action="store_true", help="catalog + graph only")
    ap.add_argument("--stats", action="store_true", help="print stats and exit")
    args = ap.parse_args()
    if args.stats:
        print_stats()
    else:
        harvest(do_embed=not args.no_embed)
