# -*- coding: utf-8 -*-
"""Generate the QIHive session summary .docx for the Documentation Brain build."""
import sys
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt

sys.stdout.reconfigure(encoding="utf-8")

OUT_DIR = Path(r"C:\QIH\shared\documentation\session_summaries")
OUT_DIR.mkdir(parents=True, exist_ok=True)
stamp = datetime.now().strftime("%Y-%m-%d_%H%M")
out = OUT_DIR / f"QIHive_Summary_{stamp}.docx"

# Pull live numbers from the index
sys.path.insert(0, r"C:\QIH\engine\brain")
from core.db import open_brain_db
c = open_brain_db()
docs = c.execute("SELECT COUNT(*) FROM docs").fetchone()[0]
emb = c.execute("SELECT COUNT(*) FROM docs WHERE embedded=1").fetchone()[0]
edges = c.execute("SELECT COUNT(*) FROM doc_relationships").fetchone()[0]
stale = c.execute("SELECT COUNT(*) FROM docs WHERE stale=1").fetchone()[0]
c.close()

d = Document()
d.add_heading("QI Hive — Session Summary", 0)
p = d.add_paragraph()
p.add_run(f"Documentation Brain & Librarian — {datetime.now():%Y-%m-%d %H:%M}").bold = True

d.add_heading("✅ Completed This Session", 1)
for line in [
    "Evaluated documentation across the whole QI ecosystem (936 docs, ~24 MB, 9 projects).",
    "Designed a 3-layer model: federated storage + central index + Librarian agent (TheBrain-inspired Plex).",
    "Built doc_harvester.py — catalogs every doc, derives a typed knowledge graph, embeds into qi_docs (Chroma).",
    "Added migration 2026_06_18_doc_index.sql — new `docs` and `doc_relationships` tables in qi_brain.db.",
    "Created the hive-librarian sub-agent (find / curate / flag stale / dedupe / enforce compliance).",
    "Wired the harvester into the nightly reconciler as a fold-in step (runs daily).",
    "Enhanced the Guide feature: added PART 12 (Documentation Brain) to QI_Claude_Manager_Guide.md; "
    "fixed stale C:\\UNIVERSAL path in the dashboard guide blurb.",
    "Cleanup: archived 3 stale/duplicate governance docs in C:\\QIH\\docs\\ to _archive\\ with a pointer README.",
    f"First index build: {docs} docs cataloged, {edges} graph edges, {emb} embedded, {stale} stale flagged.",
]:
    d.add_paragraph(line, style="List Bullet")

d.add_heading("🔄 Next Up", 1)
for line in [
    "Add docs/ + standard logs to TubeScout and PersonalSong (flagged non-compliant by the index).",
    "Resolve OpenClaw dual docs/ + Documentation/ folder.",
    "Add a Brain API endpoint + dashboard tile for the doc knowledge graph (the clickable Plex).",
    "Review the 175 duplicate-groups (mostly Maia_Archive copies) and archive canonically.",
    "Have hive-librarian cross-check the 25 stale logs against active projects and dispatch refreshes.",
]:
    d.add_paragraph(line, style="List Bullet")

d.add_heading("🚀 In Development", 1)
for line in [
    "The Plex visual graph view (re-centering node navigation) on the Hive dashboard.",
    "Linking docs to decisions/features/sessions nodes (currently edges cover doc→project + wikilinks + mentions).",
]:
    d.add_paragraph(line, style="List Bullet")

d.add_heading("🌅 Future Enhancements", 1)
for line in [
    "Per-chunk embeddings for long docs (currently one vector per doc, first 8k chars).",
    "Auto-supersede edges (detect 'superseded by' / ARCHIVED to link old→new docs).",
    "Librarian-driven weekly 'documentation health' report to Renne.",
]:
    d.add_paragraph(line, style="List Bullet")

d.add_heading("📁 Documents Updated", 1)
tbl = d.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
tbl.rows[0].cells[0].text = "File"
tbl.rows[0].cells[1].text = "Change"
for f, ch in [
    (r"C:\QIH\engine\brain\doc_harvester.py", "NEW — index builder + graph + embed"),
    (r"C:\QIH\engine\brain\migrations\2026_06_18_doc_index.sql", "NEW — docs + doc_relationships tables"),
    (r"C:\APPS\CLAUDE\.claude\agents\hive-librarian.md", "NEW — Librarian sub-agent"),
    (r"C:\QIH\ecosystem\QI_Claude_Manager_Guide.md", "MOD — added PART 12 (Documentation Brain)"),
    (r"C:\QIH\engine\hive\dashboard\server.py", "MOD — fixed guide path, added Doc Brain note"),
    (r"C:\QIH\tools\nightly_reconcile.py", "MOD — doc-harvest fold-in"),
    (r"C:\QIH\docs\_archive", "NEW — archived 3 stale/duplicate governance docs"),
]:
    r = tbl.add_row().cells
    r[0].text = f
    r[1].text = ch

d.save(str(out))
print(f"SAVED: {out}")
